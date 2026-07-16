import os
import html
import io
import json
import subprocess
import tempfile
import re
import sqlite3
import streamlit as st
import pandas as pd
import plotly.graph_objects as go
from datetime import datetime, date, timedelta
from zoneinfo import ZoneInfo
from urllib.parse import urlsplit, urlunsplit
import logging
import unicodedata
from collections import Counter
from typing import Optional, Tuple, Dict, List
from dataclasses import dataclass
from pathlib import Path
import database as database_module
from database import (
    criar_tabelas, create_user, login_user, get_all_users, delete_user_by_admin,
    add_conta, get_contas, create_shared_report,
    get_shared_report, list_shared_reports, revoke_shared_report,
    verify_admin_password, create_password_reset_code, reset_password_with_code,
    bootstrap_admin_from_env
)

# =========================
# LOGGING CONFIG
# =========================
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)


def _connect_db_compat():
    if hasattr(database_module, "connect_db"):
        return database_module.connect_db()
    db_path = getattr(database_module, "DATABASE", None)
    if not db_path:
        data_dir = os.environ.get("DASHBOARD_DATA_DIR")
        base_dir = Path(data_dir).expanduser() if data_dir else Path(__file__).resolve().parent
        db_path = base_dir / "database.db"
    db_path = Path(db_path).expanduser()
    db_path.parent.mkdir(parents=True, exist_ok=True)
    return sqlite3.connect(str(db_path))


def _garantir_tabela_relatorios_salvos():
    conn = _connect_db_compat()
    cursor = conn.cursor()
    cursor.execute("""
        CREATE TABLE IF NOT EXISTS saved_reports (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            user_id INTEGER NOT NULL,
            client_name TEXT NOT NULL,
            title TEXT NOT NULL,
            payload TEXT NOT NULL,
            summary TEXT,
            period_start DATE,
            period_end DATE,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY(user_id) REFERENCES users(id)
        )
    """)
    conn.commit()
    conn.close()


def _create_saved_report_fallback(
    user_id,
    client_name,
    title,
    payload,
    summary=None,
    period_start=None,
    period_end=None,
):
    _garantir_tabela_relatorios_salvos()
    conn = _connect_db_compat()
    cursor = conn.cursor()
    cursor.execute("""
        INSERT INTO saved_reports (
            user_id, client_name, title, payload, summary, period_start, period_end
        )
        VALUES (?, ?, ?, ?, ?, ?, ?)
    """, (
        user_id,
        client_name,
        title,
        payload,
        summary,
        period_start,
        period_end,
    ))
    conn.commit()
    conn.close()


def _list_saved_reports_fallback(user_id):
    _garantir_tabela_relatorios_salvos()
    conn = _connect_db_compat()
    cursor = conn.cursor()
    cursor.execute("""
        SELECT id, client_name, title, summary, period_start, period_end, created_at
        FROM saved_reports
        WHERE user_id = ?
        ORDER BY created_at DESC
    """, (user_id,))
    reports = cursor.fetchall()
    conn.close()
    return reports


def _get_saved_report_fallback(report_id, user_id):
    _garantir_tabela_relatorios_salvos()
    conn = _connect_db_compat()
    cursor = conn.cursor()
    cursor.execute("""
        SELECT id, client_name, title, payload, summary, period_start, period_end, created_at
        FROM saved_reports
        WHERE id = ? AND user_id = ?
    """, (report_id, user_id))
    report = cursor.fetchone()
    conn.close()
    return report


def _delete_saved_report_fallback(report_id, user_id):
    _garantir_tabela_relatorios_salvos()
    conn = _connect_db_compat()
    cursor = conn.cursor()
    cursor.execute("""
        DELETE FROM saved_reports
        WHERE id = ? AND user_id = ?
    """, (report_id, user_id))
    changed = cursor.rowcount > 0
    conn.commit()
    conn.close()
    return changed


create_saved_report = getattr(
    database_module,
    "create_saved_report",
    _create_saved_report_fallback,
)
list_saved_reports = getattr(
    database_module,
    "list_saved_reports",
    _list_saved_reports_fallback,
)
get_saved_report = getattr(
    database_module,
    "get_saved_report",
    _get_saved_report_fallback,
)
delete_saved_report = getattr(
    database_module,
    "delete_saved_report",
    _delete_saved_report_fallback,
)

# =========================
# CLIENTES
# =========================
CLIENTES_RELATORIO = [
    "Luciana Matoso",
    "Andressa Manfroi",
    "Francini Pereira",
    "Vanessa Secco",
    "Taci Buzato",
    "Patricia Terron",
    "Dani Liranço",
    "Milene Buzato",
    "Anna Sarkis",
    "Amanda Ribeiro",
    "Ana Hirata",
    "Gabriella Souza",
    "Jhulia Padilha",
    "Fernanda Tecchio",
    "Kleicy Abreu",
    "Dr Italo",
    "Patricia Leite",
    "Rafaela Riso",
]

MESES_RELATORIO = [
    ("01", "Janeiro"),
    ("02", "Fevereiro"),
    ("03", "Março"),
    ("04", "Abril"),
    ("05", "Maio"),
    ("06", "Junho"),
    ("07", "Julho"),
    ("08", "Agosto"),
    ("09", "Setembro"),
    ("10", "Outubro"),
    ("11", "Novembro"),
    ("12", "Dezembro"),
]

# =========================
# DATA MODELS
# =========================
@dataclass
class User:
    id: int
    username: str
    is_admin: bool = False

    def initials(self) -> str:
        return self.username[0].upper() if self.username else "?"

# =========================
# CONFIG
# =========================
st.set_page_config(
    page_title="Dashboard Financeiro",
    page_icon="💰",
    layout="wide",
    initial_sidebar_state="expanded"
)

criar_tabelas()
bootstrap_admin_from_env()

# =========================
# CARREGA CSS
# =========================
def load_css():
    try:
        css_path = Path(__file__).resolve().parent / "styles.css"
        with css_path.open("r", encoding="utf-8") as f:
            st.markdown(f"<style>{f.read()}</style>", unsafe_allow_html=True)
    except FileNotFoundError:
        logger.warning("styles.css not found")
        st.markdown("<style>.stApp { background-color: #0D0D1A; }</style>", unsafe_allow_html=True)

load_css()

st.markdown(
    """
    <style>
    .page-title {
        margin: 0 0 .35rem 0 !important;
        color: #f3f5ff !important;
        font-size: clamp(2rem, 3.1vw, 3.15rem) !important;
        font-weight: 850 !important;
        line-height: 1.05 !important;
        letter-spacing: 0 !important;
    }
    .page-subtitle {
        margin-top: .35rem !important;
        margin-bottom: 1.55rem !important;
        color: #a7aac9 !important;
        font-size: clamp(1rem, 1.25vw, 1.22rem) !important;
        font-weight: 650 !important;
        line-height: 1.45 !important;
    }
    .stMainBlockContainer,
    [data-testid="stMainBlockContainer"],
    .block-container {
        max-width: 1680px !important;
        padding: 4.25rem clamp(28px, 4vw, 72px) 3rem !important;
    }
    .kpi-card {
        min-height: 142px;
        padding: 1.25rem 1.3rem;
        background: linear-gradient(145deg, #121426, #0f1120);
        border: 1px solid #272744;
        border-top: 3px solid #4b4d65;
        border-radius: 13px;
        box-shadow: 0 12px 32px rgba(0, 0, 0, .18);
    }
    .kpi-card.green { border-top-color: #2fc792; }
    .kpi-card.neutral { border-top-color: #7b7e98; }
    .kpi-card.blue { border-top-color: #68a9ed; }
    .kpi-card.purple { border-top-color: #7167dc; }
    .kpi-card.red { border-top-color: #f07f91; }
    .kpi-label {
        margin-bottom: .65rem;
        color: #a4a6bc;
        font-size: .82rem;
        font-weight: 700;
        letter-spacing: .055em;
        text-transform: uppercase;
    }
    .kpi-value {
        color: #e7e9f3;
        font-family: "DM Mono", Consolas, monospace;
        font-size: clamp(1.3rem, 1.7vw, 1.75rem);
        font-weight: 700;
        line-height: 1.2;
        overflow-wrap: anywhere;
    }
    .kpi-value.green,
    .valor-pos,
    .variacao-pos { color: #2fc792 !important; }
    .kpi-value.red,
    .valor-neg,
    .variacao-neg { color: #f07f91 !important; }
    .kpi-value.blue { color: #68a9ed !important; }
    .kpi-footer {
        margin-top: .8rem;
        color: #77799d;
        font-size: .86rem;
        line-height: 1.4;
    }
    .comp-row {
        display: flex;
        align-items: center;
        justify-content: space-between;
        gap: 1rem;
        margin: .55rem 0 .35rem;
        color: #b7b9cc;
        font-size: .92rem;
        font-weight: 700;
    }
    .comp-bar-bg {
        width: 100%;
        height: 6px;
        overflow: hidden;
        background: #1b1d2d;
        border-radius: 99px;
        margin-bottom: .82rem;
    }
    .comp-bar-fill {
        height: 100%;
        min-width: 4px;
        border-radius: inherit;
    }
    .fin-table,
    .dre-table {
        width: 100%;
        overflow: hidden;
        color: #cfd1df;
        background: #10121f;
        border: 1px solid #272744;
        border-collapse: separate;
        border-spacing: 0;
        border-radius: 12px;
        font-size: .95rem;
    }
    .fin-table th,
    .dre-table th {
        padding: .85rem 1rem;
        color: #8e91ad;
        background: #15172a;
        border-bottom: 1px solid #272744;
        font-size: .78rem;
        font-weight: 700;
        letter-spacing: .06em;
        text-align: left;
        text-transform: uppercase;
    }
    .fin-table td,
    .dre-table td {
        padding: .8rem 1rem;
        border-bottom: 1px solid #202236;
    }
    .assistant-top-button [data-testid="stButton"] button,
    div[data-testid="stButton"] button[kind="primary"] {
        border-radius: 999px !important;
        background: linear-gradient(135deg, #7167dc, #4f8cff) !important;
        border: 1px solid rgba(160, 165, 255, .38) !important;
        box-shadow: 0 14px 34px rgba(0, 0, 0, .30) !important;
        font-weight: 800 !important;
    }
    [data-testid="stDialog"] {
        background: #0f1120 !important;
    }
    [data-testid="stChatMessage"] {
        background: #111326;
        border: 1px solid #252844;
        border-radius: 10px;
    }
    .machine-delta {
        display: inline-flex;
        align-items: center;
        margin-top: .75rem;
        padding: .28rem .62rem;
        border-radius: 999px;
        font-family: "DM Mono", Consolas, monospace;
        font-size: .78rem;
        font-weight: 800;
    }
    .machine-delta.positive {
        color: #35e09c;
        background: rgba(47, 199, 146, .16);
    }
    .machine-delta.negative {
        color: #ff7185;
        background: rgba(240, 127, 145, .16);
    }
    .machine-delta.neutral {
        color: #a7aac9;
        background: rgba(167, 170, 201, .12);
    }
    </style>
    """,
    unsafe_allow_html=True,
)

# =========================
# HELPERS
# =========================
def fmt_brl(valor: float, sinal: bool = False) -> str:
    if not isinstance(valor, (int, float)):
        return "R$ 0,00"
    prefix = "+R$ " if sinal and valor >= 0 else "R$ "
    return f"{prefix}{valor:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")


def fmt_brl_saida(valor: float) -> str:
    valor_num = pd.to_numeric(pd.Series([valor]), errors="coerce").fillna(0).iloc[0]
    return f"- {fmt_brl(abs(float(valor_num)))}"


def fmt_brl_chat(valor: float, sinal: bool = False) -> str:
    return fmt_brl(valor, sinal=sinal).replace("$", r"\$")


def texto_html(valor, limite: Optional[int] = None) -> str:
    texto = "" if pd.isna(valor) else str(valor)
    if limite is not None:
        texto = texto[:limite]
    return html.escape(texto)


def parse_valor_br(v) -> Optional[float]:
    try:
        v = str(v).replace("R$", "").replace(" ", "").strip()
        if not v:
            return None
        if "," in v:
            v = v.replace(".", "").replace(",", ".")
        return pd.to_numeric(v, errors="coerce")
    except Exception as e:
        logger.error(f"Error parsing valor: {v} - {e}")
        return None

def tag_forma(forma: str) -> str:
    if not isinstance(forma, str):
        forma = ""
    mapa = {
        "transferência": "tag-blue",
        "boleto": "tag-orange",
        "dinheiro": "tag-green",
        "cartão": "tag-red",
        "pix": "tag-blue",
    }
    cls = "tag-gray"
    for k, v in mapa.items():
        if k in forma.lower():
            cls = v
            break
    return f'<span class="tag {cls}">{html.escape(forma)}</span>'

def tag_categoria(cat: str) -> str:
    if not isinstance(cat, str):
        cat = "N/A"
    return f'<span class="tag tag-gray">{html.escape(cat)}</span>'


def classificar_forma_recebimento(row) -> str:
    tipo = normalizar_texto(row.get("tipo_recebimento", ""))
    if tipo:
        mapa_tipo = {
            "pix": "PIX",
            "cartao": "Cartão",
            "antecipacao": "Antecipação",
            "credito bancario": "Banco direto",
            "clinipay": "Clinipay",
            "pagamento direto": "Pagamento direto",
        }
        return mapa_tipo.get(tipo, str(row.get("tipo_recebimento", "")).strip())

    texto = " ".join(
        normalizar_texto(row.get(coluna, ""))
        for coluna in [
            "forma", "conta_destino", "descricao", "memo", "fonte",
            "tipo_transacao", "nome", "detalhe",
        ]
    )
    if "pix" in texto:
        return "PIX"
    if "boleto" in texto:
        return "Boleto"
    if "dinheiro" in texto or "caixa" in texto:
        return "Dinheiro"
    if (
        "maquininha" in texto or "infinite" in texto or "infinity" in texto
        or "rede" in texto or "redecard" in texto or "visa" in texto
        or "master" in texto or "cartao" in texto or "cartão" in texto
    ):
        return "Maquininhas"
    if "banco" in texto or "sicoob" in texto or "sicredi" in texto:
        return "Banco direto"
    return "Outros"


def normalizar_texto(valor) -> str:
    texto = unicodedata.normalize("NFKD", str(valor or ""))
    return "".join(c for c in texto if not unicodedata.combining(c)).lower().strip()


def agrupar_servico_vendido(servico, tipo_item=None) -> str:
    nome = str(servico or "").strip()
    texto = normalizar_texto(nome)
    tipo = normalizar_texto(tipo_item)
    texto_busca = f"{texto} {tipo}".strip()
    if "tirzepa" in texto_busca:
        return "TIRZEPATIDA"
    if tipo == "produto":
        return "PRODUTO"
    if re.search(
        r"\b(?:terapia|terapias|tarapia)\s+injetave(?:l|is)\b|"
        r"injetave(?:l|is)|dupla\s+intramuscular|mesoterapia",
        texto_busca,
    ):
        return "INJETÁVEIS"
    if (
        "programa" in texto_busca
        or "protocolo" in texto_busca
        or re.search(r"\b6\s+meses?\b", texto_busca)
    ):
        return "PROGRAMA"
    if "implante" in texto_busca:
        return "IMPLANTES"
    if re.search(r"\bconsulta\b|avaliacao|avaliação|atendimento", texto_busca):
        return "CONSULTA"
    return "Outros"


def agrupar_procedimento_clinicorp(procedimento) -> str:
    nome = str(procedimento or "").strip()
    texto = normalizar_texto(nome)
    if not texto:
        return "Serviço não identificado"

    grupos = [
        (r"confissao\s+de\s+divida|confissão\s+de\s+dívida", "Outras receitas"),
        (
            r"produto|booster\s+firmador|serum|sérum|creme|protetor|stick|"
            r"sabonete|cicatri|pelicula\s+dermica|película\s+dermica|"
            r"cosmebeauty|skincare|reparador|suplemento|complex\s+suplemento|"
            r"\bcinta\b|detox|funil|capsulas|cápsulas|max\s+suplemento",
            "Produtos",
        ),
        (r"ultralift", "Ultralift"),
        (r"\bmetodo\b|\bmétodo\b", "Método"),
        (r"microagulh", "Microagulhamento"),
        (
            r"design|limpeza\s+de\s+pele|limpeza|demais\s+procedimentos|"
            r"estetica|estética",
            "Estética",
        ),
        (
            r"\bspa\b|massagem|hidratacao|hidratação|hydra|aqua",
            "Procedimentos/Serviços",
        ),
        (r"bioimped", "Bioimpedância"),
        (r"\bblur\b", "Blur"),
        (r"rejuvenize", "Rejuvenize"),
        (r"preench", "Preenchimentos"),
        (r"bioestimul", "Bioestimuladores"),
        (r"\btoxina\b|botox|botulin", "Toxina Botulínica"),
        (r"skin\s*booster|skinbooster", "Skinbooster"),
        (r"laser", "Laser"),
        (r"peeling", "Peelings"),
        (r"fios?\s+de\s+sustentacao|fio\s+pdo|\bpdo\b", "Fios de Sustentação"),
        (r"ultrassom|ultraformer|liftera", "Ultrassom Microfocado"),
        (r"implante", "Implantes"),
        (r"injetave(?:l|is)|intramuscular|mesoterapia", "Injetáveis"),
        (r"programa|protocolo", "Programas"),
        (r"\bconsulta\b|avaliacao|avaliação|atendimento", "Consultas"),
        (r"\bretorno\b", "Retornos"),
    ]
    for padrao, rotulo in grupos:
        if re.search(padrao, texto):
            return rotulo
    return nome


def nome_banco_brasileiro(codigo_ou_nome) -> str:
    banco = str(codigo_ou_nome or "").strip()
    if not banco:
        return "Banco"
    banco_limpo = re.sub(r"\D", "", banco)
    bancos = {
        "001": "Banco do Brasil",
        "033": "Santander",
        "041": "Banrisul",
        "077": "Banco Inter",
        "104": "Caixa Econômica Federal",
        "136": "Unicred",
        "197": "Stone",
        "208": "BTG Pactual",
        "212": "Banco Original",
        "237": "Bradesco",
        "260": "Nubank",
        "290": "PagBank",
        "323": "Mercado Pago",
        "336": "C6 Bank",
        "341": "Itaú",
        "380": "PicPay",
        "422": "Banco Safra",
        "655": "Banco Votorantim",
        "748": "Sicredi",
        "756": "Sicoob",
    }
    return bancos.get(banco_limpo.zfill(3), banco)


def parse_data_flexivel(valores) -> pd.Series:
    """Converte datas em texto e datas seriais do Excel."""
    datas = pd.to_datetime(valores, dayfirst=True, errors="coerce")
    numeros = pd.to_numeric(valores, errors="coerce")
    datas_excel = pd.to_datetime(
        numeros,
        unit="D",
        origin="1899-12-30",
        errors="coerce",
    )
    mascara_excel = numeros.between(20000, 80000)
    return datas.mask(mascara_excel, datas_excel)


def classificar_grupo_custo(df: pd.DataFrame) -> pd.Series:
    """Classifica retirada, custo fixo e custo variável por texto financeiro."""
    if df is None or df.empty:
        return pd.Series(dtype="object")

    categorias = df.get(
        "categoria",
        pd.Series("", index=df.index),
    ).fillna("").map(normalizar_texto)
    descricoes = df.get(
        "descricao",
        pd.Series("", index=df.index),
    ).fillna("").map(normalizar_texto)
    tipos = df.get(
        "tipo",
        pd.Series("Custo Variável", index=df.index),
    ).fillna("Custo Variável").map(normalizar_texto)
    texto = categorias + " " + descricoes + " " + tipos
    categorias_fixas_modelo = [
        "13 salario",
        "agua e saneamento",
        "aluguel",
        "alvara de funcionamento",
        "condominio",
        "crm",
        "csll",
        "cursos e treinamentos",
        "despesa administrativa",
        "descontos incondicionais obtidos",
        "despesa estrutural",
        "despesa operacional",
        "fgts e multa de fgts",
        "fretes pagos",
        "honorarios advocaticios",
        "honorarios consultoria",
        "honorarios contabeis",
        "inss patronal",
        "inss sobre pro-labore",
        "inss sobre salarios",
        "irpj",
        "lanches e refeicoes",
        "manutencao predial",
        "marketing e publicidade",
        "materiais de escritorio",
        "medicina do trabalho",
        "moveis utensilios e instalacoes administrativos",
        "outras despesas nao considerar",
        "pessoal e beneficios",
        "plano de saude colaboradores",
        "salarios",
        "tarifas bancarias",
        "uniformes",
        "vale-transporte",
        "sistemas",
    ]
    categorias_variaveis_modelo = [
        "cofins",
        "comissoes de vendedores",
        "despesas tributarias",
        "devolucao de vendas",
        "energia eletrica",
        "insumos medicamentos",
        "insumos injetaveis",
        "iss sobre faturamento",
        "materiais aplicados na prestacao de servico",
        "materiais e medicamentos",
        "materiais para revenda",
        "outras despesas operacionais",
        "pis",
        "prestacao de servicos",
        "software licenca de uso",
        "taxas de maquininha e boletos",
        "tributario e financeiro",
        "uso e consumo",
        "utensilios acessorios",
        "mentoria",
        "despesas a identificar",
        "holding",
        "cartao de credito",
        "investimentos",
        "gratificacoes",
    ]
    padrao_fixo_modelo = "|".join(
        re.escape(categoria) for categoria in categorias_fixas_modelo
    )
    padrao_variavel_modelo = "|".join(
        re.escape(categoria) for categoria in categorias_variaveis_modelo
    )
    fixo_modelo = categorias.str.contains(padrao_fixo_modelo, regex=True)
    variavel_modelo = categorias.str.contains(padrao_variavel_modelo, regex=True)
    antecipacao_lucro = texto.str.contains(
        r"antecipacao\s+de\s+lucro|antecipação\s+de\s+lucro",
        regex=True,
    )
    pro_labore = texto.str.contains(
        r"pro[\s-]*labore|pró[\s-]*labore",
        regex=True,
    )
    retirada = (
        categorias.str.contains(
        r"retirada\s*(de\s*)?lucro|distribuicao\s*(de\s*)?lucro|"
        r"socio|sócio",
        regex=True,
        )
        | descricoes.str.contains(
        r"retirada\s*(de\s*)?lucro|distribuicao\s*(de\s*)?lucro|"
        r"socio|sócio",
        regex=True,
        )
    )
    grupo_fixo = (
        tipos.str.fullmatch(r".*fixo.*")
        | categorias.str.fullmatch(r".*fixo.*")
    )
    grupo_variavel = (
        tipos.str.fullmatch(r".*variavel.*")
        | categorias.str.fullmatch(r".*variavel.*")
    )
    fixo_por_texto = grupo_fixo | (~grupo_variavel & texto.str.contains(
        r"aluguel|condominio|condomínio|energia|luz|agua|água|telefone|"
        r"internet|honorarios|honorários|contabil|contábil|software|"
        r"sistema|mensalidade|salario|salário|ordenado|fgts|inss|"
        r"imposto|simples nacional|seguro|plano medico|plano médico|"
        r"tarifa bancaria|tarifa bancária|taxa condominial|ponto eletronico|"
        r"ponto eletrônico|marketing|propaganda|publicidade|limpeza|"
        r"outras despesas com pessoal|irrf",
        regex=True,
    ))
    fixo = fixo_modelo | (fixo_por_texto & ~variavel_modelo)

    return pd.Series(
        [
            "Antecipação de Lucro" if a else
            "Pró-labore" if p else
            "Retirada de Lucro" if r else
            "Custo Fixo" if f else
            "Custo Variável"
            for a, p, r, f in zip(antecipacao_lucro, pro_labore, retirada, fixo)
        ],
        index=df.index,
    )


def filtrar_fornecedores_producao(df: Optional[pd.DataFrame]) -> pd.DataFrame:
    """Mantém gastos ligados diretamente à produção clínica."""
    if df is None or df.empty:
        return pd.DataFrame()

    base = df.copy()
    fornecedores_catalogo = {
        "ADYEN LATIN AMERICA": ["adyen latin america"],
        "ALEXANDRE SANTANA / RETATRUTIDE": [
            "alexandre santana",
            "alexandre santana retatrutide",
            "retatrutide",
        ],
        "BALK PRODUTOS HOSPITALARES": ["balk produtos hospitalares"],
        "BIO MEDS PHARMACEUTICA LTDA": [
            "bio meds pharmaceutica",
            "bio meds pharmaceutica ltda",
            "biomeds",
        ],
        "BIOMETIK": ["biometik"],
        "CELLGENIC BRASIL": ["cellgenic brasil", "cellgenic br"],
        "CENTRAL FARMA IPATINGA": [
            "central farma ipatinga",
            "central farmaipatinga",
        ],
        "CLIVIVA CENTRO MEDICO": ["cliviva centro medico"],
        "COMERCIAL RAMOS MATERIAIS": [
            "comercial ramos materiais",
            "comercial ramos materiais cirurgicos",
            "comercial ramos",
        ],
        "DERMEGE": ["dermege"],
        "DROGARIA SÃO PAULO": ["drogaria sao paulo"],
        "DROGASIL": ["drogasil"],
        "DUX COMPANY": ["dux company"],
        "FASTSHOP": ["fastshop", "fast shop"],
        "FLEXSIV": ["flexsiv"],
        "FELIPE FERNANDES": ["felipe fernandes"],
        "GPZ COMERCIAL LTDA": ["gpz comercial", "gpz comercial ltda"],
        "GENESYSMED": ["genesym", "genesysmed"],
        "HEALTH TECH FARMACIA": [
            "health tech farmacia",
            "health tech farmacia de manipu",
            "health tech",
        ],
        "INFINITY SUPLEMENTOS": ["infinity suplementos"],
        "ELMECO": [
            "elmeco",
            "elmeco produ",
            "elmeco produtos medicos e farmaceuticos",
        ],
        "ECWC TECNOLOGIA": ["ecwc tecnologia"],
        "ESSENTIAL": ["essential"],
        "TRIQUE COMERCIO": ["trique comercio"],
        "SUPERMED": ["supermed", "supremed"],
        "STIN PHARMA EXCELENCIA": ["stin pharma excelencia"],
        "QUANTITY SERVIÇOS E COMERCIO": [
            "quantity servicos e comercio",
            "quantity",
        ],
        "INTERNATIONAL SKIN SOLUTION": ["international skin solution"],
        "INSTITUTO DE IMUNOLOGIA": ["instituto de imunologia"],
        "INP INDUSTRIAS DE ALIMENTOS": ["inp industrias de alimentos"],
        "MAGAZINE MEDICA": ["magazine medica"],
        "MEDPHARMA MEDICAMENTOS": ["medpharma medicamentos"],
        "MEICOS": ["meicos"],
        "MERCADO LIVRE": ["mercado livre", "mercado pago"],
        "NITRATUS PHARMA": ["nitratus pharma"],
        "NECK PROFISSIONAIS": ["neck profissionais", "neck"],
        "LA VIE LEGACY LABS PRODUTOS": ["la vie legacy labs produtos"],
        "ONCO PROD DISTRIBUIDO": [
            "onco prod distribuido",
            "onco prod distribuidora",
            "onco prod",
        ],
        "TOSKANI MEDSAN": ["toskani medsan", "toskani"],
        "PHARMACIA ESSENTIA": ["pharmacia essentia"],
    }
    texto = (
        base.get("categoria", pd.Series("", index=base.index)).fillna("").map(normalizar_texto)
        + " "
        + base.get("descricao", pd.Series("", index=base.index)).fillna("").map(normalizar_texto)
        + " "
        + base.get("fornecedor", pd.Series("", index=base.index)).fillna("").map(normalizar_texto)
        + " "
        + base.get("tipo", pd.Series("", index=base.index)).fillna("").map(normalizar_texto)
    )
    fornecedor_canonico = pd.Series("", index=base.index, dtype="object")
    for nome_padrao, aliases in fornecedores_catalogo.items():
        padrao_alias = "|".join(
            re.escape(normalizar_texto(alias))
            for alias in aliases
        )
        if not padrao_alias:
            continue
        mascara_alias = texto.str.contains(padrao_alias, regex=True)
        fornecedor_canonico = fornecedor_canonico.mask(
            (fornecedor_canonico == "") & mascara_alias,
            nome_padrao,
        )
    fornecedor_listado = fornecedor_canonico != ""
    fora_producao = texto.str.contains(
        r"simples\s+nacional|imposto|tributo|salario|salarios|ordenado|"
        r"repasse\s+medico|pro[\s-]*labore|retirada|ajuste\s+de\s+caixa|"
        r"\bitau\b|sicoob|sicredi|banco|tarifa|aluguel|internet|energia|"
        r"telefone|honorario|contabil|software|marketing|"
        r"despesa\s+administrativa|matheus\s+\|\s+mercado\s+pago",
        regex=True,
    )
    base["fornecedor"] = fornecedor_canonico.where(
        fornecedor_canonico != "",
        base.get("fornecedor", pd.Series("", index=base.index)),
    )
    return base[fornecedor_listado & ~fora_producao].copy()


def filtrar_por_periodo(
    df: Optional[pd.DataFrame],
    inicio: pd.Timestamp,
    fim: pd.Timestamp,
) -> pd.DataFrame:
    if df is None or df.empty or "data" not in df.columns:
        return pd.DataFrame()
    datas = pd.to_datetime(df["data"], dayfirst=True, errors="coerce")
    return df.loc[datas.between(inicio, fim, inclusive="both")].copy()


def remover_lancamentos_de_saldo(
    df: Optional[pd.DataFrame],
) -> pd.DataFrame:
    if df is None or df.empty:
        return pd.DataFrame()
    if "memo" not in df.columns:
        return df.copy()

    memorandos = df["memo"].fillna("").map(normalizar_texto)
    lancamento_saldo = memorandos.str.contains(
        r"\bsaldo\b|saldo\s+total\s+disponivel|saldo\s+movimentacao|"
        r"saldo\s+aplic|saldo\s+aplicacao|saldo\s+conta",
        regex=True,
    )
    return df[~lancamento_saldo].copy()


def filtrar_movimentos_validos_recebimento(
    df: Optional[pd.DataFrame],
) -> pd.DataFrame:
    if df is None or df.empty:
        return pd.DataFrame()
    movimentos = df.copy()
    texto = (
        movimentos.get("memo", pd.Series("", index=movimentos.index))
        .fillna("")
        .map(normalizar_texto)
        + " "
        + movimentos.get("descricao", pd.Series("", index=movimentos.index))
        .fillna("")
        .map(normalizar_texto)
        + " "
        + movimentos.get("fonte", pd.Series("", index=movimentos.index))
        .fillna("")
        .map(normalizar_texto)
        + " "
        + movimentos.get("tipo_transacao", pd.Series("", index=movimentos.index))
        .fillna("")
        .map(normalizar_texto)
    )
    excluir = texto.str.contains(
        r"transferencia\s+entre\s+contas|transf\.?\s+entre\s+contas|"
        r"transf(?:erencia)?\s+(?:mesma\s+)?titularidade|"
        r"mesma\s+tit\.?|mesma\s+titularidade|"
        r"transf(?:erencia)?\s+(?:entre|p/|para)\s+(?:minhas\s+)?contas?|"
        r"transf(?:erencia)?\s+(?:conta\s+propria|propria|mesma\s+empresa)|"
        r"credito\s+transf(?:erencia)?\s+(?:propria|entre\s+contas)|"
        r"\bted\b.*(?:mesma\s+titularidade|conta\s+propria)|"
        r"\bdoc\b.*(?:mesma\s+titularidade|conta\s+propria)|"
        r"resgate\s+(?:de\s+)?(?:aplicacao|aplicação|investimento)|"
        r"aplicacao\s+(?:automatica|financeira)|aplicação\s+(?:automatica|financeira)|"
        r"credito\s+(?:de\s+)?resgate|resgate\s+automatico|resgate\s+automático|"
        r"saldo\s+do\s+dia|saldo\s+anterior|saldo\s+total|"
        r"saldo\s+disponivel|saldo\s+movimentacao|"
        r"antecipacao\s+de\s+cartao|antecipacao\s+de\s+maquininha|"
        r"antecipa(?:cao|ção)?\s+(?:de\s+)?(?:recebiveis|recebíveis)|"
        r"\bantecipa\b",
        regex=True,
    )
    if "eh_antecipacao" in movimentos.columns:
        excluir = excluir | movimentos["eh_antecipacao"].fillna(False).astype(bool)
    return movimentos[~excluir].copy()


def separar_recebimentos_conta_corrente_clinicorp(
    df_banco: Optional[pd.DataFrame],
) -> Tuple[pd.DataFrame, pd.DataFrame]:
    if df_banco is None or df_banco.empty or "valor" not in df_banco.columns:
        return pd.DataFrame(), pd.DataFrame()

    creditos_brutos = df_banco[
        pd.to_numeric(df_banco["valor"], errors="coerce").fillna(0) > 0
    ].copy()
    if creditos_brutos.empty:
        return pd.DataFrame(), pd.DataFrame()

    tipos = creditos_brutos.get(
        "tipo_transacao",
        pd.Series("", index=creditos_brutos.index),
    ).fillna("").map(normalizar_texto)
    arquivos = creditos_brutos.get(
        "_arquivo_origem",
        pd.Series("", index=creditos_brutos.index),
    ).fillna("").map(normalizar_texto)
    fontes = creditos_brutos.get(
        "fonte",
        pd.Series("", index=creditos_brutos.index),
    ).fillna("").map(normalizar_texto)
    conta_corrente_clinicorp = (
        tipos.str.contains("entrada", regex=False)
        & (
            arquivos.str.contains("conta corrente", regex=False)
            | fontes.str.contains("bancos", regex=False)
        )
    )
    if conta_corrente_clinicorp.any():
        operacionais = creditos_brutos[conta_corrente_clinicorp].copy()
        ignorados = creditos_brutos[~conta_corrente_clinicorp].copy()
        return operacionais, ignorados

    creditos = filtrar_movimentos_validos_recebimento(creditos_brutos)
    if creditos.empty:
        return pd.DataFrame(), creditos_brutos.copy()
    return pd.DataFrame(), creditos.copy()


def remover_repasses_maquininha_do_banco(
    df: Optional[pd.DataFrame],
) -> pd.DataFrame:
    if df is None or df.empty or "memo" not in df.columns:
        return pd.DataFrame() if df is None else df.copy()

    movimentos = df.copy()
    memorandos = (
        movimentos.get("memo", pd.Series("", index=movimentos.index))
        .fillna("").map(normalizar_texto)
        + " "
        + movimentos.get("nome", pd.Series("", index=movimentos.index))
        .fillna("").map(normalizar_texto)
        + " "
        + movimentos.get("detalhe", pd.Series("", index=movimentos.index))
        .fillna("").map(normalizar_texto)
        + " "
        + movimentos.get("fonte", pd.Series("", index=movimentos.index))
        .fillna("").map(normalizar_texto)
        + " "
        + movimentos.get("tipo_ofx", pd.Series("", index=movimentos.index))
        .fillna("").map(normalizar_texto)
        + " "
        + movimentos.get("_arquivo_origem", pd.Series("", index=movimentos.index))
        .fillna("").map(normalizar_texto)
    )
    valores = pd.to_numeric(movimentos.get("valor", 0), errors="coerce").fillna(0)
    repasse_maquininha = memorandos.str.contains(
        r"recebimento\s+rede|redecard|cielo|stone|getnet|pagseguro|"
        r"mercado\s*pago|maquininha|visa|master|mast|elo|"
        r"infinite\s*pay|infinitepay|statements?",
        regex=True,
    )
    return movimentos[~repasse_maquininha].copy()


def remover_repasses_clinipay_do_banco(
    df_banco: Optional[pd.DataFrame],
    df_clinipay: Optional[pd.DataFrame],
) -> pd.DataFrame:
    if df_banco is None or df_banco.empty:
        return pd.DataFrame() if df_banco is None else df_banco.copy()
    if df_clinipay is None or df_clinipay.empty or "memo" not in df_banco.columns:
        return df_banco.copy()

    movimentos = df_banco.copy()
    memorandos = movimentos["memo"].fillna("").map(normalizar_texto)
    valores = pd.to_numeric(movimentos.get("valor", 0), errors="coerce").fillna(0)
    repasse_clinipay = (valores > 0) & memorandos.str.contains(
        r"clinipay|clini\s*pay|clinicorp\s*pay|clinicorp",
        regex=True,
    )
    return movimentos[~repasse_clinipay].copy()


def remover_fluxo_duplicado_clinipay(
    df_fluxo: Optional[pd.DataFrame],
    df_clinipay: Optional[pd.DataFrame],
) -> pd.DataFrame:
    if df_fluxo is None or df_fluxo.empty:
        return pd.DataFrame() if df_fluxo is None else df_fluxo.copy()
    if df_clinipay is None or df_clinipay.empty:
        return df_fluxo.copy()
    if "data" not in df_fluxo.columns or "valor" not in df_fluxo.columns:
        return df_fluxo.copy()
    if "data" not in df_clinipay.columns or "valor" not in df_clinipay.columns:
        return df_fluxo.copy()

    clinipay = df_clinipay.copy()
    clinipay_valores = pd.to_numeric(
        clinipay["valor"],
        errors="coerce",
    ).fillna(0)
    clinipay_datas = pd.to_datetime(
        clinipay["data"],
        errors="coerce",
    ).dt.date
    chaves_clinipay = Counter(
        (data, round(float(valor), 2))
        for data, valor in zip(clinipay_datas, clinipay_valores)
        if pd.notna(data) and float(valor) > 0
    )
    if not chaves_clinipay:
        return df_fluxo.copy()

    fluxo = df_fluxo.copy()
    fluxo_valores = pd.to_numeric(
        fluxo["valor"],
        errors="coerce",
    ).fillna(0)
    fluxo_datas = pd.to_datetime(
        fluxo["data"],
        errors="coerce",
    ).dt.date
    manter = []
    for data, valor in zip(fluxo_datas, fluxo_valores):
        chave = (data, round(float(valor), 2))
        if pd.notna(data) and float(valor) > 0 and chaves_clinipay.get(chave, 0) > 0:
            chaves_clinipay[chave] -= 1
            manter.append(False)
        else:
            manter.append(True)
    return fluxo.loc[manter].copy()


def calcular_antecipacoes(
    df_extrato: Optional[pd.DataFrame],
) -> Dict[str, float]:
    resultado = {"recebido": 0.0, "custo": 0.0, "liquido": 0.0}
    if df_extrato is None or df_extrato.empty:
        return resultado

    extrato = df_extrato.copy()
    if {
        "valor_bruto",
        "taxas_antecipacao",
        "eh_antecipacao",
    }.issubset(extrato.columns):
        antecipadas = extrato[
            extrato["eh_antecipacao"].fillna(False).astype(bool)
        ].copy()
        if "status_antecipacao" in antecipadas.columns:
            status = antecipadas["status_antecipacao"].map(normalizar_texto)
            antecipadas = antecipadas[status.str.contains("pago", regex=False)]
        recebido = float(antecipadas["valor_bruto"].fillna(0).sum())
        custo = float(
            antecipadas["taxas_antecipacao"].fillna(0).abs().sum()
        )
        liquido = float(antecipadas["valor"].fillna(0).sum())
        return {
            "recebido": recebido,
            "custo": custo,
            "liquido": liquido,
        }

    memorandos = extrato["memo"].fillna("").map(normalizar_texto)
    fontes = extrato.get(
        "fonte",
        pd.Series("", index=extrato.index),
    ).fillna("").map(normalizar_texto)
    eh_antecipacao = fontes.str.contains(
        "antecipacao de cartao",
        regex=False,
    ) | memorandos.str.contains(
        r"antecipa|antecip\.|recebiveis|recebíveis",
        regex=True,
    )
    eh_custo = memorandos.str.contains(
        r"taxa|tarifa|custo|encargo|desconto|fee",
        regex=True,
    )
    creditos = extrato.loc[
        eh_antecipacao
        & (extrato["valor"] > 0)
        & ~eh_custo,
        "valor",
    ]
    debitos = extrato.loc[
        eh_antecipacao & (extrato["valor"] < 0),
        "valor",
    ]
    custos_positivos = extrato.loc[
        eh_antecipacao
        & (extrato["valor"] > 0)
        & eh_custo,
        "valor",
    ]
    recebido = float(creditos.sum())
    custo = float(debitos.abs().sum() + custos_positivos.sum())
    return {
        "recebido": recebido,
        "custo": custo,
        "liquido": recebido - custo,
    }


def calcular_antecipacoes_maquininha(
    df_maquininha: Optional[pd.DataFrame],
) -> Dict[str, float]:
    resultado = {"recebido": 0.0, "custo": 0.0, "liquido": 0.0}
    if df_maquininha is None or df_maquininha.empty:
        return resultado

    base = df_maquininha.copy()
    antecipada = base.get(
        "antecipada_maquininha",
        pd.Series("", index=base.index),
    ).fillna("").map(normalizar_texto)
    mascara_antecipada = antecipada.str.contains(
        r"\bsim\b|antecipad",
        regex=True,
    )
    if not mascara_antecipada.any():
        return resultado

    antecipadas = base[mascara_antecipada].copy()
    liquido = pd.to_numeric(
        antecipadas.get("valor", 0),
        errors="coerce",
    ).fillna(0).abs()
    bruto = pd.to_numeric(
        antecipadas.get("valor_bruto_maquininha", 0),
        errors="coerce",
    ).fillna(0).abs()
    bruto = bruto.mask((bruto <= 0) & (liquido > 0), liquido)
    custo = (bruto - liquido).clip(lower=0)
    custo_taxa = pd.to_numeric(
        antecipadas.get("taxa_maquininha", 0),
        errors="coerce",
    ).fillna(0).abs()
    custo = custo.mask(custo <= 0, custo_taxa)
    return {
        "recebido": float(bruto.sum()),
        "custo": float(custo.sum()),
        "liquido": float(liquido.sum()),
    }


def conciliar_recebimentos(
    df_vendas: Optional[pd.DataFrame],
    df_extrato: Optional[pd.DataFrame],
    df_clinipay: Optional[pd.DataFrame] = None,
    df_fluxo_caixa: Optional[pd.DataFrame] = None,
) -> Dict[str, object]:
    """Calcula recebimentos identificados sem completar com créditos genéricos."""
    vazio = {
        "total": 0.0,
        "pix": 0.0,
        "cartao": 0.0,
        "antecipacao": 0.0,
        "outros": 0.0,
        "bancario": 0.0,
        "clinipay": 0.0,
        "direto": 0.0,
        "vendas_conciliadas": pd.DataFrame(),
        "creditos_conciliados": pd.DataFrame(),
    }
    if (
        df_vendas is None or df_vendas.empty
    ):
        return vazio

    vendas = df_vendas.copy().reset_index(drop=True)
    if df_extrato is not None and not df_extrato.empty:
        creditos = df_extrato[df_extrato["valor"] > 0].copy().reset_index(drop=True)
        creditos = filtrar_movimentos_validos_recebimento(creditos)
    else:
        creditos = pd.DataFrame(columns=["data", "valor", "memo"])
    if vendas.empty:
        return vazio

    if not creditos.empty:
        creditos["data"] = pd.to_datetime(creditos["data"], errors="coerce")
        creditos["memo_norm"] = creditos["memo"].fillna("").map(normalizar_texto)
        creditos["fonte_norm"] = creditos.get(
            "fonte",
            pd.Series("", index=creditos.index),
        ).fillna("").map(normalizar_texto)
    else:
        creditos["memo_norm"] = pd.Series(dtype="object")
        creditos["fonte_norm"] = pd.Series(dtype="object")

    mascara_pix = creditos["memo_norm"].str.contains(
        r"pix[_\s-]*(?:cred|cre)|recebimento\s+pix", regex=True
    )
    if "eh_antecipacao" in creditos.columns:
        marcada_antecipacao = (
            creditos["eh_antecipacao"].fillna(False).astype(bool)
        )
    else:
        marcada_antecipacao = creditos["fonte_norm"].str.contains(
            "antecipacao de cartao",
            regex=False,
        )
    mascara_antecipacao = marcada_antecipacao | creditos["memo_norm"].str.contains(
        r"antecipa|antecip\.|recebiveis|recebíveis",
        regex=True,
    )
    mascara_custo_antecipacao = creditos["memo_norm"].str.contains(
        r"taxa|tarifa|custo|encargo|desconto|fee",
        regex=True,
    )
    mascara_antecipacao = (
        mascara_antecipacao & ~mascara_custo_antecipacao
    )
    mascara_cartao = creditos["memo_norm"].str.contains(
        r"visa|master|mast|elo|cartao|cartão|credito|crédito|antec|"
        r"recebimento\s+rede|redecard|cielo|stone|getnet|pagseguro|"
        r"mercado\s*pago",
        regex=True,
    ) & ~mascara_pix & ~mascara_antecipacao
    fonte_venda_amigotech = (
        "fonte_venda" in vendas.columns
        and vendas["fonte_venda"].fillna("").map(normalizar_texto)
        .eq("amigotech").any()
    )
    fonte_venda_conta_azul = (
        "fonte_venda" in vendas.columns
        and vendas["fonte_venda"].fillna("").map(normalizar_texto)
        .eq("conta azul").any()
    )
    fonte_venda_com_banco_direto = (
        "fonte_venda" in vendas.columns
        and vendas["fonte_venda"].fillna("").map(normalizar_texto)
        .isin(["belle software", "amigotech", "conta azul"]).any()
    )
    mascara_credito_bancario = (
        fonte_venda_com_banco_direto
        & ~(mascara_pix | mascara_antecipacao | mascara_cartao)
    )

    pix = creditos[mascara_pix].copy()
    cartoes = creditos[mascara_cartao].copy()
    creditos_bancarios = creditos[mascara_credito_bancario].copy()
    pix["tipo_recebimento"] = "PIX"
    cartoes["tipo_recebimento"] = "Cartão"
    creditos_bancarios["tipo_recebimento"] = "Crédito bancário"
    if df_clinipay is not None and not df_clinipay.empty:
        clinipay = df_clinipay[df_clinipay["valor"] > 0].copy()
        clinipay = filtrar_movimentos_validos_recebimento(clinipay)
        clinipay["data"] = pd.to_datetime(clinipay["data"], errors="coerce")
        clinipay["tipo_recebimento"] = "Clinipay"
    else:
        clinipay = pd.DataFrame(columns=creditos.columns)
    if df_fluxo_caixa is not None and not df_fluxo_caixa.empty:
        pagamentos_diretos = df_fluxo_caixa[df_fluxo_caixa["valor"] > 0].copy()
        pagamentos_diretos = filtrar_movimentos_validos_recebimento(
            pagamentos_diretos
        )
        pagamentos_diretos["data"] = pd.to_datetime(
            pagamentos_diretos["data"], errors="coerce"
        )
        pagamentos_diretos["tipo_recebimento"] = "Pagamento direto"
    else:
        pagamentos_diretos = pd.DataFrame(columns=creditos.columns)
    creditos_identificados = pd.concat(
        [
            pix,
            cartoes,
            creditos_bancarios,
            clinipay,
            pagamentos_diretos,
        ],
        ignore_index=True,
        sort=False,
    )
    creditos_identificados = creditos_identificados.sort_values(
        "data", na_position="last"
    )
    total_creditos = float(creditos_identificados["valor"].sum())
    if fonte_venda_amigotech or fonte_venda_conta_azul:
        coluna_venda = "valor_recebido" if "valor_recebido" in vendas.columns else "valor"
        total_vendas_conciliaveis = float(
            pd.to_numeric(vendas[coluna_venda], errors="coerce")
            .fillna(0)
            .sum()
        )
        total_conciliado = min(total_vendas_conciliaveis, total_creditos)
    else:
        total_conciliado = total_creditos

    saldo_conciliavel = total_conciliado
    linhas_conciliadas = []
    total_pix = 0.0
    total_cartao = 0.0
    total_antecipacao = 0.0
    total_bancario = 0.0
    total_clinipay = 0.0
    total_direto = 0.0
    for _, credito in creditos_identificados.iterrows():
        if saldo_conciliavel <= 0:
            break
        valor_alocado = min(float(credito["valor"]), saldo_conciliavel)
        linha = credito.copy()
        linha["valor"] = valor_alocado
        linhas_conciliadas.append(linha)
        if credito["tipo_recebimento"] == "PIX":
            total_pix += valor_alocado
        elif credito["tipo_recebimento"] == "Antecipação":
            total_antecipacao += valor_alocado
        elif credito["tipo_recebimento"] == "Clinipay":
            total_clinipay += valor_alocado
        elif credito["tipo_recebimento"] == "Pagamento direto":
            total_direto += valor_alocado
        elif credito["tipo_recebimento"] == "Crédito bancário":
            total_bancario += valor_alocado
        else:
            total_cartao += valor_alocado
        saldo_conciliavel -= valor_alocado

    creditos_conciliados = (
        pd.DataFrame(linhas_conciliadas)
        if linhas_conciliadas else pd.DataFrame(columns=creditos.columns)
    )
    return {
        "total": total_conciliado,
        "pix": total_pix,
        "cartao": total_cartao,
        "antecipacao": total_antecipacao,
        "outros": 0.0,
        "bancario": total_bancario,
        "clinipay": total_clinipay,
        "direto": total_direto,
        "vendas_conciliadas": vendas,
        "creditos_conciliados": creditos_conciliados,
    }


def conciliar_pagamentos(
    df_contas: Optional[pd.DataFrame],
    df_extrato: Optional[pd.DataFrame],
) -> Dict[str, object]:
    """Concilia contas pagas e mantém despesas importadas do gerencial."""
    vazio = {
        "total": 0.0,
        "fixos": 0.0,
        "variaveis": 0.0,
        "retiradas": 0.0,
        "pro_labore": 0.0,
        "antecipacoes_lucro": 0.0,
        "total_planilha": 0.0,
        "total_banco": 0.0,
        "diferenca_banco": 0.0,
        "contas_conciliadas": pd.DataFrame(),
        "debitos_extrato": pd.DataFrame(),
    }
    if df_contas is None or df_contas.empty:
        return vazio

    contas = df_contas.copy()
    if "grupo_custo" not in contas.columns:
        contas["grupo_custo"] = classificar_grupo_custo(contas)
    texto_contas = (
        contas.get("categoria", pd.Series("", index=contas.index))
        .fillna("").map(normalizar_texto)
        + " "
        + contas.get("descricao", pd.Series("", index=contas.index))
        .fillna("").map(normalizar_texto)
        + " "
        + contas.get("forma", pd.Series("", index=contas.index))
        .fillna("").map(normalizar_texto)
    )
    movimentacao_entre_contas = texto_contas.str.contains(
        r"^\s*1\.1\.\d|1\.1\.1\.\d+.*(?:sicoob|sicredi|itau|itaú|"
        r"bradesco|santander|banco|infinite\s*pay|infinitepay)|"
        r"transferencia\s+entre\s+contas|transf\.?\s+entre\s+contas|"
        r"movimentacao\s+conta|movimentação\s+conta|repasse\s+entre\s+contas",
        regex=True,
    ) & ~texto_contas.str.contains(
        r"tarifa|taxa|juros|financiamento|emprestimo|empréstimo",
        regex=True,
    )
    contas = contas[~movimentacao_entre_contas].copy()
    vazio["contas_conciliadas"] = contas.iloc[0:0].copy()

    fonte_norm = contas.get(
        "fonte",
        pd.Series("", index=contas.index),
    ).fillna("").map(normalizar_texto)
    mascara_importada = fonte_norm.str.contains(
        "gerencial de resultados belle",
        regex=False,
    )
    contas_importadas = contas[mascara_importada].copy()
    contas_conciliaveis = contas[~mascara_importada].copy()
    formas_conciliaveis = contas_conciliaveis.get(
        "forma",
        pd.Series("", index=contas_conciliaveis.index),
    ).fillna("").map(normalizar_texto)
    descricoes_conciliaveis = contas_conciliaveis.get(
        "descricao",
        pd.Series("", index=contas_conciliaveis.index),
    ).fillna("").map(normalizar_texto)
    categorias_conciliaveis = contas_conciliaveis.get(
        "categoria",
        pd.Series("", index=contas_conciliaveis.index),
    ).fillna("").map(normalizar_texto)
    mascara_dinheiro = (
        formas_conciliaveis.str.contains(r"dinheiro|caixa", regex=True)
        | descricoes_conciliaveis.str.contains(r"pagamento\s+em\s+dinheiro", regex=True)
        | categorias_conciliaveis.str.contains(r"dinheiro|caixa", regex=True)
    )
    contas_dinheiro = contas_conciliaveis[mascara_dinheiro].copy()
    contas_conciliaveis = contas_conciliaveis[~mascara_dinheiro].copy()

    linhas = []
    totais = {
        "Custo Fixo": 0.0,
        "Custo Variável": 0.0,
        "Retirada de Lucro": 0.0,
        "Pró-labore": 0.0,
        "Antecipação de Lucro": 0.0,
    }
    for _, conta in contas_importadas.sort_values(
        "data", na_position="last"
    ).iterrows():
        linha = conta.copy()
        linha["valor_original"] = float(conta["valor"])
        linhas.append(linha)
        grupo = conta.get("grupo_custo", "Custo Variável")
        totais[grupo] = totais.get(grupo, 0.0) + float(conta["valor"])
    total_importado = float(contas_importadas["valor"].sum())

    for _, conta in contas_dinheiro.sort_values(
        "data", na_position="last"
    ).iterrows():
        linha = conta.copy()
        linha["valor_original"] = float(conta["valor"])
        linha["forma"] = linha.get("forma", "") or "Dinheiro"
        linhas.append(linha)
        grupo = conta.get("grupo_custo", "Custo Variável")
        totais[grupo] = totais.get(grupo, 0.0) + float(conta["valor"])
    total_dinheiro = float(contas_dinheiro["valor"].sum())

    contas_conciliaveis = contas_conciliaveis.sort_values(
        "data", na_position="last"
    )
    for _, conta in contas_conciliaveis.iterrows():
        linha = conta.copy()
        linha["valor_original"] = float(conta["valor"])
        linhas.append(linha)
        grupo = conta.get("grupo_custo", "Custo Variável")
        totais[grupo] = totais.get(grupo, 0.0) + float(conta["valor"])

    total_planilha = float(
        total_importado
        + total_dinheiro
        + contas_conciliaveis["valor"].sum()
    )

    if df_extrato is None or df_extrato.empty:
        return {
            "total": total_planilha,
            "fixos": totais["Custo Fixo"],
            "variaveis": totais["Custo Variável"],
            "retiradas": totais["Retirada de Lucro"],
            "pro_labore": totais["Pró-labore"],
            "antecipacoes_lucro": totais["Antecipação de Lucro"],
            "total_planilha": total_planilha,
            "total_banco": 0.0,
            "diferenca_banco": total_planilha,
            "contas_conciliadas": pd.DataFrame(linhas),
            "debitos_extrato": pd.DataFrame(),
        }
    debitos = df_extrato[df_extrato["valor"] < 0].copy()
    if not debitos.empty and "memo" in debitos.columns:
        texto_debitos = (
            debitos.get("memo", pd.Series("", index=debitos.index))
            .fillna("").map(normalizar_texto)
            + " "
            + debitos.get("nome", pd.Series("", index=debitos.index))
            .fillna("").map(normalizar_texto)
            + " "
            + debitos.get("detalhe", pd.Series("", index=debitos.index))
            .fillna("").map(normalizar_texto)
            + " "
            + debitos.get("fonte", pd.Series("", index=debitos.index))
            .fillna("").map(normalizar_texto)
            + " "
            + debitos.get("tipo_ofx", pd.Series("", index=debitos.index))
            .fillna("").map(normalizar_texto)
            + " "
            + debitos.get("tipo_transacao", pd.Series("", index=debitos.index))
            .fillna("").map(normalizar_texto)
            + " "
            + debitos.get("_arquivo_origem", pd.Series("", index=debitos.index))
            .fillna("").map(normalizar_texto)
        )
        debitos_nao_despesa = texto_debitos.str.contains(
            r"transferencia\s+entre\s+contas|transf\.?\s+entre\s+contas|"
            r"saldo|aplicacao|aplic\.|resgate|investimento|"
            r"movimentacao\s+conta|maquininha|infinite\s*pay|infinitepay|"
            r"\bmaquininha\b|\bmaq\b|tipo_ofx\s*maquininha|"
            r"deposito\s+infinitepay|dep[oó]sito\s+infinitepay|"
            r"statements?.*pix.*enviado|pix.*enviado.*statements?",
            regex=True,
        )
        debitos = debitos[~debitos_nao_despesa].copy()
    total_debitos = (
        float(debitos["valor"].abs().sum())
        if not debitos.empty else 0.0
    )

    return {
        "total": total_planilha,
        "fixos": totais["Custo Fixo"],
        "variaveis": totais["Custo Variável"],
        "retiradas": totais["Retirada de Lucro"],
        "pro_labore": totais["Pró-labore"],
        "antecipacoes_lucro": totais["Antecipação de Lucro"],
        "total_planilha": total_planilha,
        "total_banco": total_debitos,
        "diferenca_banco": total_planilha - total_debitos,
        "contas_conciliadas": pd.DataFrame(linhas),
        "debitos_extrato": debitos,
    }


def resposta_assistente(
    pergunta: str,
    periodo_label: str,
    recebimentos_periodo: float,
    vendas_periodo: float,
    despesas_periodo: float,
    custos_fixos_periodo: float,
    custos_variaveis_periodo: float,
    retiradas_periodo: float,
    antecipacoes_lucro_periodo: float = 0.0,
) -> str:
    texto = normalizar_texto(pergunta)
    saidas_lucro_periodo = retiradas_periodo + antecipacoes_lucro_periodo
    resultado_operacional_periodo = (
        recebimentos_periodo - despesas_periodo - saidas_lucro_periodo
    )
    resultado_final_periodo = resultado_operacional_periodo
    margem_operacional = (
        resultado_operacional_periodo / recebimentos_periodo * 100
        if recebimentos_periodo > 0 else 0
    )

    if any(p in texto for p in ["receita", "recebimento", "entrada", "faturei"]):
        return (
            f"No período {periodo_label}, os recebimentos totalizam "
            f"{fmt_brl_chat(recebimentos_periodo)}."
        )
    if "retirada" in texto or "distribuicao" in texto:
        return (
            f"A retirada de lucro no período é {fmt_brl_chat(retiradas_periodo)}. "
            "Ela é identificada pela categoria e apresentada separadamente dos "
            "demais custos variáveis."
        )
    if any(p in texto for p in ["despesa", "gasto", "custo"]):
        return (
            f"As despesas somam {fmt_brl_chat(despesas_periodo)}.\n\n"
            f"Custos fixos: {fmt_brl_chat(custos_fixos_periodo)}.\n\n"
            f"Custos variáveis: {fmt_brl_chat(custos_variaveis_periodo)}.\n\n"
            f"A retirada de lucro, apresentada à parte, é "
            f"{fmt_brl_chat(retiradas_periodo)} e a antecipação de lucro é "
            f"{fmt_brl_chat(antecipacoes_lucro_periodo)}."
        )
    if any(p in texto for p in ["resultado", "lucro", "prejuizo", "margem"]):
        situacao = "positivo" if resultado_operacional_periodo >= 0 else "negativo"
        margem_txt = f"{margem_operacional:.1f}".replace(".", ",")
        return (
            f"Resultado: {fmt_brl_chat(resultado_operacional_periodo)} "
            f"({situacao}), com margem de {margem_txt}%.\n\n"
            "A conta usada é recebimentos menos despesas, retirada de lucro "
            "e antecipação de lucro.\n\n"
            f"Resultado final: "
            f"{fmt_brl_chat(resultado_final_periodo)}."
        )
    if any(p in texto for p in ["periodo", "mensal", "bimestral", "trimestral", "anual"]):
        return f"A análise exibida considera o período {periodo_label}."
    if any(p in texto for p in ["importar", "arquivo", "excel", "ofx"]):
        if st.session_state.get("share_mode"):
            return (
                "Este relatório está em modo somente leitura. Apenas o responsável "
                "pelo painel pode importar ou substituir arquivos."
            )
        return (
            "Use **Importar Arquivos** no menu lateral. A planilha de contas "
            "precisa conter valor, vencimento e categoria; o extrato deve estar "
            "no formato OFX."
        )
    if "fechamento" in texto:
        return (
            "O Fechamento resume recebimentos, custos fixos, custos variáveis, "
            "resultado operacional e mostra a retirada de lucro separadamente."
        )

    return (
        "Ainda não tenho uma resposta padrão para essa pergunta. "
        "Entre em contato com o suporte financeiro para uma orientação específica."
    )


@st.dialog("💬 Assistente Financeiro", width="large")
def abrir_assistente(
    periodo_label,
    recebimentos_periodo,
    vendas_periodo,
    despesas_periodo,
    custos_fixos_periodo,
    custos_variaveis_periodo,
    retiradas_periodo,
    antecipacoes_lucro_periodo=0.0,
):
    st.caption(
        "Pergunte sobre receitas, despesas, resultado, retirada, período ou importação."
    )

    for item in st.session_state.historico_chat[-6:]:
        with st.chat_message(item["papel"]):
            st.markdown(item["texto"])

    with st.form("form_assistente_financeiro", clear_on_submit=True):
        pergunta_chat = st.text_input(
            "Sua pergunta",
            placeholder="Ex.: Qual foi o resultado deste período?",
        )
        enviar_chat = st.form_submit_button("Perguntar", use_container_width=True)

    if enviar_chat and pergunta_chat.strip():
        resposta_chat = resposta_assistente(
            pergunta_chat,
            periodo_label,
            recebimentos_periodo,
            vendas_periodo,
            despesas_periodo,
            custos_fixos_periodo,
            custos_variaveis_periodo,
            retiradas_periodo,
            antecipacoes_lucro_periodo,
        )
        st.session_state.historico_chat.extend([
            {"papel": "user", "texto": pergunta_chat.strip()},
            {"papel": "assistant", "texto": resposta_chat},
        ])
        st.rerun(scope="fragment")


def validar_entrada_usuario(username: str, password: str) -> Tuple[bool, str]:
    if not username or not password:
        return False, "❌ Usuário e senha são obrigatórios"
    if len(username) < 3:
        return False, "❌ Usuário deve ter pelo menos 3 caracteres"
    if len(password) < 8:
        return False, "❌ Senha deve ter pelo menos 8 caracteres"
    if any(char in username for char in ["'", '"', ";", "--"]):
        return False, "❌ Usuário contém caracteres inválidos"
    return True, ""


def ler_tabela_flexivel(uploaded_file) -> pd.DataFrame:
    """Lê exportações tabulares de Excel, Google Sheets, CSV e ODS."""
    extensao = Path(uploaded_file.name).suffix.lower()
    uploaded_file.seek(0)
    if extensao == ".csv":
        try:
            return pd.read_csv(
                uploaded_file,
                sep=None,
                engine="python",
                encoding="utf-8-sig",
            )
        except UnicodeDecodeError:
            uploaded_file.seek(0)
            return pd.read_csv(
                uploaded_file,
                sep=None,
                engine="python",
                encoding="latin-1",
            )
    if extensao in {".xlsx", ".xls", ".ods"}:
        try:
            return pd.read_excel(uploaded_file)
        except ImportError as e:
            if extensao == ".xls" and "xlrd" in str(e).lower():
                return converter_xls_com_excel(uploaded_file)
            raise
    raise ValueError(
        "Formato não suportado. Exporte o documento como XLSX, XLS, CSV ou ODS."
    )


def converter_xls_com_excel(
    uploaded_file,
    header=0,
    sheet_name=0,
) -> pd.DataFrame:
    """Converte .xls pelo Excel instalado quando xlrd não está disponível."""
    if os.name != "nt":
        raise RuntimeError(
            "Este arquivo usa o formato antigo .xls. Instale xlrd>=2.0.1 "
            "no servidor ou exporte como .xlsx."
        )

    uploaded_file.seek(0)
    with tempfile.TemporaryDirectory(prefix="dashboard_xls_") as pasta:
        origem = Path(pasta) / "origem.xls"
        destino = Path(pasta) / "convertido.xlsx"
        origem.write_bytes(uploaded_file.read())

        script = """
$ErrorActionPreference = 'Stop'
$excel = $null
$workbook = $null
try {
    $excel = New-Object -ComObject Excel.Application
    $excel.Visible = $false
    $excel.DisplayAlerts = $false
    $workbook = $excel.Workbooks.Open($env:DASHBOARD_XLS_ORIGEM)
    $workbook.SaveAs($env:DASHBOARD_XLS_DESTINO, 51)
    $workbook.Close($false)
} finally {
    if ($workbook -ne $null) {
        [Runtime.InteropServices.Marshal]::ReleaseComObject($workbook) |
            Out-Null
    }
    if ($excel -ne $null) {
        $excel.Quit()
        [Runtime.InteropServices.Marshal]::ReleaseComObject($excel) |
            Out-Null
    }
}
"""
        ambiente = os.environ.copy()
        ambiente["DASHBOARD_XLS_ORIGEM"] = str(origem)
        ambiente["DASHBOARD_XLS_DESTINO"] = str(destino)
        powershell = (
            Path(os.environ.get("ProgramFiles", r"C:\Program Files"))
            / "PowerShell" / "7" / "pwsh.exe"
        )
        executavel = str(powershell) if powershell.exists() else "powershell.exe"
        resultado = subprocess.run(
            [executavel, "-NoProfile", "-NonInteractive", "-Command", script],
            env=ambiente,
            capture_output=True,
            text=True,
            timeout=60,
            check=False,
        )
        if resultado.returncode != 0 or not destino.exists():
            raise RuntimeError(
                "Não foi possível abrir o Excel antigo automaticamente. "
                "Feche e reabra o Streamlit pela sua sessão do Windows ou "
                "salve o arquivo como .xlsx no Excel."
            )
        return pd.read_excel(destino, header=header, sheet_name=sheet_name)


def ler_todas_tabelas_flexiveis(uploaded_file, header=0) -> List[pd.DataFrame]:
    """Lê todas as abas possíveis, mantendo compatibilidade com CSV e XLS."""
    extensao = Path(uploaded_file.name).suffix.lower()
    uploaded_file.seek(0)

    if extensao == ".csv":
        return [ler_tabela_flexivel(uploaded_file)]

    if extensao in {".xlsx", ".xls", ".ods"}:
        try:
            planilhas = pd.read_excel(
                uploaded_file,
                sheet_name=None,
                header=header,
            )
            return [
                quadro for quadro in planilhas.values()
                if quadro is not None and not quadro.empty
            ]
        except ImportError as e:
            if extensao == ".xls" and "xlrd" in str(e).lower():
                uploaded_file.seek(0)
                convertido = converter_xls_com_excel(
                    uploaded_file,
                    header=header,
                    sheet_name=None,
                )
                return [
                    quadro for quadro in convertido.values()
                    if quadro is not None and not quadro.empty
                ]
            raise

    return [ler_tabela_flexivel(uploaded_file)]


def promover_cabecalho_por_termos(
    df: pd.DataFrame,
    termos: List[str],
) -> pd.DataFrame:
    """Encontra a linha de cabeçalho em exportações com linhas acima da tabela."""
    if df is None or df.empty:
        return pd.DataFrame()

    termos_norm = [normalizar_texto(termo) for termo in termos]
    for idx, linha in df.iterrows():
        texto_linha = " ".join(
            normalizar_texto(valor)
            for valor in linha.tolist()
            if str(valor).strip() and str(valor).lower() != "nan"
        )
        if all(termo in texto_linha for termo in termos_norm):
            promovido = df.iloc[idx + 1:].copy()
            promovido.columns = [
                str(valor).strip() if str(valor).strip() else f"coluna_{pos}"
                for pos, valor in enumerate(linha.tolist())
            ]
            return promovido.dropna(how="all")

    return df


def ler_tabelas_amigotech(uploaded_file) -> List[pd.DataFrame]:
    candidatos = []
    candidatos.extend(ler_todas_tabelas_flexiveis(uploaded_file))
    candidatos.extend(ler_todas_tabelas_flexiveis(uploaded_file, header=None))

    quadros = []
    vistos = set()
    for bruto in candidatos:
        df = promover_cabecalho_por_termos(
            bruto,
            ["data de pagamento", "tipo", "valor"],
        )
        if df is None or df.empty:
            continue
        df = df.loc[:, ~df.columns.duplicated()].copy()
        df.columns = [str(col).strip() for col in df.columns]
        assinatura = tuple(normalizar_texto(col) for col in df.columns)
        if assinatura in vistos:
            continue
        vistos.add(assinatura)
        quadros.append(df)
    return quadros


def processar_tabela_generica(uploaded_file) -> pd.DataFrame:
    """Preserva relatórios ainda não usados nos cálculos do dashboard."""
    try:
        df = ler_tabela_flexivel(uploaded_file)
        df = df.loc[:, ~df.columns.duplicated()].copy()
        df.columns = [str(col).strip() for col in df.columns]
        df = df.dropna(how="all")
        if df.empty:
            st.warning(f"⚠️ {uploaded_file.name}: nenhum registro encontrado.")
        return df
    except Exception as e:
        logger.error("Erro ao processar relatório genérico: %s", e)
        st.error(f"❌ Erro ao processar {uploaded_file.name}: {str(e)}")
        return pd.DataFrame()


def processar_contas_receber(uploaded_file) -> pd.DataFrame:
    try:
        import re

        quadros = []
        candidatos = []
        candidatos.extend(ler_todas_tabelas_flexiveis(uploaded_file))

        for bruto in candidatos:
            df = promover_cabecalho_por_termos(
                bruto,
                ["vcto", "valor"],
            )
            df = df.loc[:, ~df.columns.duplicated()].copy()
            mapa = {normalizar_texto(col): col for col in df.columns}

            def localizar(*nomes):
                for nome in nomes:
                    coluna = mapa.get(normalizar_texto(nome))
                    if coluna is not None:
                        return coluna
                return None

            col_valor_bruto = localizar(
                "valor", "valor a receber", "valor total", "total", "saldo"
            )
            col_valor_liquido = localizar(
                "valor liq.", "valor líquido", "valor liquido"
            )
            col_data_confirmacao = localizar(
                "data conf.", "data conf", "data confirmação",
                "data confirmacao", "data de confirmação",
                "data de confirmacao",
            )
            col_data = localizar(
                "vencimento", "vcto.", "vcto", "data de vencimento",
                "recebimento", "data"
            )
            col_descricao = localizar(
                "cliente", "paciente", "titular", "descrição", "descricao",
                "fornecedor", "pagador"
            )
            col_status = localizar("status", "situação", "situacao")
            col_conta_destino = localizar(
                "conta destino", "conta bancária", "conta bancaria",
                "conta", "destino"
            )
            col_observacao = localizar(
                "observação", "observacao", "obs", "observacoes",
                "observações"
            )
            if not (col_valor_bruto or col_valor_liquido) or not col_data:
                continue

            def data_transferencia_observacao(valor):
                texto = "" if pd.isna(valor) else str(valor)
                texto_norm = normalizar_texto(texto)
                if "transf" not in texto_norm and "transferencia" not in texto_norm:
                    return pd.NaT
                match = re.search(r"(\d{1,2}/\d{1,2}/\d{2,4})", texto)
                if not match:
                    return pd.NaT
                return pd.to_datetime(
                    match.group(1),
                    dayfirst=True,
                    errors="coerce",
                )

            df["valor"] = (
                df[col_valor_bruto].apply(parse_valor_br)
                if col_valor_bruto else df[col_valor_liquido].apply(parse_valor_br)
            )
            df["valor_recebido"] = (
                df[col_valor_liquido].apply(parse_valor_br)
                if col_valor_liquido else df["valor"]
            )
            df["data"] = pd.to_datetime(
                df[col_data], dayfirst=True, errors="coerce"
            )
            df["observacao"] = (
                df[col_observacao].fillna("").astype(str).str.strip()
                if col_observacao else ""
            )
            if col_data_confirmacao:
                datas_iso = pd.to_datetime(
                    df[col_data_confirmacao],
                    errors="coerce",
                )
                datas_br = pd.to_datetime(
                    df[col_data_confirmacao],
                    dayfirst=True,
                    errors="coerce",
                )
                df["data_recebimento_belle"] = datas_iso.fillna(datas_br)
            else:
                df["data_recebimento_belle"] = df["data"]
            df["valor_recebimento_belle"] = df["valor_recebido"]
            df["transferencia_recebimento_belle"] = (
                df["observacao"].fillna("").map(normalizar_texto).str.contains(
                    r"transf|transferencia|transferência|entre\s+contas|"
                    r"conta\s+propria|conta\s+própria|mesma\s+titularidade",
                    regex=True,
                )
            )
            datas_transferencia = df["observacao"].apply(
                data_transferencia_observacao
            )
            df.loc[datas_transferencia.notna(), "data"] = (
                datas_transferencia[datas_transferencia.notna()]
            )
            df["descricao"] = (
                df[col_descricao].fillna("").astype(str).str.strip()
                if col_descricao else "Conta a receber"
            )
            df["status"] = (
                df[col_status].fillna("").astype(str).str.strip()
                if col_status else ""
            )
            df["conta_destino"] = (
                df[col_conta_destino].fillna("").astype(str).str.strip()
                if col_conta_destino else ""
            )
            df["fonte_venda"] = "Belle Software"
            df["_linha_origem_belle"] = df.index
            df = df.dropna(subset=["valor", "data"])
            df = df[df["valor"] > 0]
            if not df.empty:
                quadros.append(df)

        if not quadros:
            st.error(
                "❌ Contas a receber precisa conter uma coluna de valor e "
                "outra de vencimento/data."
            )
            return pd.DataFrame()

        return (
            pd.concat(quadros, ignore_index=True, sort=False)
            .reset_index(drop=True)
        )
    except Exception as e:
        logger.error("Erro ao processar contas a receber: %s", e)
        st.error(f"❌ Erro ao processar contas a receber: {str(e)}")
        return pd.DataFrame()


def remover_vendas_nao_identificadas_infinity(
    df: Optional[pd.DataFrame],
) -> Optional[pd.DataFrame]:
    if df is None or df.empty:
        return df
    suspeitas = filtrar_vendas_nao_identificadas_infinity(df)
    if suspeitas.empty:
        return df
    return df.drop(index=suspeitas.index).copy()


def filtrar_vendas_nao_identificadas_infinity(
    df: Optional[pd.DataFrame],
) -> pd.DataFrame:
    if df is None or df.empty or "observacao" not in df.columns:
        return pd.DataFrame()
    observacao_norm = df["observacao"].fillna("").map(normalizar_texto)
    mascara = (
        observacao_norm.str.contains(
            r"nao\s+identifiquei|nao\s+identificada|nao\s+localizada",
            regex=True,
        )
        & observacao_norm.str.contains(
            r"infinite\s*pay|infinitepay|infinity\s*pay|infinitypay",
            regex=True,
        )
    )
    return df[mascara].copy()


def processar_amigotech_receber(uploaded_file) -> pd.DataFrame:
    try:
        quadros = []
        candidatos = ler_tabelas_amigotech(uploaded_file)

        for df in candidatos:
            mapa = {normalizar_texto(col): col for col in df.columns}

            def localizar(*nomes):
                nomes_norm = [normalizar_texto(nome) for nome in nomes]
                for nome in nomes_norm:
                    coluna = mapa.get(nome)
                    if coluna is not None:
                        return coluna
                for chave, coluna in mapa.items():
                    if any(nome in chave for nome in nomes_norm):
                        return coluna
                return None

            col_valor = localizar(
                "valor liquido r$", "valor líquido r$",
                "valor liquido", "valor líquido", "valor recebido",
                "valor pago", "recebido", "valor", "total"
            )
            col_valor_bruto = localizar(
                "valor original r$", "valor bruto", "valor original",
                "valor total", "total",
                "valor"
            )
            col_data = localizar(
                "data recebimento", "data de recebimento", "recebido em",
                "data pagamento", "data de pagamento", "vencimento", "data"
            )
            col_descricao = localizar(
                "cliente", "paciente", "pagador", "responsavel",
                "responsável", "nome", "descricao", "descrição",
                "procedimento"
            )
            col_categoria = localizar("categoria")
            col_classificacao = localizar("classificacao", "classificação")
            col_forma = localizar(
                "forma de pagamento", "forma pagamento",
                "meio de pagamento", "pagamento"
            )
            col_status = localizar("status", "situacao", "situação")
            col_tipo = localizar("tipo")
            col_conta = localizar(
                "banco", "conta destino", "conta", "forma de pagamento",
                "forma pagamento", "meio de pagamento", "pagamento"
            )
            col_observacao = localizar("observacao", "observação", "obs")

            if not col_valor or not col_data:
                continue

            if col_tipo:
                tipo_norm = df[col_tipo].fillna("").map(normalizar_texto)
                df = df[tipo_norm.str.contains("entrada", regex=False)].copy()
                if df.empty:
                    continue
            if col_status:
                status_norm = df[col_status].fillna("").map(normalizar_texto)
                nao_recebido = status_norm.str.contains(
                    r"aberto|pendente|cancel|estorn|vencid|nao\s+pago|não\s+pago",
                    regex=True,
                )
                df = df[~nao_recebido].copy()
                if df.empty:
                    continue

            df["valor_bruto"] = (
                df[col_valor_bruto].apply(parse_valor_br)
                if col_valor_bruto else df[col_valor].apply(parse_valor_br)
            )
            df["valor"] = df[col_valor].apply(parse_valor_br)
            df["valor_recebido"] = df[col_valor].apply(parse_valor_br)
            df["data"] = pd.to_datetime(
                df[col_data],
                dayfirst=True,
                errors="coerce",
            )
            df["descricao"] = (
                df[col_descricao].fillna("").astype(str).str.strip()
                if col_descricao else "Recebimento Amigotech"
            )
            df.loc[df["descricao"] == "", "descricao"] = "Recebimento Amigotech"
            df["status"] = (
                df[col_status].fillna("").astype(str).str.strip()
                if col_status else "Recebido"
            )
            df["conta_destino"] = (
                df[col_conta].fillna("").astype(str).str.strip()
                if col_conta else ""
            )
            df["forma"] = (
                df[col_forma].fillna("").astype(str).str.strip()
                if col_forma else ""
            )
            df["categoria"] = (
                df[col_categoria].fillna("").astype(str).str.strip()
                if col_categoria else ""
            )
            df["classificacao"] = (
                df[col_classificacao].fillna("").astype(str).str.strip()
                if col_classificacao else ""
            )
            df["observacao"] = (
                df[col_observacao].fillna("").astype(str).str.strip()
                if col_observacao else ""
            )
            texto_categoria = (
                df["categoria"].fillna("").map(normalizar_texto)
                + " " + df["classificacao"].fillna("").map(normalizar_texto)
            )
            texto_movimento = (
                texto_categoria
                + " " + df["descricao"].fillna("").map(normalizar_texto)
                + " " + df["observacao"].fillna("").map(normalizar_texto)
            )
            tem_categoria_informada = texto_categoria.str.replace(
                "-", "", regex=False
            ).str.strip() != ""
            categoria_venda = texto_categoria.str.contains(
                r"receita\s+operacional|prestacao\s+de\s+servico|prestação\s+de\s+serviço",
                regex=True,
            )
            movimento_nao_venda = texto_movimento.str.contains(
                r"transferencia\s+entre\s+contas|transferência\s+entre\s+contas|"
                r"receita\s+financeira|rendimentos?\s+s/?aplicacao|"
                r"rendimentos?\s+s/?aplicação|"
                r"recebimento\s+rede|redecard|cielo|stone|getnet|pagseguro|"
                r"mercado\s+pago|\brede\s+(visa|mast|master|elo|amex)",
                regex=True,
            )
            df["venda_valida"] = (
                (~tem_categoria_informada | categoria_venda)
                & ~movimento_nao_venda
            )
            df["fonte_venda"] = "Amigotech"
            df["_linha_origem_amigotech"] = df.index
            df = df.dropna(subset=["valor", "data"])
            df = df[df["valor"] > 0]
            if not df.empty:
                quadros.append(df)

        if not quadros:
            st.error(
                "❌ Amigotech recebimentos precisa conter as colunas "
                "'Data de pagamento' e 'Valor Líquido R$'."
            )
            return pd.DataFrame()
        return pd.concat(quadros, ignore_index=True, sort=False).reset_index(drop=True)
    except Exception as e:
        logger.error("Erro ao processar recebimentos Amigotech: %s", e)
        st.error(f"❌ Erro ao processar recebimentos Amigotech: {str(e)}")
        return pd.DataFrame()


def processar_amigotech_pagar(uploaded_file) -> pd.DataFrame:
    try:
        quadros = []
        candidatos = ler_tabelas_amigotech(uploaded_file)

        for df in candidatos:
            mapa = {normalizar_texto(col): col for col in df.columns}

            def localizar(*nomes):
                nomes_norm = [normalizar_texto(nome) for nome in nomes]
                for nome in nomes_norm:
                    coluna = mapa.get(nome)
                    if coluna is not None:
                        return coluna
                for chave, coluna in mapa.items():
                    if any(nome in chave for nome in nomes_norm):
                        return coluna
                return None

            col_tipo = localizar("tipo")
            col_valor = localizar(
                "valor liquido r$", "valor líquido r$",
                "valor liquido", "valor líquido", "valor pago",
                "valor original", "valor", "total"
            )
            col_data_pagamento = localizar(
                "data pagamento", "data de pagamento", "pago em"
            )
            col_data_vencimento = localizar(
                "data de vencimento", "vencimento", "data"
            )
            col_grupo_custo = localizar("categoria", "custo", "tipo de custo")
            col_classificacao = localizar("classificacao", "classificação")
            col_descricao = localizar(
                "descricao", "descrição", "pago a", "recebido de",
                "fornecedor", "favorecido"
            )
            col_forma = localizar(
                "forma de pagamento", "forma pagamento",
                "meio de pagamento", "pagamento"
            )
            col_conta_pagamento = localizar("banco", "conta", "conta destino")
            col_status = localizar("status", "situacao", "situação")
            col_observacao = localizar("observacao", "observação", "obs")
            if not col_valor or not (col_data_pagamento or col_data_vencimento):
                continue

            if col_tipo:
                tipo_norm = df[col_tipo].fillna("").map(normalizar_texto)
                df = df[tipo_norm.str.contains("saida|saída", regex=True)].copy()
                if df.empty:
                    continue

            status_norm = (
                df[col_status].fillna("").map(normalizar_texto)
                if col_status else pd.Series("", index=df.index)
            )
            datas_pagamento = (
                pd.to_datetime(
                    df[col_data_pagamento],
                    dayfirst=True,
                    errors="coerce",
                )
                if col_data_pagamento else pd.Series(pd.NaT, index=df.index)
            )
            linha_paga = datas_pagamento.notna() | status_norm.str.contains(
                r"pago|paga|baixad|liquidad",
                regex=True,
            )
            linha_nao_paga = status_norm.str.contains(
                r"aberto|pendente|vencid|cancel|estorn|nao\s+pago|não\s+pago",
                regex=True,
            )
            df = df[linha_paga & ~linha_nao_paga].copy()
            datas_pagamento = datas_pagamento.loc[df.index]
            if df.empty:
                continue

            df["valor"] = df[col_valor].apply(parse_valor_br).abs()
            datas_vencimento = (
                pd.to_datetime(
                    df[col_data_vencimento],
                    dayfirst=True,
                    errors="coerce",
                )
                if col_data_vencimento else pd.Series(pd.NaT, index=df.index)
            )
            df["data"] = datas_pagamento.combine_first(datas_vencimento)
            grupo_original = (
                df[col_grupo_custo].fillna("").astype(str).str.strip()
                if col_grupo_custo else ""
            )
            classificacao = (
                df[col_classificacao].fillna("").astype(str).str.strip()
                if col_classificacao else ""
            )
            categoria = (
                classificacao
                if col_classificacao else
                grupo_original if col_grupo_custo else "Despesa Amigotech"
            )
            df["categoria"] = categoria
            descricao = (
                df[col_descricao].fillna("").astype(str).str.strip()
                if col_descricao else ""
            )
            df["descricao"] = descricao
            df.loc[df["descricao"] == "", "descricao"] = df["categoria"]
            forma = (
                df[col_forma].fillna("").astype(str).str.strip()
                if col_forma else ""
            )
            conta_pagamento = (
                df[col_conta_pagamento].fillna("").astype(str).str.strip()
                if col_conta_pagamento else ""
            )
            if col_forma and col_conta_pagamento:
                df["forma"] = (forma + " " + conta_pagamento).str.strip()
            elif col_forma:
                df["forma"] = forma
            elif col_conta_pagamento:
                df["forma"] = conta_pagamento
            else:
                df["forma"] = ""
            df["observacao"] = (
                df[col_observacao].fillna("").astype(str).str.strip()
                if col_observacao else ""
            )
            texto_financeiro = (
                df["categoria"].fillna("").map(normalizar_texto) + " "
                + df["descricao"].fillna("").map(normalizar_texto) + " "
                + df["observacao"].fillna("").map(normalizar_texto)
            )
            movimento_nao_despesa = texto_financeiro.str.contains(
                r"transferencia\s+entre\s+contas|transf\.?\s+entre\s+contas|"
                r"saldo|aplicacao|aplic\.|resgate|investimento|"
                r"movimentacao\s+conta",
                regex=True,
            )
            df = df[~movimento_nao_despesa].copy()
            df["tipo"] = grupo_original
            df.loc[df["tipo"].fillna("").astype(str).str.strip() == "", "tipo"] = (
                "Custo Variável"
            )
            df["numero_documento"] = ""
            df["nsu"] = ""
            df["estabelecimento"] = "Amigotech"
            df["fonte"] = "Amigotech"
            df["_linha_origem_amigotech"] = df.index
            df = df.dropna(subset=["valor", "data"])
            df = df[df["valor"] > 0]
            if not df.empty:
                df["grupo_custo"] = classificar_grupo_custo(df)
                quadros.append(df)

        if not quadros:
            st.error("❌ Amigotech contas a pagar precisa conter data e valor.")
            return pd.DataFrame()
        return pd.concat(quadros, ignore_index=True, sort=False).reset_index(drop=True)
    except Exception as e:
        logger.error("Erro ao processar contas a pagar Amigotech: %s", e)
        st.error(f"❌ Erro ao processar contas a pagar Amigotech: {str(e)}")
        return pd.DataFrame()


def processar_excel(uploaded_file) -> pd.DataFrame:
    try:
        df = ler_tabela_flexivel(uploaded_file)
        df = df.loc[:, ~df.columns.duplicated()].copy()
        mapa_colunas = {normalizar_texto(col): col for col in df.columns}

        def localizar(*opcoes):
            for opcao in opcoes:
                coluna = mapa_colunas.get(normalizar_texto(opcao))
                if coluna is not None:
                    return coluna
            return None

        col_valor = localizar("valor", "valor pago", "total")
        col_vencimento = localizar(
            "vencimento", "vcto.", "vcto", "data de vencimento", "data"
        )
        col_categoria = localizar(
            "categoria", "conta destino", "plano de contas", "conta"
        )
        col_descricao = localizar(
            "descrição", "descricao", "fornecedor", "favorecido"
        )
        col_forma = localizar("forma de pgto", "forma de pagamento", "pagamento")
        col_tipo = localizar("tipo", "tipo de custo")
        col_documento = localizar("nº doc", "no doc", "numero doc", "documento")
        col_nsu = localizar("nsu")
        col_estabelecimento = localizar("estabelecimento", "empresa", "unidade")

        faltantes = []
        if not col_valor:
            faltantes.append("valor")
        if not col_vencimento:
            faltantes.append("vencimento/Vcto.")
        if not col_categoria:
            faltantes.append("categoria/Conta Destino")
        if faltantes:
            st.error(
                "❌ Colunas obrigatórias não identificadas: "
                + ", ".join(faltantes)
            )
            return pd.DataFrame()

        df["valor"] = df[col_valor].apply(parse_valor_br)
        df = df.dropna(subset=["valor"])
        df = df[df["valor"] > 0]
        df["data"] = pd.to_datetime(
            df[col_vencimento],
            dayfirst=True,
            errors="coerce",
        )
        df = df.dropna(subset=["data"])
        df["categoria"] = (
            df[col_categoria]
            .fillna("sem categoria")
            .astype(str)
            .str.strip()
        )
        df["descricao"] = (
            df[col_descricao].fillna("").astype(str).str.strip()
            if col_descricao else ""
        )
        df.loc[df["descricao"] == "", "descricao"] = df["categoria"]
        df["forma"] = (
            df[col_forma].fillna("").astype(str).str.strip()
            if col_forma else ""
        )

        categorias_norm = df["categoria"].map(normalizar_texto)
        descricoes_norm = df["descricao"].map(normalizar_texto)
        conta_bancaria_destino = categorias_norm.str.contains(
            r"^\s*1\.1\.\d|^\s*1\.1\.1|sicoob|sicredi|itau|itaú|"
            r"bradesco|santander|caixa\s+economica|banco\s+do\s+brasil|"
            r"infinite\s*pay|infinitepay|maquininha",
            regex=True,
        ) & ~categorias_norm.str.contains(
            r"tarifa|taxa|juros|financiamento|emprestimo|empréstimo",
            regex=True,
        )
        descricao_transferencia = descricoes_norm.str.contains(
            r"transferencia\s+entre\s+contas|transf\.?\s+entre\s+contas|"
            r"movimentacao\s+conta|movimentação\s+conta|repasse\s+entre\s+contas",
            regex=True,
        )
        df = df[~(conta_bancaria_destino | descricao_transferencia)].copy()
        if df.empty:
            return pd.DataFrame()

        categorias_norm = df["categoria"].map(normalizar_texto)
        categorias_fixas = categorias_norm.str.contains(
            r"salario|pro-labore|aluguel|energia|agua|telefone|internet|"
            r"honorarios contabil|plano medico|software|sistema de gestao|"
            r"tarifa pacote|simples nacional|inss|fgts|relogio de ponto",
            regex=True,
        )
        df["tipo"] = (
            df[col_tipo].fillna("").astype(str).str.strip()
            if col_tipo else "Custo Variável"
        )
        if not col_tipo:
            df.loc[categorias_fixas, "tipo"] = "Custo Fixo"

        df["numero_documento"] = (
            df[col_documento].fillna("").astype(str).str.strip()
            if col_documento else ""
        )
        df["nsu"] = (
            df[col_nsu].fillna("").astype(str).str.strip()
            if col_nsu else ""
        )
        df["estabelecimento"] = (
            df[col_estabelecimento].fillna("").astype(str).str.strip()
            if col_estabelecimento else ""
        )
        df["grupo_custo"] = classificar_grupo_custo(df)
        logger.info(f"Excel processed: {len(df)} valid records")
        return df
    except Exception as e:
        logger.error(f"Error processing Excel: {e}")
        st.error(f"❌ Erro ao processar Excel: {str(e)}")
        return pd.DataFrame()


def localizar_coluna_por_termos(
    colunas,
    *termos,
    todos: bool = True,
) -> Optional[object]:
    for coluna in colunas:
        texto = normalizar_texto(coluna)
        encontrados = [normalizar_texto(termo) in texto for termo in termos]
        if (todos and all(encontrados)) or (not todos and any(encontrados)):
            return coluna
    return None


def preparar_relatorio_baixas_conta_azul(uploaded_file) -> pd.DataFrame:
    quadros = []
    for bruto in ler_todas_tabelas_flexiveis(uploaded_file):
        df = promover_cabecalho_por_termos(
            bruto,
            ["data da baixa", "valor"],
        )
        df = df.loc[:, ~df.columns.duplicated()].copy()
        df.columns = [str(col).strip() for col in df.columns]
        df = df.dropna(how="all")
        if df.empty:
            continue
        col_data = localizar_coluna_por_termos(df.columns, "data", "baixa")
        col_descricao = localizar_coluna_por_termos(
            df.columns,
            "parcela",
            todos=False,
        )
        col_conta = localizar_coluna_por_termos(df.columns, "conta", "financeira")
        col_valor = localizar_coluna_por_termos(df.columns, "valor")
        col_liquido = localizar_coluna_por_termos(
            df.columns,
            "baixado",
            todos=False,
        )
        col_juros = localizar_coluna_por_termos(df.columns, "juros")
        col_multa = localizar_coluna_por_termos(df.columns, "multa")
        col_desconto = localizar_coluna_por_termos(df.columns, "desconto")
        col_tarifa = localizar_coluna_por_termos(df.columns, "tarifa")
        if not col_valor:
            continue
        total_geral = pd.Series(False, index=df.index)
        if col_data:
            total_geral = total_geral | df[col_data].astype(str).map(normalizar_texto).eq("total geral")
        if col_descricao:
            total_geral = total_geral | df[col_descricao].astype(str).map(normalizar_texto).eq("total geral")
        total_geral = total_geral | df.astype(str).apply(
            lambda linha: linha.map(normalizar_texto).eq("total geral").any(),
            axis=1,
        )
        df = df[~total_geral].copy()
        if col_data:
            df[col_data] = pd.to_datetime(df[col_data], dayfirst=True, errors="coerce")
            df[col_data] = df[col_data].ffill()

        df["data"] = df[col_data] if col_data else pd.NaT
        df["descricao"] = (
            df[col_descricao].fillna("").astype(str).str.strip()
            if col_descricao else ""
        )
        df["conta_destino"] = (
            df[col_conta].fillna("").astype(str).str.strip()
            if col_conta else ""
        )
        df["valor"] = df[col_valor].apply(parse_valor_br)
        df["valor_juros"] = (
            df[col_juros].apply(parse_valor_br).abs()
            if col_juros else 0.0
        )
        df["valor_multa"] = (
            df[col_multa].apply(parse_valor_br).abs()
            if col_multa else 0.0
        )
        df["valor_liquido"] = (
            df[col_liquido].apply(parse_valor_br)
            if col_liquido else df["valor"]
        )
        df["valor_desconto"] = (
            df[col_desconto].apply(parse_valor_br).abs()
            if col_desconto else 0.0
        )
        df["valor_tarifa"] = (
            df[col_tarifa].apply(parse_valor_br).abs()
            if col_tarifa else 0.0
        )
        df = df.dropna(subset=["data", "valor"])
        if not df.empty:
            quadros.append(df)

    if not quadros:
        return pd.DataFrame()
    return pd.concat(quadros, ignore_index=True, sort=False).reset_index(drop=True)


def preparar_visao_conta_azul(uploaded_file, tipo: str) -> pd.DataFrame:
    quadros = []
    for df in ler_todas_tabelas_flexiveis(uploaded_file):
        df = df.loc[:, ~df.columns.duplicated()].copy()
        df.columns = [str(col).strip() for col in df.columns]
        df = df.dropna(how="all")
        if df.empty:
            continue
        colunas = list(df.columns)
        col_conta = localizar_coluna_por_termos(colunas, "conta", "banc")
        col_data = localizar_coluna_por_termos(colunas, "ultimo", "pagamento")
        if col_data is None:
            col_data = localizar_coluna_por_termos(colunas, "data", "prevista")
        if col_data is None:
            col_data = localizar_coluna_por_termos(colunas, "data", "vencimento")
        col_situacao = localizar_coluna_por_termos(colunas, "situ")
        col_categoria = localizar_coluna_por_termos(colunas, "categoria", "1")

        if tipo == "receber":
            col_nome = localizar_coluna_por_termos(colunas, "nome", "cliente")
            col_valor = (
                localizar_coluna_por_termos(colunas, "total", "recebido")
                or localizar_coluna_por_termos(colunas, "valor", "recebido")
            )
            col_desconto = localizar_coluna_por_termos(colunas, "desconto")
            col_tarifa = localizar_coluna_por_termos(colunas, "tarifa")
            col_juros = localizar_coluna_por_termos(colunas, "juros")
            col_multa = localizar_coluna_por_termos(colunas, "multa")
        else:
            col_nome = localizar_coluna_por_termos(colunas, "nome", "fornecedor")
            col_valor = (
                localizar_coluna_por_termos(colunas, "total", "pago")
                or localizar_coluna_por_termos(colunas, "valor", "pago")
            )
            col_desconto = None
            col_tarifa = None
            col_juros = None
            col_multa = None
        col_descricao = localizar_coluna_por_termos(
            colunas,
            "descri",
            todos=False,
        )
        if not col_valor or not col_data:
            continue

        resultado = pd.DataFrame(index=df.index)
        resultado["data"] = pd.to_datetime(
            df[col_data],
            dayfirst=True,
            errors="coerce",
        )
        resultado["valor"] = df[col_valor].apply(parse_valor_br).abs()
        resultado["valor_juros"] = (
            df[col_juros].apply(parse_valor_br).abs()
            if col_juros else 0.0
        )
        resultado["valor_multa"] = (
            df[col_multa].apply(parse_valor_br).abs()
            if col_multa else 0.0
        )
        resultado["valor_desconto"] = (
            df[col_desconto].apply(parse_valor_br).abs()
            if col_desconto else 0.0
        )
        resultado["valor_tarifa"] = (
            df[col_tarifa].apply(parse_valor_br).abs()
            if col_tarifa else 0.0
        )
        descricao = (
            df[col_descricao].fillna("").astype(str).str.strip()
            if col_descricao else pd.Series("", index=df.index)
        )
        nome = (
            df[col_nome].fillna("").astype(str).str.strip()
            if col_nome else pd.Series("", index=df.index)
        )
        resultado["descricao"] = descricao.mask(descricao == "", nome)
        resultado.loc[resultado["descricao"] == "", "descricao"] = (
            "Conta Azul"
        )
        if tipo == "pagar":
            resultado["fornecedor"] = nome
        elif tipo == "receber":
            resultado["cliente"] = nome
        resultado["conta_destino"] = (
            df[col_conta].fillna("").astype(str).str.strip()
            if col_conta else ""
        )
        resultado["categoria"] = (
            df[col_categoria].fillna("").astype(str).str.strip()
            if col_categoria else resultado["descricao"]
        )
        resultado["status"] = (
            df[col_situacao].fillna("").astype(str).str.strip()
            if col_situacao else ""
        )
        resultado = resultado.dropna(subset=["data", "valor"])
        resultado = resultado[resultado["valor"] > 0]
        if tipo == "pagar" and "status" in resultado.columns:
            status_norm = resultado["status"].map(normalizar_texto)
            quitado = status_norm.str.contains(
                r"quitado|recebido|pago|baixado",
                regex=True,
            )
            if quitado.any():
                resultado = resultado[quitado].copy()
        if not resultado.empty:
            quadros.append(resultado)

    if not quadros:
        return pd.DataFrame()
    return pd.concat(quadros, ignore_index=True, sort=False).reset_index(drop=True)


def localizar_valor_venda_conta_azul(colunas) -> Optional[object]:
    candidatos_liquido_venda = []
    candidatos_total = []
    candidatos_financeiro = []
    candidatos_genericos = []
    candidatos_bruto = []
    for coluna in colunas:
        texto = normalizar_texto(coluna)
        if "valor" not in texto:
            continue
        if "frete" in texto or "desconto" in texto:
            continue
        if "bruto" in texto:
            candidatos_bruto.append(coluna)
            continue
        if "financeiro" in texto:
            candidatos_financeiro.append(coluna)
            continue
        if "liqu" in texto or "quido" in texto or "l�quido" in str(coluna).lower():
            candidatos_liquido_venda.append(coluna)
            continue
        if "total" in texto:
            candidatos_total.append(coluna)
            continue
        candidatos_genericos.append(coluna)
    if candidatos_liquido_venda:
        return candidatos_liquido_venda[-1]
    if candidatos_total:
        return candidatos_total[-1]
    if candidatos_financeiro:
        return candidatos_financeiro[-1]
    if candidatos_genericos:
        return candidatos_genericos[-1]
    if candidatos_bruto:
        return candidatos_bruto[-1]
    return None


def processar_conta_azul_receber(uploaded_file) -> pd.DataFrame:
    try:
        df = preparar_visao_conta_azul(uploaded_file, "receber")
        if df.empty:
            df = preparar_relatorio_baixas_conta_azul(uploaded_file)
        if df.empty:
            st.error("❌ Contas a receber do Conta Azul não trouxe registros válidos.")
            return pd.DataFrame()
        df["valor"] = pd.to_numeric(df["valor"], errors="coerce").fillna(0).abs()
        for coluna in ["valor_juros", "valor_multa"]:
            if coluna not in df.columns:
                df[coluna] = 0.0
            df[coluna] = pd.to_numeric(df[coluna], errors="coerce").fillna(0).abs()
        df["valor"] = df["valor"] + df["valor_juros"] + df["valor_multa"]
        df["valor_recebido"] = df["valor"]
        df["status"] = "Baixado"
        df["forma"] = df["conta_destino"]
        df["memo"] = df["descricao"]
        df["tipo_recebimento"] = "Conta Azul"
        df["fonte_venda"] = "Conta Azul"
        df["fonte"] = "Conta Azul - Contas a receber"
        for coluna in ["valor_desconto", "valor_tarifa"]:
            if coluna not in df.columns:
                df[coluna] = 0.0
            df[coluna] = pd.to_numeric(df[coluna], errors="coerce").fillna(0).abs()
        return df[df["valor"] > 0].reset_index(drop=True)
    except Exception as e:
        logger.error("Erro ao processar contas a receber Conta Azul: %s", e)
        st.error(f"❌ Erro ao processar contas a receber Conta Azul: {str(e)}")
        return pd.DataFrame()


def processar_conta_azul_pagar(uploaded_file) -> pd.DataFrame:
    try:
        df = preparar_visao_conta_azul(uploaded_file, "pagar")
        if df.empty:
            df = preparar_relatorio_baixas_conta_azul(uploaded_file)
        if df.empty:
            st.error("❌ Contas a pagar do Conta Azul não trouxe registros válidos.")
            return pd.DataFrame()
        if "valor_liquido" in df.columns:
            df["valor"] = pd.to_numeric(df["valor_liquido"], errors="coerce").fillna(
                pd.to_numeric(df["valor"], errors="coerce")
            ).abs()
        else:
            df["valor"] = pd.to_numeric(df["valor"], errors="coerce").abs()
        if "categoria" not in df.columns:
            df["categoria"] = df["descricao"].replace("", "Conta Azul")
        df["forma"] = df["conta_destino"]
        df["tipo"] = "Custo Variável"
        df["fonte"] = "Conta Azul - Contas a pagar"
        df = df[df["valor"] > 0].copy()
        descricao_norm = df["descricao"].fillna("").map(normalizar_texto)
        df = df[~descricao_norm.str.contains(
            r"recebimento\s+conta\s+pf|recebimento\s+pf",
            regex=True,
        )].copy()
        if not df.empty:
            df["grupo_custo"] = classificar_grupo_custo(df)
        return df.reset_index(drop=True)
    except Exception as e:
        logger.error("Erro ao processar contas a pagar Conta Azul: %s", e)
        st.error(f"❌ Erro ao processar contas a pagar Conta Azul: {str(e)}")
        return pd.DataFrame()


def processar_conta_azul_vendas(uploaded_file) -> pd.DataFrame:
    try:
        quadros = []
        for df in ler_todas_tabelas_flexiveis(uploaded_file):
            df = df.loc[:, ~df.columns.duplicated()].copy()
            df.columns = [str(col).strip() for col in df.columns]
            mapa = {normalizar_texto(col): col for col in df.columns}

            def localizar(*opcoes):
                opcoes_norm = [normalizar_texto(opcao) for opcao in opcoes]
                for opcao in opcoes:
                    coluna = mapa.get(normalizar_texto(opcao))
                    if coluna is not None:
                        return coluna
                for chave, coluna in mapa.items():
                    if any(opcao in chave for opcao in opcoes_norm):
                        return coluna
                return None

            col_numero = localizar("número da venda", "numero da venda")
            col_cliente = localizar("cliente")
            col_data = localizar("data da venda", "data")
            col_tipo_item = localizar("tipo de item", "tipo item")
            col_servico = localizar(
                "serviço", "servico", "produto", "descrição",
                "descricao", "categoria", "nome do produto", "nome do serviço",
                "nome do servico", "nome do item", "item vendido"
            )
            col_valor = localizar_valor_venda_conta_azul(df.columns)
            if not col_data or not col_valor:
                continue
            if col_numero:
                total_geral = (
                    df[col_numero]
                    .fillna("")
                    .astype(str)
                    .map(normalizar_texto)
                    .eq("total geral")
                )
                df = df[~total_geral].copy()
            df["data"] = pd.to_datetime(df[col_data], dayfirst=True, errors="coerce")
            df["valor"] = df[col_valor].apply(parse_valor_br)
            df["valor_liquido"] = df["valor"]
            cliente = (
                df[col_cliente].fillna("").astype(str).str.strip()
                if col_cliente else pd.Series("", index=df.index)
            )
            numero = (
                df[col_numero].fillna("").astype(str).str.strip()
                if col_numero else pd.Series("", index=df.index)
            )
            servico = (
                df[col_servico].fillna("").astype(str).str.strip()
                if col_servico else pd.Series("", index=df.index)
            )
            tipo_item = (
                df[col_tipo_item].fillna("").astype(str).str.strip()
                if col_tipo_item else pd.Series("", index=df.index)
            )
            df["descricao"] = (
                "Venda " + numero.replace("nan", "") + " - " + cliente
                + " - " + servico.replace("nan", "")
            ).str.strip(" -")
            df.loc[df["descricao"] == "", "descricao"] = "Venda Conta Azul"
            df["servico_vendido"] = servico
            df.loc[df["servico_vendido"] == "", "servico_vendido"] = (
                "Serviço não identificado"
            )
            df["tipo_item_vendido"] = tipo_item
            df["status"] = "Aprovada"
            df["fonte_venda"] = "Conta Azul"
            df = df.dropna(subset=["data", "valor"])
            df = df[df["valor"] > 0]
            if not df.empty:
                quadros.append(df)

        if not quadros:
            st.error("❌ Vendas do Conta Azul precisa conter data e valor.")
            return pd.DataFrame()
        return pd.concat(quadros, ignore_index=True, sort=False).reset_index(drop=True)
    except Exception as e:
        logger.error("Erro ao processar vendas Conta Azul: %s", e)
        st.error(f"❌ Erro ao processar vendas Conta Azul: {str(e)}")
        return pd.DataFrame()


def processar_gerencial_resultados_belle(uploaded_file) -> pd.DataFrame:
    try:
        import re

        candidatos = []
        candidatos.extend(ler_todas_tabelas_flexiveis(uploaded_file))
        candidatos.extend(ler_todas_tabelas_flexiveis(uploaded_file, header=None))
        candidatos = [
            df.dropna(how="all").copy()
            for df in candidatos
            if df is not None and not df.empty
        ]
        if not candidatos:
            st.warning(f"⚠️ {uploaded_file.name}: nenhum registro encontrado.")
            return pd.DataFrame()

        alvo = normalizar_texto(
            "DESPESAS COM OPE. DE CARTÃO DE CRÉDITO"
        )
        alvo_vendas = normalizar_texto("Venda de Planos")

        def valores_numericos_linha(linha: pd.Series) -> List[float]:
            valores = []
            for coluna, valor_original in linha.items():
                valor = parse_valor_br(valor_original)
                if pd.isna(valor) or float(valor) == 0:
                    continue
                nome_coluna = normalizar_texto(coluna)
                if any(
                    termo in nome_coluna
                    for termo in [
                        "codigo", "cod", "conta", "descricao",
                        "historico", "grupo", "classificacao",
                    ]
                ):
                    continue
                valores.append(float(valor))
            return valores

        def valor_proximo_da_conta(linha: pd.Series, alvo_conta: str) -> float:
            valores = list(linha.tolist())
            for pos, valor in enumerate(valores):
                texto = normalizar_texto(valor)
                if texto != alvo_conta and alvo_conta not in texto:
                    continue

                for passo in range(1, len(valores)):
                    for idx in (pos + passo, pos - passo):
                        if idx < 0 or idx >= len(valores):
                            continue
                        valor_parseado = parse_valor_br(valores[idx])
                        if pd.notna(valor_parseado) and float(valor_parseado) != 0:
                            return abs(float(valor_parseado))
            return 0.0

        def linha_eh_taxa_cartao(linha: pd.Series) -> bool:
            texto = " ".join(
                normalizar_texto(valor)
                for valor in linha.tolist()
                if str(valor).strip() and str(valor).lower() != "nan"
            )
            return (
                alvo in texto
                or (
                    "despesas com" in texto
                    and "cartao" in texto
                    and ("ope" in texto or "operacao" in texto)
                )
            )

        registros = []
        encontrou_conta = False
        vendas_gerencial = 0.0
        meses = {
            "janeiro": 1,
            "fevereiro": 2,
            "marco": 3,
            "abril": 4,
            "maio": 5,
            "junho": 6,
            "julho": 7,
            "agosto": 8,
            "setembro": 9,
            "outubro": 10,
            "novembro": 11,
            "dezembro": 12,
        }
        data_padrao = pd.Timestamp(
            st.session_state.get("periodo_inicio_global", date.today())
        )
        for df in candidatos:
            texto_arquivo = " ".join(
                normalizar_texto(valor)
                for valor in df.astype(str).to_numpy().ravel().tolist()
                if str(valor).strip()
                and str(valor).lower() != "nan"
            )
            match = re.search(
                r"mensal\s*:\s*([a-z]+).*?ano\s*:\s*(20\d{2})",
                texto_arquivo,
            )
            if match and match.group(1) in meses:
                data_padrao = pd.Timestamp(
                    int(match.group(2)),
                    meses[match.group(1)],
                    1,
                )
                break

        for df in candidatos:
            df = df.loc[:, ~df.columns.duplicated()].copy()
            linhas_vendas = df[
                df.apply(
                    lambda linha: alvo_vendas in " ".join(
                        normalizar_texto(valor)
                        for valor in linha.tolist()
                        if str(valor).strip()
                        and str(valor).lower() != "nan"
                    ),
                    axis=1,
                )
            ].copy()
            for _, linha_venda in linhas_vendas.iterrows():
                valor_venda = valor_proximo_da_conta(
                    linha_venda,
                    alvo_vendas,
                )
                if valor_venda > 0:
                    vendas_gerencial = max(
                        vendas_gerencial,
                        valor_venda,
                    )

            linhas = df[df.apply(linha_eh_taxa_cartao, axis=1)].copy()
            if linhas.empty:
                continue

            encontrou_conta = True
            colunas_data = [
                coluna for coluna in df.columns
                if normalizar_texto(coluna) in {"data", "competencia", "mes"}
            ]

            for _, linha in linhas.iterrows():
                valores = valores_numericos_linha(linha)

                if not valores:
                    continue
                valor_lancamento = max(abs(valor) for valor in valores)
                recebimento_liquido = (
                    vendas_gerencial - valor_lancamento
                    if vendas_gerencial > 0 else 0.0
                )

                data_lancamento = data_padrao
                for coluna in colunas_data:
                    data_candidata = pd.to_datetime(
                        linha[coluna], dayfirst=True, errors="coerce"
                    )
                    if not pd.isna(data_candidata):
                        data_lancamento = data_candidata
                        break

                registros.append({
                    "data": data_lancamento,
                    "valor": valor_lancamento,
                    "categoria": "DESPESAS COM OPE. DE CARTÃO DE CRÉDITO",
                    "descricao": "Gerencial de resultados Belle",
                    "forma": "Cartão de crédito",
                    "tipo": "Custo Variável",
                    "numero_documento": "",
                    "nsu": "",
                    "estabelecimento": "",
                    "grupo_custo": "Custo Variável",
                    "fonte": "Gerencial de resultados Belle",
                    "vendas_gerencial": vendas_gerencial,
                    "recebimento_liquido_gerencial": recebimento_liquido,
                })

        if not encontrou_conta:
            st.warning(
                f"⚠️ {uploaded_file.name}: conta 'DESPESAS COM OPE. DE "
                "CARTÃO DE CRÉDITO' não encontrada."
            )
            return pd.DataFrame()

        resultado = pd.DataFrame(registros).drop_duplicates()
        if resultado.empty:
            st.warning(
                f"⚠️ {uploaded_file.name}: a conta foi encontrada, mas sem "
                "valor numérico aproveitável."
            )
            return pd.DataFrame()

        return resultado
    except Exception as e:
        logger.error("Erro ao processar gerencial Belle: %s", e)
        st.error(f"❌ Erro ao processar gerencial de resultados: {str(e)}")
        return pd.DataFrame()


def processar_orcamentos(uploaded_file) -> pd.DataFrame:
    try:
        df = ler_tabela_flexivel(uploaded_file)
        df = df.loc[:, ~df.columns.duplicated()].copy()
        mapa_colunas = {normalizar_texto(col): col for col in df.columns}

        def localizar(*opcoes):
            for opcao in opcoes:
                coluna = mapa_colunas.get(normalizar_texto(opcao))
                if coluna is not None:
                    return coluna
            return None

        col_status = localizar("status", "situação", "situacao", "estado")
        col_valor = localizar(
            "valor total com desconto", "valor aprovado", "valor total",
            "valor", "total",
            "valor do orçamento", "valor do orcamento"
        )
        col_data = localizar(
            "data aprovação", "data de aprovação", "data aprovacao",
            "data de aprovacao", "data", "data criação", "data criacao",
            "emissão", "emissao",
            "data do orçamento", "data do orcamento", "criado em"
        )
        col_cliente = localizar(
            "paciente", "cliente", "nome do cliente", "nome do paciente"
        )
        col_procedimentos = localizar(
            "tipo de procedimento", "tipo procedimento",
            "procedimentos", "procedimento", "descrição", "descricao",
            "orçamento", "orcamento", "projeto"
        )

        faltantes = []
        if not col_status:
            faltantes.append("status/situação")
        if not col_valor:
            faltantes.append("valor/valor total")
        if faltantes:
            st.error(
                "❌ A planilha de orçamentos precisa das colunas: "
                + ", ".join(faltantes)
            )
            return pd.DataFrame()

        status_normalizado = df[col_status].fillna("").map(normalizar_texto)
        status_aprovado = status_normalizado.str.contains(
            r"aprovad|approved", regex=True
        )
        df = df[status_aprovado].copy()
        if df.empty:
            encontrados = ", ".join(
                sorted(status_normalizado[status_normalizado != ""].unique())[:8]
            )
            st.warning(
                "⚠️ Nenhum orçamento com status aprovado foi encontrado."
                + (f" Status identificados: {encontrados}." if encontrados else "")
            )
            return pd.DataFrame()
        df["status"] = df[col_status].fillna("").astype(str).str.strip()
        df["valor"] = df[col_valor].apply(parse_valor_br)
        df["valor_liquido"] = df["valor"]
        df = df.dropna(subset=["valor"])
        df = df[df["valor"] > 0]
        df["data"] = (
            pd.to_datetime(df[col_data], dayfirst=True, errors="coerce")
            if col_data else pd.NaT
        )
        if not col_data:
            st.warning(
                "⚠️ A planilha não possui coluna de data. As vendas serão listadas "
                "em Detalhes, mas não entrarão no filtro da Visão Financeira."
            )
        clientes = (
            df[col_cliente].fillna("").astype(str).str.strip()
            if col_cliente else pd.Series("", index=df.index)
        )
        procedimentos = (
            df[col_procedimentos].fillna("").astype(str).str.strip()
            if col_procedimentos else pd.Series("", index=df.index)
        )
        df["descricao"] = (
            clientes + " — " + procedimentos
        ).str.strip(" —")
        df.loc[df["descricao"] == "", "descricao"] = "Orçamento aprovado"
        df["servico_vendido"] = procedimentos
        df.loc[df["servico_vendido"] == "", "servico_vendido"] = (
            "Serviço não identificado"
        )
        df["fonte_venda"] = "Clinicorp"
        logger.info("Approved budgets processed: %s", len(df))
        return df[[
            "data", "descricao", "valor", "valor_liquido", "status", "fonte_venda",
            "servico_vendido",
        ]]
    except Exception as e:
        logger.error("Error processing budgets: %s", e)
        st.error(f"❌ Erro ao processar orçamentos: {str(e)}")
        return pd.DataFrame()


def processar_clinipay(uploaded_file) -> pd.DataFrame:
    try:
        df = ler_tabela_flexivel(uploaded_file)
        df.columns = [normalizar_texto(col) for col in df.columns]
        df = df.loc[:, ~df.columns.duplicated()].copy()

        def localizar(*termos):
            for termo in termos:
                termo_norm = normalizar_texto(termo)
                for coluna in df.columns:
                    if termo_norm == coluna or termo_norm in coluna:
                        return coluna
            return None

        col_forma = localizar("forma")
        col_valor_bruto = localizar("valor")
        col_valor_liquido = localizar("vl final", "valor final", "valor liquido", "valor líquido")
        col_taxas = localizar("taxas", "taxa")
        if col_forma and col_valor_bruto and col_valor_liquido and col_taxas:
            col_data = localizar("recebimento", "pagamento", "data")
            col_paciente = localizar("paciente", "cliente")
            col_status = localizar("status")
            col_transacao = localizar("transa", "transacao", "transação")
            col_parcela = localizar("parcela")
            col_juros = localizar("multa/juros", "juros", "multa")

            if col_status:
                status_norm = df[col_status].fillna("").map(normalizar_texto)
                df = df[status_norm.str.contains("liquidado", regex=False)].copy()
            if df.empty:
                return pd.DataFrame()

            df["data"] = (
                pd.to_datetime(df[col_data], dayfirst=True, errors="coerce")
                if col_data else pd.NaT
            )
            df["valor_bruto_clinipay"] = df[col_valor_bruto].apply(parse_valor_br)
            df["valor"] = df[col_valor_liquido].apply(parse_valor_br)
            df["taxa_clinipay"] = df[col_taxas].apply(parse_valor_br).fillna(0).abs()
            df["juros_clinipay"] = (
                df[col_juros].apply(parse_valor_br).fillna(0).abs()
                if col_juros else 0.0
            )
            df = df.dropna(subset=["data", "valor_bruto_clinipay", "valor"])
            df = df[df["valor_bruto_clinipay"] > 0]
            forma = df[col_forma].fillna("Clinipay").astype(str).str.strip()
            paciente = (
                df[col_paciente].fillna("").astype(str).str.strip()
                if col_paciente else pd.Series("", index=df.index)
            )
            transacao = (
                df[col_transacao].fillna("").astype(str).str.strip()
                if col_transacao else pd.Series("", index=df.index)
            )
            df["memo"] = (forma + " - " + paciente).str.strip(" -")
            df.loc[df["memo"] == "", "memo"] = "Recebimento Clinipay"
            df["forma_clinipay"] = forma
            df["status_clinipay"] = (
                df[col_status].fillna("").astype(str).str.strip()
                if col_status else ""
            )
            df["parcela_clinipay"] = (
                df[col_parcela].fillna("").astype(str).str.strip()
                if col_parcela else ""
            )
            df["transacao_clinipay"] = transacao
            df["tipo_ofx"] = "CLINIPAY"
            df["fonte"] = "Clinipay"
            df["_tipo_relatorio_clinipay"] = "recebiveis"
            df["percentual_taxa_clinipay"] = (
                df["taxa_clinipay"] / df["valor_bruto_clinipay"] * 100
            ).where(df["valor_bruto_clinipay"] > 0, 0)
            base_juros = df["valor_bruto_clinipay"] - df["juros_clinipay"]
            df["percentual_juros_clinipay"] = (
                df["juros_clinipay"] / base_juros * 100
            ).where(base_juros > 0, 0)
            logger.info("Clinipay detailed entries processed: %s", len(df))
            return df[
                [
                    "data", "valor", "memo", "tipo_ofx", "fonte",
                    "valor_bruto_clinipay", "taxa_clinipay",
                    "juros_clinipay", "percentual_taxa_clinipay",
                    "percentual_juros_clinipay", "forma_clinipay",
                    "status_clinipay", "parcela_clinipay",
                    "transacao_clinipay", "_tipo_relatorio_clinipay",
                ]
            ]

        obrigatorias = ["data", "tipo", "descricao", "valor"]
        faltantes = [col for col in obrigatorias if col not in df.columns]
        if faltantes:
            st.error(
                "❌ O extrato Clinipay precisa das colunas: "
                + ", ".join(faltantes)
            )
            return pd.DataFrame()

        tipos = df["tipo"].fillna("").map(normalizar_texto)
        col_status_simples = localizar("status", "situacao", "situação", "estado")
        if col_status_simples:
            status_norm = df[col_status_simples].fillna("").map(normalizar_texto)
            status_valido = status_norm.str.contains(
                r"liquidado|recebido|pago|confirmado",
                regex=True,
            )
            status_invalidado = status_norm.str.contains(
                r"a\s+receber|vencid|pendente|aberto|cancelad|estorn",
                regex=True,
            )
            df = df[status_valido & ~status_invalidado].copy()
        df["data"] = pd.to_datetime(df["data"], dayfirst=True, errors="coerce")
        df["valor"] = df["valor"].apply(parse_valor_br)
        df = df.dropna(subset=["data", "valor"])
        tipos = df["tipo"].fillna("").map(normalizar_texto)
        saida_extrato = (
            df["valor"].lt(0)
            | tipos.str.contains(
                r"saida|saque|debito|tarifa|taxa|custo|estorno",
                regex=True,
            )
        )
        df = df[saida_extrato].copy()
        if df.empty:
            return pd.DataFrame()
        df["memo"] = df["descricao"].fillna("Recebimento Clinipay").astype(str)
        df["tipo_ofx"] = "CLINIPAY"
        df["fonte"] = "Clinipay"
        df["valor_bruto_clinipay"] = 0.0
        df["taxa_clinipay"] = 0.0
        df["juros_clinipay"] = 0.0
        df["percentual_taxa_clinipay"] = 0.0
        df["percentual_juros_clinipay"] = 0.0
        df["forma_clinipay"] = "Extrato Clinipay"
        df["status_clinipay"] = ""
        df["parcela_clinipay"] = ""
        df["transacao_clinipay"] = ""
        df["_tipo_relatorio_clinipay"] = "extrato"
        logger.info("Clinipay statement outflows processed: %s", len(df))
        return df[
            [
                "data", "valor", "memo", "tipo_ofx", "fonte",
                "valor_bruto_clinipay", "taxa_clinipay",
                "juros_clinipay", "percentual_taxa_clinipay",
                "percentual_juros_clinipay", "forma_clinipay",
                "status_clinipay", "parcela_clinipay",
                "transacao_clinipay", "_tipo_relatorio_clinipay",
            ]
        ]
    except Exception as e:
        logger.error("Error processing Clinipay: %s", e)
        st.error(f"❌ Erro ao processar Clinipay: {str(e)}")
        return pd.DataFrame()


def processar_fluxo_caixa(uploaded_file) -> pd.DataFrame:
    """Extrai apenas pagamentos diretos, evitando duplicar PIX e cartões."""
    try:
        df = ler_tabela_flexivel(uploaded_file)
        df.columns = [normalizar_texto(col) for col in df.columns]
        df = df.loc[:, ~df.columns.duplicated()].copy()

        obrigatorias = ["data", "descricao", "valor", "tipo"]
        faltantes = [col for col in obrigatorias if col not in df.columns]
        if faltantes:
            st.error(
                "❌ O Fluxo de Caixa precisa das colunas: "
                + ", ".join(faltantes)
            )
            return pd.DataFrame()

        tipos = df["tipo"].fillna("").map(normalizar_texto)
        descricoes = df["descricao"].fillna("").map(normalizar_texto)
        pagamentos_diretos = (
            tipos.eq("entrada")
            & descricoes.str.contains("pagamento de tratamento", regex=False)
        )
        df = df[pagamentos_diretos].copy()
        df["data"] = pd.to_datetime(df["data"], dayfirst=True, errors="coerce")
        df["valor"] = df["valor"].apply(parse_valor_br)
        df = df.dropna(subset=["data", "valor"])
        df = df[df["valor"] > 0]
        nome_coluna = "nome" if "nome" in df.columns else None
        df["memo"] = (
            "Pagamento direto — " + df[nome_coluna].fillna("").astype(str)
            if nome_coluna else "Pagamento direto de tratamento"
        )
        df["tipo_ofx"] = "PAGAMENTO_DIRETO"
        df["fonte"] = "Fluxo de Caixa"
        logger.info("Direct treatment payments processed: %s", len(df))
        return df[["data", "valor", "memo", "tipo_ofx", "fonte"]]
    except Exception as e:
        logger.error("Error processing cash flow: %s", e)
        st.error(f"❌ Erro ao processar Fluxo de Caixa: {str(e)}")
        return pd.DataFrame()


def processar_ofx(
    uploaded_file,
    fonte: str = "Extrato bancário",
) -> pd.DataFrame:
    try:
        import re

        raw_content = uploaded_file.read()
        try:
            content = raw_content.decode("utf-8-sig")
        except UnicodeDecodeError:
            content = raw_content.decode("latin-1")

        def tag_valor(bloco: str, tag: str) -> str:
            match = re.search(
                rf"<{tag}>\s*([^<\r\n]+)",
                bloco,
                flags=re.IGNORECASE,
            )
            return match.group(1).strip() if match else ""

        transacoes = []
        blocos = re.findall(
            r"<STMTTRN>(.*?)(?=<STMTTRN>|</BANKTRANLIST>|</STMTTRN>)",
            content,
            flags=re.IGNORECASE | re.DOTALL,
        )
        if not blocos:
            blocos = re.findall(
                r"<TRNTYPE>.*?(?=<TRNTYPE>|</BANKTRANLIST>)",
                content,
                flags=re.IGNORECASE | re.DOTALL,
            )
        for bloco in blocos:
            try:
                tipo_t = tag_valor(bloco, "TRNTYPE") or "OTHER"
                dt = tag_valor(bloco, "DTPOSTED")[:8]
                amt = tag_valor(bloco, "TRNAMT")
                memo = (
                    tag_valor(bloco, "MEMO")
                    or tag_valor(bloco, "NAME")
                    or tag_valor(bloco, "FITID")
                )
                if not dt or not amt:
                    continue
                data = datetime.strptime(dt, "%Y%m%d")
                valor = float(amt.replace(",", "."))
                if valor != 0:
                    transacoes.append({
                        "data": data,
                        "valor": valor,
                        "memo": memo,
                        "tipo_ofx": tipo_t,
                        "fonte": fonte,
                    })
            except ValueError as e:
                logger.warning(f"Skipping malformed OFX entry: {e}")
                continue
        saldo_match = re.search(r"<BALAMT>\s*([^<\r\n]+)", content, re.I)
        data_saldo_match = re.search(r"<DTASOF>\s*([^<\r\n]+)", content, re.I)
        conta_match = re.search(r"<ACCTID>\s*([^<\r\n]+)", content, re.I)
        banco_match = re.search(r"<BANKID>\s*([^<\r\n]+)", content, re.I)
        saldo_ofx = (
            float(saldo_match.group(1).strip().replace(",", "."))
            if saldo_match else None
        )
        data_saldo = (
            data_saldo_match.group(1).strip()[:8]
            if data_saldo_match else ""
        )
        df = pd.DataFrame(transacoes) if transacoes else pd.DataFrame()
        if not df.empty:
            df["_saldo_ofx"] = saldo_ofx
            df["_data_saldo_ofx"] = data_saldo
            df["_conta_ofx"] = (
                conta_match.group(1).strip() if conta_match else "Conta OFX"
            )
            df["_banco_ofx"] = nome_banco_brasileiro(
                banco_match.group(1).strip() if banco_match else "Banco"
            )
            df["_tipo_conta"] = "Conta bancária"
        logger.info(f"OFX processed: {len(transacoes)} valid transactions")
        return df
    except Exception as e:
        logger.error(f"Error processing OFX: {e}")
        st.error(f"❌ Erro ao processar OFX: {str(e)}")
        return pd.DataFrame()


def processar_excel_bancos(uploaded_file) -> pd.DataFrame:
    try:
        def ler_linhas_cruas() -> pd.DataFrame:
            extensao = Path(uploaded_file.name).suffix.lower()
            uploaded_file.seek(0)
            if extensao == ".csv":
                try:
                    return pd.read_csv(uploaded_file, header=None, sep=None, engine="python", encoding="utf-8-sig")
                except UnicodeDecodeError:
                    uploaded_file.seek(0)
                    return pd.read_csv(uploaded_file, header=None, sep=None, engine="python", encoding="latin-1")
            if extensao in {".xlsx", ".xls", ".ods"}:
                try:
                    return pd.read_excel(uploaded_file, header=None)
                except ImportError as e:
                    if extensao == ".xls" and "xlrd" in str(e).lower():
                        return converter_xls_com_excel(uploaded_file, header=None)
                    raise
            return ler_tabela_flexivel(uploaded_file)

        linhas = ler_linhas_cruas()
        if linhas.empty:
            return pd.DataFrame()
        linhas.columns = range(len(linhas.columns))

        def texto_linha(idx: int) -> str:
            return " ".join(
                normalizar_texto(valor)
                for valor in linhas.loc[idx].tolist()
                if pd.notna(valor)
            )

        def valor_linha_por_rotulo(padrao: str) -> Optional[float]:
            for idx in linhas.index:
                if re.search(padrao, texto_linha(idx), flags=re.I):
                    valores = [
                        parse_valor_br(valor)
                        for valor in linhas.loc[idx].tolist()[1:]
                    ]
                    valores = [
                        float(valor)
                        for valor in valores
                        if pd.notna(valor)
                    ]
                    if valores:
                        return valores[0]
            return None

        def texto_apos_rotulo(padrao: str, padrao_texto: str = "") -> str:
            for idx in linhas.index:
                texto = texto_linha(idx)
                if re.search(padrao, texto, flags=re.I):
                    for valor in linhas.loc[idx].tolist()[1:]:
                        if pd.notna(valor) and str(valor).strip():
                            return str(valor).strip()
            return padrao_texto

        banco = texto_apos_rotulo(r"cooperativa|banco", "Banco")
        conta = texto_apos_rotulo(r"\bconta\b", "Conta Excel")
        saldo_banco = valor_linha_por_rotulo(
            r"saldo\s+atual|saldo\s+disponivel|saldo\s+da\s+conta"
        )
        saldo_invest = valor_linha_por_rotulo(
            r"saldo\s+de\s+investimentos|investimentos\s+com\s+resgate|"
            r"saldo\s+aplic"
        )

        data_saldo = ""
        for idx in linhas.index:
            texto = texto_linha(idx)
            if "saldo da conta em" in texto:
                match = re.search(r"(\d{2}/\d{2}/\d{4})", str(linhas.loc[idx].tolist()))
                if match:
                    data_saldo = pd.to_datetime(
                        match.group(1),
                        dayfirst=True,
                        errors="coerce",
                    ).strftime("%Y%m%d")
                break
        if not data_saldo:
            match_nome = re.search(
                r"(\d{2})[-_/](\d{2})[-_/](\d{4})",
                uploaded_file.name,
            )
            if match_nome:
                data_saldo = f"{match_nome.group(3)}{match_nome.group(2)}{match_nome.group(1)}"

        linha_cabecalho = None
        for idx in linhas.index:
            texto = texto_linha(idx)
            if "data" in texto and "descri" in texto and "valor" in texto:
                linha_cabecalho = idx
                break

        quadros = []
        def chave_coluna_banco(valor) -> str:
            return re.sub(r"[^a-z0-9]", "", normalizar_texto(valor))

        for idx in linhas.index:
            chaves = [chave_coluna_banco(valor) for valor in linhas.loc[idx].tolist()]
            if not any(chave in {"saldoliquido", "saldobruto"} for chave in chaves):
                continue

            def indice_chave(*nomes):
                nomes_chave = {chave_coluna_banco(nome) for nome in nomes}
                for pos, chave in enumerate(chaves):
                    if chave in nomes_chave:
                        return pos
                return None

            col_saldo_posicao = indice_chave("SaldoLiquido", "Saldo Liquido")
            if col_saldo_posicao is None:
                col_saldo_posicao = indice_chave("SaldoBruto", "Saldo Bruto")
            col_conta_posicao = indice_chave("agContaDac", "Conta", "Agência Conta")
            col_tipo_posicao = indice_chave("Tipo")
            col_ativo_posicao = indice_chave("Ativo")

            registros_posicao = []
            for idx_pos in range(idx + 1, len(linhas)):
                linha = linhas.loc[idx_pos]
                saldo_posicao = (
                    parse_valor_br(linha.iloc[col_saldo_posicao])
                    if col_saldo_posicao is not None else None
                )
                if saldo_posicao is None or pd.isna(saldo_posicao):
                    continue
                conta_posicao = (
                    str(linha.iloc[col_conta_posicao]).strip()
                    if col_conta_posicao is not None and pd.notna(linha.iloc[col_conta_posicao])
                    else conta
                )
                tipo_posicao = (
                    str(linha.iloc[col_tipo_posicao]).strip()
                    if col_tipo_posicao is not None and pd.notna(linha.iloc[col_tipo_posicao])
                    else "Investimentos"
                )
                ativo_posicao = (
                    str(linha.iloc[col_ativo_posicao]).strip()
                    if col_ativo_posicao is not None and pd.notna(linha.iloc[col_ativo_posicao])
                    else ""
                )
                nome_investimento = " - ".join(
                    parte
                    for parte in [tipo_posicao, ativo_posicao, conta_posicao]
                    if parte
                )
                registros_posicao.append({
                    "data": pd.to_datetime(data_saldo, format="%Y%m%d", errors="coerce"),
                    "valor": 0.0,
                    "memo": "Posição consolidada de investimentos",
                    "tipo_ofx": "SALDO_INVESTIMENTO",
                    "fonte": "Bancos",
                    "_saldo_ofx": float(saldo_posicao),
                    "_data_saldo_ofx": data_saldo,
                    "_conta_ofx": nome_investimento or "Investimentos",
                    "_banco_ofx": banco,
                    "_tipo_conta": "Investimento",
                })
            if registros_posicao:
                quadros.append(pd.DataFrame(registros_posicao))
            break

        if saldo_banco is not None and pd.notna(saldo_banco):
            quadros.append(pd.DataFrame([{
                "data": pd.to_datetime(data_saldo, format="%Y%m%d", errors="coerce"),
                "valor": 0.0,
                "memo": "Saldo atual disponível em conta",
                "tipo_ofx": "SALDO_BANCO",
                "fonte": "Bancos",
                "_saldo_ofx": float(saldo_banco),
                "_data_saldo_ofx": data_saldo,
                "_conta_ofx": conta,
                "_banco_ofx": banco,
                "_tipo_conta": "Conta bancária",
            }]))

        cabecalho_lancamentos_futuros = False
        if linha_cabecalho is not None:
            inicio_contexto = max(0, linha_cabecalho - 3)
            contexto = " ".join(
                texto_linha(idx)
                for idx in range(inicio_contexto, linha_cabecalho + 1)
            )
            cabecalho_lancamentos_futuros = bool(
                re.search(r"lancamentos\s+futuros|proximos\s+\d+\s+dias", contexto)
            )

        if linha_cabecalho is not None and not cabecalho_lancamentos_futuros:
            cabecalhos = [
                normalizar_texto(valor)
                for valor in linhas.loc[linha_cabecalho].tolist()
            ]

            def indice_coluna(*termos):
                for pos, nome in enumerate(cabecalhos):
                    if any(termo in nome for termo in termos):
                        return pos
                return None

            col_data = indice_coluna("data")
            col_desc = indice_coluna("descri", "historico", "lançamento", "lancamento")
            col_nome = indice_coluna("nome", "cliente", "paciente")
            col_tipo_movimento = indice_coluna("tipo")
            col_tipo_lancamento = indice_coluna(
                "tipo lancamento",
                "tipo lançamento",
                "tipo lan",
            )
            col_detalhe = indice_coluna("detalhe", "detail")
            col_doc = indice_coluna("documento", "doc")
            col_valor = indice_coluna("valor")
            col_saldo = indice_coluna("saldo")
            eh_extrato_maquininha_excel = (
                col_tipo_movimento is not None
                and col_detalhe is not None
                and col_nome is not None
                and any("tipo de transacao" in nome for nome in cabecalhos)
                and any("hora" in nome for nome in cabecalhos)
            )
            registros = []
            for idx in range(linha_cabecalho + 1, len(linhas)):
                linha = linhas.loc[idx]
                data = (
                    pd.to_datetime(
                        linha.iloc[col_data],
                        dayfirst=True,
                        errors="coerce",
                    )
                    if col_data is not None else pd.NaT
                )
                if pd.isna(data):
                    continue
                valor = (
                    parse_valor_br(linha.iloc[col_valor])
                    if col_valor is not None else None
                )
                if valor is None or pd.isna(valor):
                    continue
                descricao = (
                    str(linha.iloc[col_desc]).strip()
                    if col_desc is not None and pd.notna(linha.iloc[col_desc])
                    else "Movimento bancário"
                )
                documento = (
                    str(linha.iloc[col_doc]).strip()
                    if col_doc is not None and pd.notna(linha.iloc[col_doc])
                    else ""
                )
                nome_lancamento = (
                    str(linha.iloc[col_nome]).strip()
                    if col_nome is not None and pd.notna(linha.iloc[col_nome])
                    else ""
                )
                detalhe_lancamento = (
                    str(linha.iloc[col_detalhe]).strip()
                    if col_detalhe is not None and pd.notna(linha.iloc[col_detalhe])
                    else ""
                )
                if descricao == "Movimento bancário" and (
                    nome_lancamento or detalhe_lancamento
                ):
                    descricao = " - ".join(
                        parte
                        for parte in [nome_lancamento, detalhe_lancamento]
                        if parte
                    )
                tipo_movimento = (
                    str(linha.iloc[col_tipo_movimento]).strip()
                    if col_tipo_movimento is not None and pd.notna(linha.iloc[col_tipo_movimento])
                    else ""
                )
                tipo_lancamento = (
                    str(linha.iloc[col_tipo_lancamento]).strip()
                    if col_tipo_lancamento is not None and pd.notna(linha.iloc[col_tipo_lancamento])
                    else ""
                )
                saldo = (
                    parse_valor_br(linha.iloc[col_saldo])
                    if col_saldo is not None else None
                )
                registros.append({
                    "data": data,
                    "valor": float(valor),
                    "memo": descricao,
                    "tipo_ofx": (
                        "MAQUININHA" if eh_extrato_maquininha_excel
                        else documento or "EXCEL"
                    ),
                    "fonte": (
                        "Maquininha" if eh_extrato_maquininha_excel
                        else "Bancos"
                    ),
                    "nome": nome_lancamento,
                    "detalhe": detalhe_lancamento,
                    "tipo_transacao": tipo_movimento,
                    "tipo_lancamento": tipo_lancamento,
                    "_arquivo_origem": getattr(uploaded_file, "name", ""),
                    "_saldo_ofx": saldo_banco,
                    "_data_saldo_ofx": data_saldo,
                    "_conta_ofx": conta,
                    "_banco_ofx": banco,
                    "_tipo_conta": (
                        "Maquininha" if eh_extrato_maquininha_excel
                        else "Conta bancária"
                    ),
                })
            if registros:
                quadros.append(pd.DataFrame(registros))

        if saldo_invest is not None and pd.notna(saldo_invest):
            quadros.append(pd.DataFrame([{
                "data": pd.to_datetime(data_saldo, format="%Y%m%d", errors="coerce"),
                "valor": 0.0,
                "memo": "Saldo de investimentos com resgate automático",
                "tipo_ofx": "SALDO_INVESTIMENTO",
                "fonte": "Bancos",
                "_saldo_ofx": float(saldo_invest),
                "_data_saldo_ofx": data_saldo,
                "_conta_ofx": f"Investimentos {conta}".strip(),
                "_banco_ofx": banco,
                "_tipo_conta": "Investimento",
            }]))

        if not quadros:
            st.error(
                "❌ A planilha de bancos precisa conter uma tabela de extrato "
                "ou uma linha de saldo de investimentos."
            )
            return pd.DataFrame()
        return pd.concat(quadros, ignore_index=True, sort=False)
    except Exception as e:
        logger.error("Erro ao processar planilha de bancos: %s", e)
        st.error(f"❌ Erro ao processar planilha de bancos: {str(e)}")
        return pd.DataFrame()


def processar_bancos(uploaded_file) -> pd.DataFrame:
    extensao = Path(uploaded_file.name).suffix.lower()
    if extensao == ".ofx":
        return processar_ofx(uploaded_file, "Bancos")
    return processar_excel_bancos(uploaded_file)


def processar_extrato_antecipacao(uploaded_file) -> pd.DataFrame:
    extensao = Path(uploaded_file.name).suffix.lower()
    if extensao == ".ofx":
        return processar_ofx(uploaded_file, "Antecipação de cartão")

    try:
        df = ler_tabela_flexivel(uploaded_file)
        df.columns = [normalizar_texto(col) for col in df.columns]
        df = df.loc[:, ~df.columns.duplicated()].copy()

        def localizar(*nomes):
            for nome in nomes:
                nome_normalizado = normalizar_texto(nome)
                if nome_normalizado in df.columns:
                    return nome_normalizado
            return None

        coluna_data = localizar(
            "recebimento", "pagamento", "vencimento", "data",
            "data da transação", "data transacao", "data de pagamento"
        )
        coluna_valor_final = localizar(
            "vl.final", "vl final", "valor final", "valor líquido",
            "valor liquido"
        )
        coluna_valor_bruto = localizar(
            "valor", "valor bruto", "valor da transação",
            "valor transacao", "total"
        )
        coluna_taxas = localizar("taxas", "taxa", "tarifas", "custo")
        coluna_descricao = localizar(
            "descrição", "descricao", "histórico", "historico", "lançamento",
            "lancamento", "evento", "titular"
        )
        coluna_titular = localizar("titular", "cliente")
        coluna_cpf = localizar("cpf", "cpf/cnpj")
        coluna_tipo = localizar("tipo")
        coluna_bandeira = localizar("bandeira")
        coluna_ultimos_digitos = localizar(
            "4 últimos digitos do cartão",
            "4 ultimos digitos do cartao",
            "últimos dígitos",
            "ultimos digitos",
        )
        coluna_autorizacao = localizar(
            "cód. autorização", "cod. autorizacao", "codigo autorizacao"
        )
        coluna_comprovante = localizar(
            "comp. venda", "comprovante venda", "comprovante"
        )
        coluna_status = localizar("status", "situação", "situacao")
        coluna_vencimento = localizar("vencimento")
        coluna_pagamento = localizar("pagamento")
        coluna_recebimento = localizar("recebimento")

        if not coluna_data or not (coluna_valor_final or coluna_valor_bruto):
            st.error(
                "❌ O extrato de antecipação em Excel precisa conter as colunas "
                "Recebimento/Pagamento e Valor ou Vl.Final."
            )
            return pd.DataFrame()

        df["data"] = pd.to_datetime(
            df[coluna_data],
            dayfirst=True,
            errors="coerce",
        )
        df["valor_bruto"] = (
            df[coluna_valor_bruto].apply(parse_valor_br)
            if coluna_valor_bruto else 0.0
        )
        df["taxas_antecipacao"] = (
            df[coluna_taxas].apply(parse_valor_br)
            if coluna_taxas else 0.0
        )
        df["valor"] = (
            df[coluna_valor_final].apply(parse_valor_br)
            if coluna_valor_final else
            df["valor_bruto"] - df["taxas_antecipacao"].fillna(0)
        )
        titular = (
            df[coluna_titular].fillna("").astype(str)
            if coluna_titular else pd.Series("", index=df.index)
        )
        tipo = (
            df[coluna_tipo].fillna("").astype(str)
            if coluna_tipo else pd.Series("", index=df.index)
        )
        bandeira = (
            df[coluna_bandeira].fillna("").astype(str)
            if coluna_bandeira else pd.Series("", index=df.index)
        )
        df["memo"] = (
            titular + " · " + tipo + " · " + bandeira
        ).str.strip(" ·")
        if coluna_descricao and coluna_descricao != coluna_titular:
            descricao = df[coluna_descricao].fillna("").astype(str)
            df.loc[df["memo"] == "", "memo"] = descricao
        df.loc[df["memo"] == "", "memo"] = "Antecipação de cartão"
        df["eh_antecipacao"] = df["memo"].map(normalizar_texto).str.contains(
            r"antecipa|recebiveis|recebíveis",
            regex=True,
        )
        df = df.dropna(subset=["data", "valor"])
        df = df[df["valor"] > 0]
        df["tipo_ofx"] = "ANTECIPACAO"
        df["fonte"] = "Antecipação de cartão"
        df["status_antecipacao"] = (
            df[coluna_status].fillna("").astype(str)
            if coluna_status else "Pago"
        )
        df["tipo_cartao"] = tipo
        df["titular_cartao"] = titular
        df["bandeira_cartao"] = bandeira
        df["cpf_titular"] = (
            df[coluna_cpf].fillna("").astype(str)
            if coluna_cpf else ""
        )
        df["ultimos_digitos_cartao"] = (
            df[coluna_ultimos_digitos].fillna("").astype(str)
            if coluna_ultimos_digitos else ""
        )
        df["codigo_autorizacao"] = (
            df[coluna_autorizacao].fillna("").astype(str)
            if coluna_autorizacao else ""
        )
        df["comprovante_venda"] = (
            df[coluna_comprovante].fillna("").astype(str)
            if coluna_comprovante else ""
        )
        df["vencimento_cartao"] = (
            df[coluna_vencimento].fillna("").astype(str)
            if coluna_vencimento else ""
        )
        df["pagamento_cartao"] = (
            df[coluna_pagamento].fillna("").astype(str)
            if coluna_pagamento else ""
        )
        df["recebimento_cartao"] = (
            df[coluna_recebimento].fillna("").astype(str)
            if coluna_recebimento else ""
        )
        return df[[
            "data", "valor", "valor_bruto", "taxas_antecipacao", "memo",
            "tipo_ofx", "fonte", "eh_antecipacao", "status_antecipacao",
            "tipo_cartao", "titular_cartao", "bandeira_cartao",
            "cpf_titular", "ultimos_digitos_cartao", "codigo_autorizacao",
            "comprovante_venda",
            "vencimento_cartao", "pagamento_cartao", "recebimento_cartao",
        ]]
    except Exception as e:
        logger.error("Erro ao processar extrato de antecipação: %s", e)
        st.error(f"❌ Erro ao processar antecipação: {str(e)}")
        return pd.DataFrame()


def extrair_texto_pdf(uploaded_file) -> str:
    uploaded_file.seek(0)
    pdf_bytes = uploaded_file.read()
    texto = ""

    try:
        from pypdf import PdfReader

        reader = PdfReader(io.BytesIO(pdf_bytes))
        texto = "\n".join(page.extract_text() or "" for page in reader.pages)
    except ImportError:
        try:
            from PyPDF2 import PdfReader

            reader = PdfReader(io.BytesIO(pdf_bytes))
            texto = "\n".join(page.extract_text() or "" for page in reader.pages)
        except ImportError:
            try:
                import pdfplumber

                with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
                    texto = "\n".join(
                        page.extract_text() or "" for page in pdf.pages
                    )
            except ImportError:
                texto = ""

    return texto


def extrair_texto_word(uploaded_file) -> str:
    extensao = Path(uploaded_file.name).suffix.lower()
    uploaded_file.seek(0)
    conteudo = uploaded_file.read()

    if extensao == ".docx":
        try:
            from docx import Document
        except ImportError:
            return ""

        documento = Document(io.BytesIO(conteudo))
        partes = [paragrafo.text for paragrafo in documento.paragraphs]
        for tabela in documento.tables:
            for linha in tabela.rows:
                partes.append(" ".join(celula.text for celula in linha.cells))
        return "\n".join(partes)

    if extensao != ".doc" or os.name != "nt":
        return ""

    texto_bruto = conteudo.decode("utf-8", errors="ignore")
    if "<html" in texto_bruto.lower():
        texto_html = re.sub(
            r"(?i)<br\s*/?>|</p>|</div>|</span>",
            "\n",
            texto_bruto,
        )
        texto_html = re.sub(r"<[^>]+>", " ", texto_html)
        linhas = [
            " ".join(html.unescape(linha).replace("\xa0", " ").split())
            for linha in texto_html.splitlines()
        ]
        return "\n".join(linha for linha in linhas if linha)

    with tempfile.TemporaryDirectory(prefix="dashboard_doc_") as pasta:
        origem = Path(pasta) / "origem.doc"
        destino = Path(pasta) / "extraido.txt"
        origem.write_bytes(conteudo)

        script = """
$ErrorActionPreference = 'Stop'
$word = $null
$document = $null
try {
    $word = New-Object -ComObject Word.Application
    $word.Visible = $false
    $document = $word.Documents.Open($env:DASHBOARD_DOC_ORIGEM)
    $document.SaveAs([ref]$env:DASHBOARD_DOC_DESTINO, [ref]2)
    $document.Close($false)
} finally {
    if ($document -ne $null) {
        [Runtime.InteropServices.Marshal]::ReleaseComObject($document) |
            Out-Null
    }
    if ($word -ne $null) {
        $word.Quit()
        [Runtime.InteropServices.Marshal]::ReleaseComObject($word) |
            Out-Null
    }
}
"""
        ambiente = os.environ.copy()
        ambiente["DASHBOARD_DOC_ORIGEM"] = str(origem)
        ambiente["DASHBOARD_DOC_DESTINO"] = str(destino)
        powershell = (
            Path(os.environ.get("ProgramFiles", r"C:\Program Files"))
            / "PowerShell" / "7" / "pwsh.exe"
        )
        executavel = str(powershell) if powershell.exists() else "powershell.exe"
        resultado = subprocess.run(
            [executavel, "-NoProfile", "-NonInteractive", "-Command", script],
            env=ambiente,
            capture_output=True,
            text=True,
            timeout=60,
            check=False,
        )
        if resultado.returncode != 0 or not destino.exists():
            return ""
        return destino.read_text(encoding="utf-8", errors="ignore")


def extrair_linhas_maquininha_de_texto(texto: str) -> pd.DataFrame:
    linhas = []
    padrao_data = re.compile(r"\b(\d{4}-\d{2}-\d{2}|\d{2}/\d{2}/\d{2,4})\b")
    padrao_valor = re.compile(
        r"(?<![\d/])(?:-?\s*R?\$?\s*)?\d{1,3}(?:\.\d{3})*,\d{2}(?![\d/])"
        r"|(?<![\d/])(?:-?\s*R?\$?\s*)?\d+\.\d{2}(?![\d/])",
        flags=re.IGNORECASE,
    )
    padrao_total_periodo = re.compile(
        r"total\s+do\s+periodo|total\s+do\s+período",
        flags=re.IGNORECASE,
    )
    linhas_texto = [" ".join(str(linha).split()) for linha in texto.splitlines()]
    totais_rede = []
    totais_rede_recebidos = []
    secao_rede = None
    data_secao_rede = ""
    for linha_limpa in linhas_texto:
        linha_norm = normalizar_texto(linha_limpa)
        if "lancamentos futuros" in linha_norm:
            secao_rede = None
            data_secao_rede = ""
            continue
        if "valores pagos" in linha_norm and "periodo" in linha_norm:
            secao_rede = "Recebidos crédito"
            datas_secao = list(padrao_data.finditer(linha_limpa))
            data_secao_rede = (
                datas_secao[0].group(1) if datas_secao else ""
            )
            continue
        if (
            "vendas com cartoes de credito" in linha_norm
            or "vendas com cartoes de debito" in linha_norm
        ) and "periodo" in linha_norm:
            secao_rede = (
                "Vendas crédito"
                if "credito" in linha_norm else "Vendas débito"
            )
            datas_secao = list(padrao_data.finditer(linha_limpa))
            data_secao_rede = (
                datas_secao[0].group(1) if datas_secao else ""
            )
            continue
        if "valores pagos" in linha_norm or "lancamentos futuros" in linha_norm:
            secao_rede = None
            data_secao_rede = ""
            continue
        if secao_rede and (
            "total do periodo" in linha_norm
            or "total de vendas no periodo" in linha_norm
        ):
            valores = list(padrao_valor.finditer(linha_limpa))
            if valores and data_secao_rede:
                valor = valores[-1].group(0)
                if re.match(r"^-\s+\d", valor.strip()):
                    valor = valor.strip()[1:].strip()
                totais_rede.append({
                    "Data": data_secao_rede,
                    "Hora": "",
                    "Tipo de transação": secao_rede,
                    "Nome": linha_limpa,
                    "Detalhe": "Resumo Rede",
                    "Valor": valor,
                })
                if secao_rede.startswith("Recebidos") or secao_rede == "Vendas débito":
                    totais_rede_recebidos.append(totais_rede[-1])
            secao_rede = None
            data_secao_rede = ""
            continue

    if totais_rede_recebidos:
        return pd.DataFrame(totais_rede_recebidos)

    if totais_rede:
        return pd.DataFrame(totais_rede)

    for linha_limpa in linhas_texto:
        if not linha_limpa:
            continue
        datas = list(padrao_data.finditer(linha_limpa))
        valores = list(padrao_valor.finditer(linha_limpa))
        if not valores:
            continue
        if datas:
            data = datas[0].group(1)
            texto_linha = linha_limpa
            valor = valores[-1].group(0)
            if re.match(r"^-\s+\d", valor.strip()):
                valor = valor.strip()[1:].strip()
        elif padrao_total_periodo.search(normalizar_texto(linha_limpa)):
            data = ""
            texto_linha = "Total do período"
            valor = valores[-1].group(0)
        else:
            continue
        linhas.append({
            "Data": data,
            "Hora": "",
            "Tipo de transação": texto_linha.split("  ")[0],
            "Nome": texto_linha,
            "Detalhe": "",
            "Valor": valor,
        })
    return pd.DataFrame(linhas)


def localizar_coluna_flexivel(colunas, aliases) -> Optional[str]:
    mapa = {normalizar_texto(coluna): coluna for coluna in colunas}
    aliases_norm = [normalizar_texto(alias) for alias in aliases]

    for alias in aliases_norm:
        coluna = mapa.get(alias)
        if coluna is not None:
            return coluna

    for alias in aliases_norm:
        for chave, coluna in mapa.items():
            if alias and alias in chave:
                return coluna
    return None


def ler_tabelas_maquininha(uploaded_file) -> List[pd.DataFrame]:
    extensao = Path(uploaded_file.name).suffix.lower()
    if extensao not in {".xlsx", ".xls", ".ods"}:
        return [ler_tabela_flexivel(uploaded_file)]

    uploaded_file.seek(0)
    try:
        planilhas = pd.read_excel(uploaded_file, sheet_name=None, header=None)
    except ImportError as e:
        if extensao == ".xls" and "xlrd" in str(e).lower():
            uploaded_file.seek(0)
            planilhas = converter_xls_com_excel(
                uploaded_file,
                header=None,
                sheet_name=None,
            )
            if not isinstance(planilhas, dict):
                planilhas = {"Planilha": planilhas}
        else:
            raise

    candidatas = []
    for nome_aba, bruto in planilhas.items():
        if bruto is None or bruto.empty:
            continue

        limite = min(len(bruto), 20)
        for indice_cabecalho in range(limite):
            cabecalho = bruto.iloc[indice_cabecalho].fillna("").astype(str)
            texto_cabecalho = " ".join(cabecalho.map(normalizar_texto))
            tem_data = "data" in texto_cabecalho
            tem_valor = (
                "valor" in texto_cabecalho
                or "depositado" in texto_cabecalho
                or "liquido" in texto_cabecalho
            )
            if not tem_data or not tem_valor:
                continue

            tabela = bruto.iloc[indice_cabecalho + 1:].copy()
            tabela.columns = [
                str(coluna).strip() if str(coluna).strip() else f"coluna_{idx}"
                for idx, coluna in enumerate(cabecalho, start=1)
            ]
            tabela = tabela.dropna(how="all")
            if not tabela.empty:
                tabela.attrs["aba_origem_maquininha"] = nome_aba
                candidatas.append(tabela)
            break

    if not candidatas:
        return [ler_tabela_flexivel(uploaded_file)]

    def prioridade(tabela):
        aba = normalizar_texto(tabela.attrs.get("aba_origem_maquininha", ""))
        colunas = " ".join(normalizar_texto(coluna) for coluna in tabela.columns)
        score = 0
        if "recebidos" in aba:
            score += 30
        if "pagamentos" in aba:
            score += 20
        if "valor depositado" in colunas:
            score += 10
        if "valor liquido" in colunas:
            score += 5
        return -score

    return sorted(candidatas, key=prioridade)


def normalizar_bruto_taxa_maquininha(df: Optional[pd.DataFrame]) -> pd.DataFrame:
    if df is None or df.empty:
        return pd.DataFrame()

    base = df.copy()
    valor_liquido = pd.to_numeric(
        base.get("valor", pd.Series(0, index=base.index)),
        errors="coerce",
    ).fillna(0).abs()
    taxa = pd.to_numeric(
        base.get("taxa_maquininha", pd.Series(0, index=base.index)),
        errors="coerce",
    ).fillna(0).abs()
    bruto_corrigido = (valor_liquido + taxa).where(valor_liquido > 0, 0)
    base["valor_bruto_maquininha"] = bruto_corrigido
    base["taxa_maquininha"] = taxa
    if "ponto_venda" not in base.columns:
        base["ponto_venda"] = ""
    base["ponto_venda"] = base["ponto_venda"].fillna("").astype(str).str.strip()
    if "maquininha_operadora" not in base.columns:
        origem = base.get(
            "_arquivo_origem",
            pd.Series("Maquininha", index=base.index),
        ).fillna("Maquininha").astype(str)
        base["maquininha_operadora"] = origem.map(
            lambda valor: detectar_operadora_maquininha(valor)
        )
    base["maquininha_operadora"] = (
        base["maquininha_operadora"].fillna("").astype(str).str.strip()
        .replace("", "Maquininha")
    )
    operadora_norm = base["maquininha_operadora"].map(normalizar_texto)
    ponto_equipamento = operadora_norm.isin(["pdv", "poy", "pos", "pinpad"])
    if ponto_equipamento.any():
        base.loc[ponto_equipamento & (base["ponto_venda"] == ""), "ponto_venda"] = (
            base.loc[ponto_equipamento, "maquininha_operadora"]
        )
        origem = base.get(
            "_arquivo_origem",
            pd.Series("Maquininha", index=base.index),
        ).fillna("Maquininha").astype(str)
        base.loc[ponto_equipamento, "maquininha_operadora"] = origem.loc[
            ponto_equipamento
        ].map(lambda valor: detectar_operadora_maquininha(valor))
    for coluna in ["bandeira", "parcelas", "status_maquininha"]:
        if coluna not in base.columns:
            base[coluna] = ""
        base[coluna] = base[coluna].fillna("").astype(str).str.strip()
    return base


def detectar_operadora_maquininha(nome_arquivo: str, df: Optional[pd.DataFrame] = None) -> str:
    texto = normalizar_texto(nome_arquivo)
    if df is not None and not df.empty:
        amostra = " ".join(normalizar_texto(coluna) for coluna in df.columns)
        valores = (
            df.head(5).fillna("").astype(str).agg(" ".join, axis=1).str.cat(sep=" ")
        )
        texto = f"{texto} {amostra} {normalizar_texto(valores)}"

    operadoras = [
        ("Rede Service", r"(^|[^a-z0-9])rede([^a-z0-9]|$)|redecard|rede\s*service"),
        ("Clinipay", r"clinipay|recebimentos\s+operacoes|recebimentos\s+operações"),
        ("Infinity Pay", r"infinity\s*pay|infinite\s*pay|infinitypay|infinitepay"),
        ("Stone", r"\bstone\b"),
        ("Saúde Service", r"saude\s*service|saúde\s*service"),
        ("Sipag", r"\bsipag\b"),
        ("Getnet", r"\bgetnet\b"),
        ("Sicredi", r"\bsicredi\b"),
        ("Cielo", r"\bcielo\b"),
        ("PagSeguro", r"pagseguro|pag\s*seguro"),
        ("Mercado Pago", r"mercado\s*pago"),
        ("Adyen", r"\baden\b|\badyen\b"),
    ]
    for nome, padrao in operadoras:
        if re.search(padrao, texto):
            return nome
    return "Maquininha"


def processar_infinity_pay(uploaded_file) -> pd.DataFrame:
    try:
        extensao = Path(uploaded_file.name).suffix.lower()
        if extensao == ".pdf":
            texto = extrair_texto_pdf(uploaded_file)

            if not texto.strip():
                st.error(
                    "❌ Para ler PDF de maquininha, instale a dependência "
                    "pypdf, PyPDF2 ou pdfplumber, ou envie o extrato em Excel/OFX."
                )
                return pd.DataFrame()

            df = extrair_linhas_maquininha_de_texto(texto)
        elif extensao in {".doc", ".docx"}:
            texto = extrair_texto_word(uploaded_file)
            if not texto.strip():
                st.error(
                    "❌ Para ler Word de maquininha, instale python-docx "
                    "ou envie o extrato em PDF, Excel ou OFX."
                )
                return pd.DataFrame()
            df = extrair_linhas_maquininha_de_texto(texto)
        else:
            tabelas = ler_tabelas_maquininha(uploaded_file)
            df = next((tabela for tabela in tabelas if tabela is not None), None)

        if df is None or df.empty:
            st.warning(f"⚠️ {uploaded_file.name}: nenhum registro encontrado.")
            return pd.DataFrame()

        colunas_data = [
            "data do depósito", "data do deposito",
            "data do recebimento", "data de recebimento", "data recebimento",
            "data recebida", "data prevista do recebimento",
            "data do pagamento", "data de pagamento", "data pagamento",
            "data prevista de pagamento", "data do crédito", "data credito",
            "data do credito", "data da venda", "data original da venda",
            "data venda", "data da transação", "data da transacao",
            "data", "date",
        ]
        colunas_valor = [
            "recebido (r$)", "recebido",
            "líquido (r$)", "liquido (r$)", "líquido", "liquido",
            "valor depositado", "valor líquido da parcela",
            "valor liquido da parcela", "valor líquido", "valor liquido",
            "valor recebido", "valor do recebimento", "valor creditado",
            "valor a receber", "valor da venda", "valor bruto da parcela atualizada",
            "valor bruto da parcela original", "valor bruto", "amount", "value",
            "valor",
        ]
        colunas_valor_bruto = [
            "valor da parcela (r$)", "valor da parcela",
            "valor bruto de recebimento",
            "valor a receber por beneficiario",
            "valor a receber por beneficiário",
            "valor da venda atualizado",
            "valor da venda original",
            "valor bruto da parcela atualizada",
            "valor bruto da parcela original",
            "valor bruto",
            "valor da venda",
            "valor original da venda",
            "valor original",
            "valor transação",
            "valor transacao",
        ]
        colunas_taxa = [
            "valor total das taxas descontadas",
            "valor total das taxas descontadas mdr recebimento automatico",
            "valor mdr",
            "valor taxa de recebimento automatico",
            "valor da taxa",
            "valor taxa",
            "valor tarifa",
            "valor tarifas",
            "taxas descontadas",
            "tarifa", "tarifas", "custo", "custos",
            "desconto",
            "taxa", "taxas", "mdr",
        ]
        colunas_hora = ["hora", "time"]
        colunas_tipo = [
            "tipo de transação", "tipo de transacao",
            "categoria da transação", "categoria da transacao",
            "modalidade", "produto", "forma de pagamento", "meio de pagamento",
            "tecnologia", "tipo",
            "tipo do plano",
        ]
        colunas_bandeira = [
            "bandeira", "cartão", "cartao", "brand", "arranjo", "produto bandeira",
        ]
        colunas_parcelas = [
            "número de parcelas", "numero de parcelas", "parcelas",
            "parcela", "qtd parcelas", "quantidade de parcelas",
            "plano", "installments",
        ]
        colunas_status = [
            "status da venda", "status", "situação", "situacao",
            "estado", "resultado",
        ]
        colunas_antecipada = [
            "antecipada", "antecipado", "antecipacao", "antecipação",
        ]
        colunas_operadora = [
            "operadora", "adquirente", "credenciadora",
            "instituição", "instituicao", "canal",
        ]
        colunas_ponto_venda = [
            "tipo de maquininha", "maquininha", "ponto de venda",
            "pdv", "poy", "terminal", "equipamento", "modelo",
        ]
        colunas_nome = [
            "nome", "name", "nome do estabelecimento", "cliente",
            "estabelecimento", "beneficiário", "beneficiario",
            "nome do paciente", "paciente", "descrição", "descricao",
        ]
        colunas_detalhe = [
            "detalhe", "detail", "nsu/cv", "nsu", "tid",
            "número da autorização", "numero da autorizacao", "parcela",
            "número de parcelas", "numero de parcelas",
            "número da operação", "numero da operacao",
            "código de autorização", "codigo de autorizacao",
        ]

        tabela_valida = None
        for candidata in ([df] if extensao in {".pdf", ".doc", ".docx"} else tabelas):
            if candidata is None or candidata.empty:
                continue
            candidata = candidata.loc[:, ~candidata.columns.duplicated()].copy()
            col_data = localizar_coluna_flexivel(candidata.columns, colunas_data)
            col_valor = localizar_coluna_flexivel(candidata.columns, colunas_valor)
            if col_data and col_valor:
                tabela_valida = (candidata, col_data, col_valor)
                break

        if tabela_valida is None:
            st.error(
                "❌ Extrato de maquininha precisa conter as colunas Data e Valor."
            )
            return pd.DataFrame()

        df, col_data, col_valor = tabela_valida
        col_hora = localizar_coluna_flexivel(df.columns, colunas_hora)
        col_tipo = localizar_coluna_flexivel(df.columns, colunas_tipo)
        col_bandeira = localizar_coluna_flexivel(df.columns, colunas_bandeira)
        col_parcelas = localizar_coluna_flexivel(df.columns, colunas_parcelas)
        col_status = localizar_coluna_flexivel(df.columns, colunas_status)
        col_antecipada = localizar_coluna_flexivel(df.columns, colunas_antecipada)
        col_operadora = localizar_coluna_flexivel(df.columns, colunas_operadora)
        col_ponto_venda = localizar_coluna_flexivel(df.columns, colunas_ponto_venda)
        col_nome = localizar_coluna_flexivel(df.columns, colunas_nome)
        col_detalhe = localizar_coluna_flexivel(df.columns, colunas_detalhe)
        col_valor_bruto = localizar_coluna_flexivel(df.columns, colunas_valor_bruto)
        col_taxa = localizar_coluna_flexivel(df.columns, colunas_taxa)
        if not col_data or not col_valor:
            st.error(
                "❌ Extrato de maquininha precisa conter as colunas Data e Valor."
            )
            return pd.DataFrame()

        df["data"] = parse_data_flexivel(df[col_data])
        df["valor"] = df[col_valor].apply(parse_valor_br)
        df["valor_bruto_maquininha"] = (
            df[col_valor_bruto].apply(parse_valor_br)
            if col_valor_bruto else pd.NA
        )
        df["taxa_maquininha"] = (
            df[col_taxa].apply(parse_valor_br).abs()
            if col_taxa else pd.NA
        )
        df = df.dropna(subset=["data", "valor"])
        df = df[df["valor"] != 0].copy()
        taxa_calculada = (
            pd.to_numeric(df["valor_bruto_maquininha"], errors="coerce").abs()
            - pd.to_numeric(df["valor"], errors="coerce").abs()
        )
        df["taxa_maquininha"] = (
            pd.to_numeric(df["taxa_maquininha"], errors="coerce")
            .fillna(taxa_calculada.where(taxa_calculada > 0))
            .fillna(0)
        )
        bruto_num = pd.to_numeric(
            df["valor_bruto_maquininha"],
            errors="coerce",
        ).abs()
        liquido_num = pd.to_numeric(df["valor"], errors="coerce").abs()
        taxa_num = pd.to_numeric(df["taxa_maquininha"], errors="coerce").fillna(0).abs()
        if col_taxa and "valor" not in normalizar_texto(col_taxa):
            taxa_num = taxa_num.mask(
                (taxa_num > 0) & (taxa_num <= 1) & (bruto_num > 0),
                taxa_num * bruto_num,
            )
        bruto_num = bruto_num.fillna(0)
        bruto_num = bruto_num.mask(
            (bruto_num <= 0) & (liquido_num > 0),
            liquido_num + taxa_num,
        )
        df["valor_bruto_maquininha"] = bruto_num
        df["taxa_maquininha"] = taxa_num
        tipo = (
            df[col_tipo].fillna("").astype(str).str.strip()
            if col_tipo else pd.Series("", index=df.index)
        )
        bandeira = (
            df[col_bandeira].fillna("").astype(str).str.strip()
            if col_bandeira else pd.Series("", index=df.index)
        )
        parcelas = (
            df[col_parcelas].fillna("").astype(str).str.strip()
            if col_parcelas else pd.Series("", index=df.index)
        )
        status_maquininha = (
            df[col_status].fillna("").astype(str).str.strip()
            if col_status else pd.Series("", index=df.index)
        )
        antecipada = (
            df[col_antecipada].fillna("").astype(str).str.strip()
            if col_antecipada else pd.Series("", index=df.index)
        )
        operadora_coluna = (
            df[col_operadora].fillna("").astype(str).str.strip()
            if col_operadora else pd.Series("", index=df.index)
        )
        ponto_venda = (
            df[col_ponto_venda].fillna("").astype(str).str.strip()
            if col_ponto_venda else pd.Series("", index=df.index)
        )
        operadora_detectada = detectar_operadora_maquininha(
            uploaded_file.name,
            df,
        )
        nome = (
            df[col_nome].fillna("").astype(str).str.strip()
            if col_nome else pd.Series("", index=df.index)
        )
        detalhe = (
            df[col_detalhe].fillna("").astype(str).str.strip()
            if col_detalhe else pd.Series("", index=df.index)
        )
        hora = (
            df[col_hora].fillna("").astype(str).str.strip()
            if col_hora else pd.Series("", index=df.index)
        )
        df["hora"] = hora
        df["tipo_transacao"] = tipo
        df["bandeira"] = bandeira
        df["parcelas"] = parcelas
        df["status_maquininha"] = status_maquininha
        df["antecipada_maquininha"] = antecipada
        df["_relatorio_maquininha_detalhado"] = bool(col_antecipada)
        df["ponto_venda"] = ponto_venda
        df["maquininha_operadora"] = operadora_detectada
        operadora_norm = operadora_coluna.map(normalizar_texto)
        operadora_valida = (
            (operadora_coluna != "")
            & ~operadora_norm.isin(["pdv", "poy", "pos", "pinpad"])
        )
        df.loc[operadora_valida, "maquininha_operadora"] = operadora_coluna[
            operadora_valida
        ]
        df["nome"] = nome
        df["detalhe"] = detalhe
        df["memo"] = (
            tipo + " - " + nome + " - " + detalhe
        ).str.strip(" -")
        df["fonte"] = "Maquininha"
        df["tipo_ofx"] = "MAQUININHA"
        df["conta_destino"] = "Maquininha"
        df["fluxo_infinity"] = df["valor"].apply(
            lambda valor: "Entrada" if valor > 0 else "Saída"
        )
        df["_linha_origem_infinity"] = df.index
        df = normalizar_bruto_taxa_maquininha(df)
        return df[[
            "data", "hora", "valor", "memo", "tipo_transacao", "nome",
            "detalhe", "fonte", "tipo_ofx", "conta_destino",
            "fluxo_infinity", "valor_bruto_maquininha", "taxa_maquininha",
            "maquininha_operadora", "bandeira", "parcelas", "status_maquininha",
            "antecipada_maquininha", "_relatorio_maquininha_detalhado",
            "ponto_venda",
            "_linha_origem_infinity",
        ]]
    except Exception as e:
        logger.error("Erro ao processar extrato de maquininha: %s", e)
        st.error(f"❌ Erro ao processar extrato de maquininha: {str(e)}")
        return pd.DataFrame()


def processar_multiplos_arquivos(
    arquivos,
    processador,
    *args,
) -> pd.DataFrame:
    quadros = []
    for arquivo in arquivos or []:
        if arquivo.size > 10 * 1024 * 1024:
            st.error(f"❌ {arquivo.name}: máximo de 10 MB por arquivo.")
            continue
        quadro = processador(arquivo, *args)
        if quadro is not None and not quadro.empty:
            quadro = quadro.copy()
            quadro["_arquivo_origem"] = arquivo.name
            quadros.append(quadro)

    if not quadros:
        return pd.DataFrame()

    combinado = pd.concat(quadros, ignore_index=True, sort=False)
    nome_processador = getattr(processador, "__name__", "")
    if nome_processador in {
        "processar_bancos",
        "processar_excel_bancos",
    }:
        return combinado.reset_index(drop=True)
    if nome_processador == "processar_infinity_pay":
        relatorio_detalhado = (
            combinado.get(
                "_relatorio_maquininha_detalhado",
                pd.Series(False, index=combinado.index),
            )
            .fillna(False)
            .astype(bool)
        )
        if relatorio_detalhado.any():
            texto_maquininha = (
                combinado.get("memo", pd.Series("", index=combinado.index))
                .fillna("").map(normalizar_texto)
                + " "
                + combinado.get("tipo_transacao", pd.Series("", index=combinado.index))
                .fillna("").map(normalizar_texto)
                + " "
                + combinado.get("detalhe", pd.Series("", index=combinado.index))
                .fillna("").map(normalizar_texto)
                + " "
                + combinado.get("_arquivo_origem", pd.Series("", index=combinado.index))
                .fillna("").map(normalizar_texto)
            )
            valores_maquininha = pd.to_numeric(
                combinado.get("valor", 0),
                errors="coerce",
            ).fillna(0)
            deposito_statement = (
                ~relatorio_detalhado
                & (valores_maquininha > 0)
                & texto_maquininha.str.contains(
                    r"deposito\s+de\s+vendas|dep[oó]sito\s+de\s+vendas|"
                    r"deposito\s+infinitepay|dep[oó]sito\s+infinitepay|"
                    r"statements?",
                    regex=True,
                )
            )
            combinado = combinado[~deposito_statement].copy()
        return combinado.reset_index(drop=True)

    colunas_deduplicacao = [
        coluna
        for coluna in combinado.columns
        if coluna != "_arquivo_origem"
    ]
    if colunas_deduplicacao:
        combinado = combinado.drop_duplicates(
            subset=colunas_deduplicacao,
            keep="last",
        )
    return combinado.reset_index(drop=True)


def dataframe_para_json(df: Optional[pd.DataFrame]) -> Optional[str]:
    if df is None:
        return None
    return df.to_json(orient="table", date_format="iso", force_ascii=False)


def dataframe_do_json(conteudo: Optional[str]) -> Optional[pd.DataFrame]:
    if conteudo is None:
        return None
    df = pd.read_json(io.StringIO(conteudo), orient="table")
    if "data" in df.columns:
        df["data"] = pd.to_datetime(df["data"], errors="coerce")
    return df


def criar_payload_compartilhado(user_id: int) -> str:
    try:
        contas = get_contas(user_id)
    except Exception:
        contas = []

    payload = {
        "versao": 1,
        "df_excel": dataframe_para_json(st.session_state.df_excel),
        "df_ofx": dataframe_para_json(st.session_state.df_ofx),
        "df_antecipacao": dataframe_para_json(
            st.session_state.df_antecipacao
        ),
        "df_orcamentos": dataframe_para_json(st.session_state.df_orcamentos),
        "df_clinipay": dataframe_para_json(st.session_state.df_clinipay),
        "df_fluxo_caixa": dataframe_para_json(st.session_state.df_fluxo_caixa),
        "df_infinity_pay": dataframe_para_json(
            st.session_state.df_infinity_pay
        ),
        "df_belle_receber": dataframe_para_json(
            st.session_state.df_belle_receber
        ),
        "df_belle_pagar": dataframe_para_json(
            st.session_state.df_belle_pagar
        ),
        "df_belle_gerencial": dataframe_para_json(
            st.session_state.df_belle_gerencial
        ),
        "df_amigotech_receber": dataframe_para_json(
            st.session_state.df_amigotech_receber
        ),
        "df_amigotech_pagar": dataframe_para_json(
            st.session_state.df_amigotech_pagar
        ),
        "df_conta_azul_receber": dataframe_para_json(
            st.session_state.df_conta_azul_receber
        ),
        "df_conta_azul_vendas": dataframe_para_json(
            st.session_state.df_conta_azul_vendas
        ),
        "df_conta_azul_pagar": dataframe_para_json(
            st.session_state.df_conta_azul_pagar
        ),
        "contas": [list(conta) for conta in contas],
        "periodo_inicio": str(st.session_state.periodo_inicio_global),
        "periodo_fim": str(st.session_state.periodo_fim_global),
    }
    return json.dumps(payload, ensure_ascii=False)


def carregar_payload_compartilhado(payload_json: str) -> None:
    payload = json.loads(payload_json)
    st.session_state.df_excel = dataframe_do_json(payload.get("df_excel"))
    st.session_state.df_ofx = dataframe_do_json(payload.get("df_ofx"))
    st.session_state.df_antecipacao = dataframe_do_json(
        payload.get("df_antecipacao")
    )
    st.session_state.df_orcamentos = dataframe_do_json(
        payload.get("df_orcamentos")
    )
    st.session_state.df_clinipay = dataframe_do_json(payload.get("df_clinipay"))
    st.session_state.df_fluxo_caixa = dataframe_do_json(
        payload.get("df_fluxo_caixa")
    )
    st.session_state.df_infinity_pay = dataframe_do_json(
        payload.get("df_infinity_pay")
    )
    st.session_state.df_belle_receber = dataframe_do_json(
        payload.get("df_belle_receber")
    )
    st.session_state.df_belle_pagar = dataframe_do_json(
        payload.get("df_belle_pagar")
    )
    st.session_state.df_belle_gerencial = dataframe_do_json(
        payload.get("df_belle_gerencial")
    )
    st.session_state.df_amigotech_receber = dataframe_do_json(
        payload.get("df_amigotech_receber")
    )
    st.session_state.df_amigotech_pagar = dataframe_do_json(
        payload.get("df_amigotech_pagar")
    )
    st.session_state.df_conta_azul_receber = dataframe_do_json(
        payload.get("df_conta_azul_receber")
    )
    st.session_state.df_conta_azul_vendas = dataframe_do_json(
        payload.get("df_conta_azul_vendas")
    )
    st.session_state.df_conta_azul_pagar = dataframe_do_json(
        payload.get("df_conta_azul_pagar")
    )
    st.session_state.share_accounts = payload.get("contas", [])
    if payload.get("periodo_inicio"):
        st.session_state.periodo_inicio_global = date.fromisoformat(
            payload["periodo_inicio"]
        )
    if payload.get("periodo_fim"):
        st.session_state.periodo_fim_global = date.fromisoformat(
            payload["periodo_fim"]
        )


def resumo_do_json(conteudo: Optional[str]) -> Dict[str, float]:
    if not conteudo:
        return {}
    try:
        return json.loads(conteudo)
    except json.JSONDecodeError:
        return {}


def fmt_periodo_salvo(inicio, fim) -> str:
    try:
        data_inicio = pd.to_datetime(inicio).strftime("%d/%m/%Y")
        data_fim = pd.to_datetime(fim).strftime("%d/%m/%Y")
        return f"{data_inicio} a {data_fim}"
    except Exception:
        return "Período não informado"


def rotulo_mes_referencia(mes_numero: str, ano: int) -> str:
    nome_mes = dict(MESES_RELATORIO).get(str(mes_numero).zfill(2), "")
    return f"{nome_mes} {ano}".strip()


def tipo_conta_registro(conta) -> str:
    if len(conta) >= 6 and conta[5]:
        return str(conta[5])
    return "Conta bancária"


def url_base_atual() -> str:
    configurada = os.getenv("APP_PUBLIC_URL", "").strip()
    if configurada:
        return configurada.rstrip("/")

    try:
        atual = str(st.context.url)
    except Exception:
        atual = ""

    if not atual:
        return "http://localhost:8501"

    partes = urlsplit(atual)
    return urlunsplit((partes.scheme, partes.netloc, partes.path, "", "")).rstrip("/")


# =========================
# SESSION STATE
# =========================
if "user" not in st.session_state:
    st.session_state.user = None
if "pagina" not in st.session_state:
    st.session_state.pagina = "visao"
if "df_excel" not in st.session_state:
    st.session_state.df_excel = None
if "df_ofx" not in st.session_state:
    st.session_state.df_ofx = None
if "df_antecipacao" not in st.session_state:
    st.session_state.df_antecipacao = None
if "df_orcamentos" not in st.session_state:
    st.session_state.df_orcamentos = None
if "df_clinipay" not in st.session_state:
    st.session_state.df_clinipay = None
if "df_fluxo_caixa" not in st.session_state:
    st.session_state.df_fluxo_caixa = None
if "df_infinity_pay" not in st.session_state:
    st.session_state.df_infinity_pay = None
if "df_belle_receber" not in st.session_state:
    st.session_state.df_belle_receber = None
if "df_belle_pagar" not in st.session_state:
    st.session_state.df_belle_pagar = None
if "df_belle_gerencial" not in st.session_state:
    st.session_state.df_belle_gerencial = None
if "df_amigotech_receber" not in st.session_state:
    st.session_state.df_amigotech_receber = None
if "df_amigotech_pagar" not in st.session_state:
    st.session_state.df_amigotech_pagar = None
if "df_conta_azul_receber" not in st.session_state:
    st.session_state.df_conta_azul_receber = None
if "df_conta_azul_vendas" not in st.session_state:
    st.session_state.df_conta_azul_vendas = None
if "df_conta_azul_pagar" not in st.session_state:
    st.session_state.df_conta_azul_pagar = None
if "mostrar_nova_conta" not in st.session_state:
    st.session_state.mostrar_nova_conta = False
if "mostrar_modal_upload" not in st.session_state:
    st.session_state.mostrar_modal_upload = False
if "upload_version" not in st.session_state:
    st.session_state.upload_version = 0
if "historico_chat" not in st.session_state:
    st.session_state.historico_chat = []
if "share_mode" not in st.session_state:
    st.session_state.share_mode = False
if "share_token" not in st.session_state:
    st.session_state.share_token = None
if "share_title" not in st.session_state:
    st.session_state.share_title = ""
if "share_accounts" not in st.session_state:
    st.session_state.share_accounts = []
if "mostrar_compartilhamento" not in st.session_state:
    st.session_state.mostrar_compartilhamento = False
if "mostrar_salvar_relatorio" not in st.session_state:
    st.session_state.mostrar_salvar_relatorio = False
if "relatorio_salvo_para_abrir" not in st.session_state:
    st.session_state.relatorio_salvo_para_abrir = None
if "usuario_exclusao_pendente" not in st.session_state:
    st.session_state.usuario_exclusao_pendente = None
if "periodo_inicio_global" not in st.session_state:
    st.session_state.periodo_inicio_global = date.today().replace(day=1)
if "periodo_fim_global" not in st.session_state:
    st.session_state.periodo_fim_global = date.today()

# Um link com ?share=TOKEN abre um snapshot somente leitura, sem exigir login.
token_compartilhado = st.query_params.get("share", "")
if not token_compartilhado and st.session_state.share_mode:
    st.session_state.share_mode = False
    st.session_state.share_token = None
    st.session_state.share_title = ""
    st.session_state.share_accounts = []
    st.session_state.user = None

if token_compartilhado:
    relatorio = get_shared_report(token_compartilhado)
    if not relatorio:
        st.error("Este link de relatório é inválido, expirou ou foi revogado.")
        st.stop()

    if token_compartilhado != st.session_state.share_token:
        _, owner_id, titulo, payload_json, _, _ = relatorio
        try:
            carregar_payload_compartilhado(payload_json)
        except Exception as e:
            logger.exception("Falha ao abrir relatório compartilhado: %s", e)
            st.error("Não foi possível carregar este relatório.")
            st.stop()

        st.session_state.share_mode = True
        st.session_state.share_token = token_compartilhado
        st.session_state.share_title = titulo
        st.session_state.user = User(
            id=owner_id,
            username=titulo,
            is_admin=False,
        )
        st.session_state.pagina = "visao"
        st.session_state.mostrar_modal_upload = False

# =========================
# LOGIN
# =========================
if not st.session_state.user:
    from login_ui import render_login

    render_login(User, login_user, reset_password_with_code, logger)

# =========================
# SIDEBAR
# =========================
user: User = st.session_state.user
if st.session_state.relatorio_salvo_para_abrir and not st.session_state.share_mode:
    relatorio = get_saved_report(
        st.session_state.relatorio_salvo_para_abrir,
        user.id,
    )
    st.session_state.relatorio_salvo_para_abrir = None
    if relatorio:
        carregar_payload_compartilhado(relatorio[3])
        st.session_state.pagina = "visao"
        st.session_state.mostrar_modal_upload = False
        st.rerun()

current_month = (
    datetime.now(ZoneInfo("America/Sao_Paulo")).strftime("%B %Y")
    .replace("January","Janeiro").replace("February","Fevereiro")
    .replace("March","Março").replace("April","Abril")
    .replace("May","Maio").replace("June","Junho")
    .replace("July","Julho").replace("August","Agosto")
    .replace("September","Setembro").replace("October","Outubro")
    .replace("November","Novembro").replace("December","Dezembro")
)


@st.dialog("🔗 Compartilhar relatório", width="large")
def abrir_compartilhamento():
    st.caption(
        "Crie uma cópia somente leitura dos dados atuais. O cliente não poderá "
        "importar arquivos, editar contas nem acessar a administração."
    )
    titulo = st.text_input(
        "Nome do relatório",
        value=f"Relatório financeiro — {current_month}",
    )
    validade = st.selectbox(
        "Validade do link",
        ["7 dias", "30 dias", "90 dias", "Sem expiração"],
        index=1,
    )
    base_url = st.text_input(
        "Endereço público do app",
        value=url_base_atual(),
        help=(
            "Para funcionar fora do seu computador, este endereço precisa ser "
            "o domínio onde o Streamlit está publicado."
        ),
    )

    host = urlsplit(base_url).hostname or ""
    if host in {"localhost", "127.0.0.1", "::1"}:
        st.warning(
            "Este endereço é local. O link só abrirá neste computador até o app "
            "ser publicado em um servidor."
        )

    if st.button(
        "Criar link compartilhável",
        type="primary",
        use_container_width=True,
    ):
        if not titulo.strip() or not base_url.strip():
            st.error("Informe o nome do relatório e o endereço público do app.")
        else:
            dias = {
                "7 dias": 7,
                "30 dias": 30,
                "90 dias": 90,
                "Sem expiração": None,
            }[validade]
            expira_em = (
                (datetime.utcnow() + timedelta(days=dias)).strftime(
                    "%Y-%m-%d %H:%M:%S"
                )
                if dias else None
            )
            token = create_shared_report(
                user.id,
                titulo.strip(),
                criar_payload_compartilhado(user.id),
                expira_em,
            )
            link = f"{base_url.rstrip('/')}?share={token}"
            st.session_state.ultimo_link_compartilhado = link

    link_criado = st.session_state.get("ultimo_link_compartilhado")
    if link_criado:
        st.success("Link criado. Envie este endereço ao cliente:")
        st.code(link_criado, language=None)
        st.link_button(
            "Abrir link em uma nova aba",
            link_criado,
            use_container_width=True,
        )

    relatorios = list_shared_reports(user.id)
    ativos = [r for r in relatorios if r[4] is None]
    if ativos:
        st.markdown("#### Links criados")
        for report_id, nome, criado_em, expira_em, _ in ativos[:10]:
            c1, c2 = st.columns([4, 1])
            with c1:
                validade_texto = (
                    f"expira em {expira_em}" if expira_em else "sem expiração"
                )
                st.caption(f"{nome} · criado em {criado_em} · {validade_texto}")
            with c2:
                if st.button("Revogar", key=f"revoke_share_{report_id}"):
                    revoke_shared_report(report_id, user.id)
                    st.rerun()


def abrir_importacao():
    st.session_state.mostrar_modal_upload = True


def limpar_dados_importados():
    for chave in [
        "df_excel",
        "df_ofx",
        "df_antecipacao",
        "df_orcamentos",
        "df_clinipay",
        "df_fluxo_caixa",
        "df_infinity_pay",
        "df_belle_receber",
        "df_belle_pagar",
        "df_belle_gerencial",
        "df_amigotech_receber",
        "df_amigotech_pagar",
        "df_conta_azul_receber",
        "df_conta_azul_vendas",
        "df_conta_azul_pagar",
    ]:
        st.session_state[chave] = None
    st.session_state.upload_version += 1
    st.session_state.historico_chat = []


def existem_dados_importados() -> bool:
    return any(
        st.session_state.get(chave) is not None
        for chave in [
            "df_excel",
            "df_ofx",
            "df_antecipacao",
            "df_orcamentos",
            "df_clinipay",
            "df_fluxo_caixa",
            "df_infinity_pay",
            "df_belle_receber",
            "df_belle_pagar",
            "df_belle_gerencial",
            "df_amigotech_receber",
            "df_amigotech_pagar",
            "df_conta_azul_receber",
            "df_conta_azul_vendas",
            "df_conta_azul_pagar",
        ]
    )


def solicitar_compartilhamento():
    st.session_state.mostrar_compartilhamento = True


def solicitar_salvar_relatorio():
    st.session_state.mostrar_salvar_relatorio = True


def navegar_para(pagina):
    st.session_state.pagina = pagina
    st.session_state.mostrar_modal_upload = False


def sair_do_app():
    logger.info("User logged out: %s", st.session_state.user.username)
    st.session_state.user = None
    st.session_state.pagina = "visao"
    st.session_state.mostrar_modal_upload = False


with st.sidebar:
    nome_exibicao = html.escape(user.username)
    iniciais_exibicao = html.escape(user.initials())
    badge_admin = (
        "🔒 SOMENTE LEITURA"
        if st.session_state.share_mode
        else "👑 ADMIN" if user.is_admin else ""
    )
    st.markdown(f"""
    <div style="display:flex;align-items:center;gap:10px;margin-bottom:1.5rem;
                padding-bottom:1rem;border-bottom:1px solid #1E1E3A">
        <div style="width:34px;height:34px;border-radius:50%;
                    background:linear-gradient(135deg,#534AB7,#7F77DD);
                    color:#fff;display:flex;align-items:center;justify-content:center;
                    font-size:0.8rem;font-weight:600">{iniciais_exibicao}</div>
        <div>
            <div style="font-size:0.85rem;font-weight:600;color:#E2E8F0">
                {nome_exibicao} {badge_admin}</div>
            <div style="font-size:0.7rem;color:#4A4A7A">{current_month}</div>
        </div>
    </div>
    """, unsafe_allow_html=True)

    if not st.session_state.share_mode:
        st.button(
            "📂  Importar Arquivos",
            use_container_width=True,
            key="btn_upload",
            on_click=abrir_importacao,
        )
        st.button(
            "🔗  Compartilhar Relatório",
            use_container_width=True,
            key="btn_share",
            on_click=solicitar_compartilhamento,
        )
        st.button(
            "💾  Salvar Relatório",
            use_container_width=True,
            key="btn_save_report",
            on_click=solicitar_salvar_relatorio,
        )
    else:
        st.info("Visualização compartilhada. Os dados não podem ser alterados.")

    st.markdown("---")
    with st.expander("📅 Período global", expanded=False):
        st.date_input(
            "De",
            key="periodo_inicio_global",
            format="DD/MM/YYYY",
        )
        st.date_input(
            "Até",
            key="periodo_fim_global",
            format="DD/MM/YYYY",
        )
        if (
            st.session_state.periodo_inicio_global
            > st.session_state.periodo_fim_global
        ):
            st.error("A data inicial deve ser anterior à data final.")
        else:
            st.caption(
                f"{st.session_state.periodo_inicio_global.strftime('%d/%m/%Y')}"
                " até "
                f"{st.session_state.periodo_fim_global.strftime('%d/%m/%Y')}"
            )

    paginas = [
        ("visao",     "📊", "Visão Financeira"),
        ("fechamento","📋", "Fechamento"),
        ("fornecedores","🧾", "Fornecedores"),
        ("comparativo","📈", "Comparativo"),
        ("detalhes",  "🔍", "Detalhes"),
        ("saldo",     "🏦", "Saldo Bancário"),
    ]
    if not st.session_state.share_mode:
        paginas.insert(3, ("relatorios", "💾", "Relatórios"))
    if user.is_admin and not st.session_state.share_mode:
        paginas.append(("admin", "⚙️", "Painel Admin"))

    for key, icone, label in paginas:
        st.button(
            f"{icone}  {label}",
            key=f"nav_{key}",
            use_container_width=True,
            on_click=navegar_para,
            args=(key,),
        )

    st.markdown("---")
    if not st.session_state.share_mode:
        st.button("🚪 Sair", use_container_width=True, on_click=sair_do_app)

if st.session_state.mostrar_compartilhamento:
    st.session_state.mostrar_compartilhamento = False
    abrir_compartilhamento()

# =========================
# MODAL DE UPLOAD
# =========================
if st.session_state.mostrar_modal_upload:
    st.markdown("""
    <style>
    .upload-title-card {
        background: #111326;
        border: 1px solid #2A2A4A;
        border-radius: 14px;
        padding: 24px 28px;
        margin-bottom: 24px;
        box-shadow: 0 18px 45px rgba(0,0,0,.28);
    }
    </style>
    """, unsafe_allow_html=True)

    with st.container():
        st.markdown("""
        <div class="upload-title-card">
            <div style="font-size:1.1rem;font-weight:700;color:#E2E8F0;margin-bottom:4px">
                📂 Importar Arquivos
            </div>
            <div style="font-size:0.78rem;color:#4A4A7A;margin-bottom:0">
                Faça upload dos documentos financeiros para análise no painel
            </div>
        </div>
        """, unsafe_allow_html=True)

        formatos_planilha = ["xlsx", "xls", "csv", "ods"]
        upload_version = st.session_state.upload_version

        col_limpar, col_status = st.columns([1, 3])
        with col_limpar:
            if st.button(
                "🧹 Limpar dados",
                type="secondary",
                use_container_width=True,
                key=f"limpar_uploads_{upload_version}",
            ):
                limpar_dados_importados()
                st.success("Dados da cliente atual foram limpos.")
                st.rerun()
        with col_status:
            st.caption(
                "Use Limpar dados antes de importar arquivos de outra cliente."
            )
        if existem_dados_importados():
            st.warning(
                "Já existem dados carregados nesta sessão. Para analisar outra "
                "cliente, clique em Limpar dados antes de enviar novos arquivos."
            )

        with st.expander("🟣 Belle Software", expanded=True):
            belle_receber_tab, belle_pagar_tab, belle_gerencial_tab = st.tabs([
                "Contas a receber",
                "Contas a pagar",
                "Gerencial de resultados",
            ])
            with belle_receber_tab:
                belle_receber_up = st.file_uploader(
                    "Contas a receber",
                    type=formatos_planilha,
                    key=f"belle_receber_up_{upload_version}",
                    accept_multiple_files=True,
                )
                if belle_receber_up:
                    st.session_state.df_belle_receber = (
                        processar_multiplos_arquivos(
                            belle_receber_up, processar_contas_receber
                        )
                    )
                    if not st.session_state.df_belle_receber.empty:
                        st.success(
                            f"✅ {len(st.session_state.df_belle_receber)} "
                            "contas a receber"
                        )
            with belle_pagar_tab:
                belle_pagar_up = st.file_uploader(
                    "Contas a pagar",
                    type=formatos_planilha,
                    key=f"belle_pagar_up_{upload_version}",
                    accept_multiple_files=True,
                )
                if belle_pagar_up:
                    st.session_state.df_belle_pagar = (
                        processar_multiplos_arquivos(
                            belle_pagar_up, processar_excel
                        )
                    )
                    if not st.session_state.df_belle_pagar.empty:
                        st.success(
                            f"✅ {len(st.session_state.df_belle_pagar)} "
                            "contas a pagar"
                        )
            with belle_gerencial_tab:
                belle_gerencial_up = st.file_uploader(
                    "Gerencial de resultados",
                    type=formatos_planilha,
                    key=f"belle_gerencial_up_{upload_version}",
                    accept_multiple_files=True,
                )
                if belle_gerencial_up:
                    st.session_state.df_belle_gerencial = (
                        processar_multiplos_arquivos(
                            belle_gerencial_up,
                            processar_gerencial_resultados_belle,
                        )
                    )
                    if not st.session_state.df_belle_gerencial.empty:
                        total_gerencial = (
                            st.session_state.df_belle_gerencial["valor"].sum()
                        )
                        st.success(
                            "✅ Despesa com operação de cartão adicionada: "
                            f"{fmt_brl(total_gerencial)}"
                        )

        with st.expander("🔵 Clinicorp", expanded=False):
            clinic1, clinic2, clinic3, clinic4, clinic5 = st.tabs([
                "Contas a pagar",
                "Orçamentos",
                "Extrato Clinipay",
                "Conta corrente",
                "Cartão de crédito",
            ])
            with clinic1:
                clinic_pagar_up = st.file_uploader(
                    "Contas a pagar do Clinicorp",
                    type=formatos_planilha,
                    key=f"clinic_pagar_up_{upload_version}",
                    accept_multiple_files=True,
                )
                if clinic_pagar_up:
                    st.session_state.df_excel = processar_multiplos_arquivos(
                        clinic_pagar_up, processar_excel
                    )
            with clinic2:
                orc_up = st.file_uploader(
                    "Orçamentos",
                    type=formatos_planilha,
                    key=f"orc_up_{upload_version}",
                    accept_multiple_files=True,
                )
                if orc_up:
                    st.session_state.df_orcamentos = processar_multiplos_arquivos(
                        orc_up, processar_orcamentos
                    )
            with clinic3:
                clinipay_up = st.file_uploader(
                    "Extrato Clinipay",
                    type=formatos_planilha,
                    key=f"clinipay_up_{upload_version}",
                    accept_multiple_files=True,
                )
                if clinipay_up:
                    st.session_state.df_clinipay = processar_multiplos_arquivos(
                        clinipay_up, processar_clinipay
                    )
            with clinic4:
                fc_up = st.file_uploader(
                    "Conta corrente do Clinicorp",
                    type=["ofx", *formatos_planilha],
                    key=f"fc_up_{upload_version}",
                    accept_multiple_files=True,
                )
                if fc_up:
                    st.session_state.df_fluxo_caixa = processar_multiplos_arquivos(
                        fc_up, processar_bancos
                    )
            with clinic5:
                antecipacao_up = st.file_uploader(
                    "Cartão de crédito / antecipações",
                    type=["ofx", *formatos_planilha],
                    key=f"antecipacao_up_{upload_version}",
                    accept_multiple_files=True,
                )
                if antecipacao_up:
                    st.session_state.df_antecipacao = (
                        processar_multiplos_arquivos(
                            antecipacao_up, processar_extrato_antecipacao
                        )
                    )

        with st.expander("🟠 Amigotech", expanded=False):
            amigo_receber_tab, amigo_pagar_tab = st.tabs([
                "Entradas / recebimentos",
                "Saídas / despesas",
            ])
            with amigo_receber_tab:
                amigotech_receber_up = st.file_uploader(
                    "Entradas do Amigotech",
                    type=formatos_planilha,
                    key=f"amigotech_receber_up_{upload_version}",
                    accept_multiple_files=True,
                )
                if amigotech_receber_up:
                    st.session_state.df_amigotech_receber = (
                        processar_multiplos_arquivos(
                            amigotech_receber_up,
                            processar_amigotech_receber,
                        )
                    )
                    if not st.session_state.df_amigotech_receber.empty:
                        total_amigo_receber = pd.to_numeric(
                            st.session_state.df_amigotech_receber[
                                "valor_recebido"
                            ],
                            errors="coerce",
                        ).fillna(0).sum()
                        st.success(
                            "✅ Entradas Amigotech importadas: "
                            f"{fmt_brl(total_amigo_receber)}"
                        )
            with amigo_pagar_tab:
                amigotech_pagar_up = st.file_uploader(
                    "Saídas do Amigotech",
                    type=formatos_planilha,
                    key=f"amigotech_pagar_up_{upload_version}",
                    accept_multiple_files=True,
                )
                if amigotech_pagar_up:
                    st.session_state.df_amigotech_pagar = (
                        processar_multiplos_arquivos(
                            amigotech_pagar_up,
                            processar_amigotech_pagar,
                        )
                    )
                    if not st.session_state.df_amigotech_pagar.empty:
                        total_amigo_pagar = (
                            st.session_state.df_amigotech_pagar["valor"].sum()
                        )
                        st.success(
                            "✅ Saídas Amigotech importadas: "
                            f"{fmt_brl(total_amigo_pagar)}"
                        )

        with st.expander("🔷 Conta Azul", expanded=False):
            conta_azul_receber_tab, conta_azul_vendas_tab, conta_azul_pagar_tab = st.tabs([
                "Contas a receber",
                "Vendas",
                "Contas a pagar",
            ])
            with conta_azul_receber_tab:
                conta_azul_receber_up = st.file_uploader(
                    "Contas a receber do Conta Azul",
                    type=formatos_planilha,
                    key=f"conta_azul_receber_up_{upload_version}",
                    accept_multiple_files=True,
                )
                if conta_azul_receber_up:
                    st.session_state.df_conta_azul_receber = (
                        processar_multiplos_arquivos(
                            conta_azul_receber_up,
                            processar_conta_azul_receber,
                        )
                    )
                    if not st.session_state.df_conta_azul_receber.empty:
                        total_ca_receber = (
                            st.session_state.df_conta_azul_receber["valor"].sum()
                        )
                        despesas_ca_receber = (
                            st.session_state.df_conta_azul_receber[
                                ["valor_desconto", "valor_tarifa"]
                            ]
                            .sum(numeric_only=True)
                            .sum()
                            if {
                                "valor_desconto",
                                "valor_tarifa",
                            }.issubset(st.session_state.df_conta_azul_receber.columns)
                            else 0.0
                        )
                        st.success(
                            "✅ Contas a receber Conta Azul importadas: "
                            f"{fmt_brl(total_ca_receber)}"
                            + (
                                " • descontos/tarifas em despesas: "
                                f"{fmt_brl(despesas_ca_receber)}"
                                if despesas_ca_receber > 0 else ""
                            )
                        )
            with conta_azul_vendas_tab:
                conta_azul_vendas_up = st.file_uploader(
                    "Vendas do Conta Azul",
                    type=formatos_planilha,
                    key=f"conta_azul_vendas_up_{upload_version}",
                    accept_multiple_files=True,
                )
                if conta_azul_vendas_up:
                    st.session_state.df_conta_azul_vendas = (
                        processar_multiplos_arquivos(
                            conta_azul_vendas_up,
                            processar_conta_azul_vendas,
                        )
                    )
                    if not st.session_state.df_conta_azul_vendas.empty:
                        total_ca_vendas = (
                            st.session_state.df_conta_azul_vendas["valor"].sum()
                        )
                        st.success(
                            "✅ Vendas Conta Azul importadas: "
                            f"{fmt_brl(total_ca_vendas)}"
                        )
            with conta_azul_pagar_tab:
                conta_azul_pagar_up = st.file_uploader(
                    "Contas a pagar do Conta Azul",
                    type=formatos_planilha,
                    key=f"conta_azul_pagar_up_{upload_version}",
                    accept_multiple_files=True,
                )
                if conta_azul_pagar_up:
                    st.session_state.df_conta_azul_pagar = (
                        processar_multiplos_arquivos(
                            conta_azul_pagar_up,
                            processar_conta_azul_pagar,
                        )
                    )
                    if not st.session_state.df_conta_azul_pagar.empty:
                        total_ca_pagar = (
                            st.session_state.df_conta_azul_pagar["valor"].sum()
                        )
                        st.success(
                            "✅ Contas a pagar Conta Azul importadas: "
                            f"{fmt_brl(total_ca_pagar)}"
                        )

        with st.expander("💳 Extratos de maquininhas em PDF/Excel", expanded=False):
            st.caption(
                "Para OFX de maquininha, use o campo único de Bancos abaixo."
            )
            infinity_up = st.file_uploader(
                "Extrato de maquininha em PDF/Excel/Word",
                type=["pdf", "doc", "docx", *formatos_planilha],
                key=f"infinity_up_{upload_version}",
                accept_multiple_files=True,
            )
            if infinity_up:
                st.session_state.df_infinity_pay = (
                    processar_multiplos_arquivos(
                        infinity_up,
                        processar_infinity_pay,
                    )
                )
                if not st.session_state.df_infinity_pay.empty:
                    entradas = st.session_state.df_infinity_pay.loc[
                        st.session_state.df_infinity_pay["valor"] > 0,
                        "valor",
                    ].sum()
                    saidas = st.session_state.df_infinity_pay.loc[
                        st.session_state.df_infinity_pay["valor"] < 0,
                        "valor",
                    ].abs().sum()
                    st.success(
                        "✅ Extrato de maquininha importado: "
                        f"{fmt_brl(entradas)} em entradas e "
                        f"{fmt_brl_saida(saidas)} em saídas"
                    )

        with st.expander("🏦 Bancos", expanded=False):
            ofx_up = st.file_uploader(
                "Arquivos de bancos em OFX ou Excel",
                type=["ofx", *formatos_planilha],
                key=f"ofx_up_{upload_version}",
                accept_multiple_files=True,
            )
            if ofx_up:
                st.session_state.df_ofx = processar_multiplos_arquivos(
                    ofx_up,
                    processar_bancos,
                )

        st.markdown("---")
        if st.button("✅  Fechar e aplicar dados", type="primary", use_container_width=True, key="fechar_modal"):
            st.session_state.mostrar_modal_upload = False
            st.rerun()

    st.stop()

# =========================
# DADOS PROCESSADOS
# =========================
if (
    st.session_state.df_conta_azul_pagar is not None
    and not st.session_state.df_conta_azul_pagar.empty
):
    conta_azul_pagar_ativa = st.session_state.df_conta_azul_pagar.copy()
    if "descricao" in conta_azul_pagar_ativa.columns:
        descricao_norm = (
            conta_azul_pagar_ativa["descricao"].fillna("").map(normalizar_texto)
        )
        conta_azul_pagar_ativa = conta_azul_pagar_ativa[
            ~descricao_norm.str.contains(
                r"recebimento\s+conta\s+pf|recebimento\s+pf",
                regex=True,
            )
        ].copy()
    quadros_contas_pagar = [conta_azul_pagar_ativa]
elif (
    st.session_state.df_amigotech_pagar is not None
    and not st.session_state.df_amigotech_pagar.empty
):
    quadros_contas_pagar = [st.session_state.df_amigotech_pagar]
elif (
    st.session_state.df_belle_pagar is not None
    and not st.session_state.df_belle_pagar.empty
) or (
    st.session_state.df_belle_gerencial is not None
    and not st.session_state.df_belle_gerencial.empty
):
    quadros_contas_pagar = [
        quadro
        for quadro in [
            st.session_state.df_belle_pagar,
            st.session_state.df_belle_gerencial,
        ]
        if quadro is not None and not quadro.empty
    ]
else:
    quadros_contas_pagar = [
        quadro
        for quadro in [st.session_state.df_excel]
        if quadro is not None and not quadro.empty
    ]
df_exc = (
    pd.concat(quadros_contas_pagar, ignore_index=True, sort=False)
    .drop_duplicates()
    .reset_index(drop=True)
    if quadros_contas_pagar else None
)
df_ofx = st.session_state.df_ofx
df_ofx_movimentos = remover_lancamentos_de_saldo(df_ofx)
df_antecipacao = st.session_state.df_antecipacao
if (
    st.session_state.df_belle_receber is not None
    and not st.session_state.df_belle_receber.empty
):
    df_orcamentos_bruto = st.session_state.df_belle_receber.copy()
    df_orcamentos = df_orcamentos_bruto.copy()
    fonte_vendas_ativa = "Contas a receber da Belle"
elif (
    st.session_state.df_orcamentos is not None
    and not st.session_state.df_orcamentos.empty
):
    df_orcamentos_bruto = st.session_state.df_orcamentos.copy()
    df_orcamentos = df_orcamentos_bruto.copy()
    fonte_vendas_ativa = "Orçamentos do Clinicorp"
elif (
    st.session_state.df_amigotech_receber is not None
    and not st.session_state.df_amigotech_receber.empty
):
    df_orcamentos_bruto = st.session_state.df_amigotech_receber.copy()
    df_orcamentos = df_orcamentos_bruto.copy()
    fonte_vendas_ativa = "Entradas do Amigotech"
elif (
    st.session_state.df_conta_azul_vendas is not None
    and not st.session_state.df_conta_azul_vendas.empty
):
    df_orcamentos_bruto = st.session_state.df_conta_azul_vendas.copy()
    df_orcamentos = df_orcamentos_bruto.copy()
    fonte_vendas_ativa = "Vendas do Conta Azul"
else:
    df_orcamentos_bruto = None
    df_orcamentos = None
    fonte_vendas_ativa = "Nenhuma fonte de vendas"
sistema_vendas_ativo = normalizar_texto(fonte_vendas_ativa)
eh_amigotech_ativo = "amigotech" in sistema_vendas_ativo
rotulo_caixa_recebimentos = (
    "Caixa/Dinheiro" if eh_amigotech_ativo else "Caixa"
)
df_orcamentos = remover_vendas_nao_identificadas_infinity(df_orcamentos)
df_clinipay = st.session_state.df_clinipay
df_fluxo_caixa = st.session_state.df_fluxo_caixa
df_infinity_pay = st.session_state.df_infinity_pay
df_conta_azul_receber = st.session_state.df_conta_azul_receber

quadros_extrato = [
    quadro
    for quadro in [df_ofx_movimentos, df_antecipacao]
    if quadro is not None and not quadro.empty
]
df_extrato_analise = (
    pd.concat(quadros_extrato, ignore_index=True, sort=False)
    if quadros_extrato else pd.DataFrame()
)

inicio_periodo = pd.Timestamp(st.session_state.periodo_inicio_global)
fim_periodo = pd.Timestamp(st.session_state.periodo_fim_global)
periodo_valido = inicio_periodo <= fim_periodo
periodo_label = (
    f"{inicio_periodo.strftime('%d/%m/%Y')} a "
    f"{fim_periodo.strftime('%d/%m/%Y')}"
)

if df_exc is not None and not df_exc.empty:
    df_exc = df_exc.copy()
    df_exc["grupo_custo"] = classificar_grupo_custo(df_exc)
    custos_fixos     = df_exc[df_exc["grupo_custo"] == "Custo Fixo"]["valor"].sum()
    custos_variaveis = df_exc[df_exc["grupo_custo"] == "Custo Variável"]["valor"].sum()
    retiradas        = df_exc[df_exc["grupo_custo"] == "Retirada de Lucro"]["valor"].sum()
    pro_labore       = df_exc[df_exc["grupo_custo"] == "Pró-labore"]["valor"].sum()
    antecipacoes_lucro = df_exc[
        df_exc["grupo_custo"] == "Antecipação de Lucro"
    ]["valor"].sum()
    despesas_total   = custos_fixos + custos_variaveis + pro_labore
else:
    despesas_total = custos_fixos = custos_variaveis = retiradas = pro_labore = antecipacoes_lucro = 0.0

df_exc_periodo = (
    filtrar_por_periodo(df_exc, inicio_periodo, fim_periodo)
    if periodo_valido else pd.DataFrame()
)
df_ofx_periodo = (
    filtrar_por_periodo(df_extrato_analise, inicio_periodo, fim_periodo)
    if periodo_valido else pd.DataFrame()
)
df_banco_periodo = (
    filtrar_por_periodo(df_ofx_movimentos, inicio_periodo, fim_periodo)
    if periodo_valido else pd.DataFrame()
)
df_antecipacao_periodo = (
    filtrar_por_periodo(df_antecipacao, inicio_periodo, fim_periodo)
    if periodo_valido else pd.DataFrame()
)
df_orcamentos_periodo = (
    filtrar_por_periodo(df_orcamentos, inicio_periodo, fim_periodo)
    if periodo_valido else pd.DataFrame()
)
df_orcamentos_bruto_periodo = (
    filtrar_por_periodo(df_orcamentos_bruto, inicio_periodo, fim_periodo)
    if periodo_valido else pd.DataFrame()
)
df_vendas_suspeitas_infinity_periodo = (
    filtrar_vendas_nao_identificadas_infinity(df_orcamentos_bruto_periodo)
)
df_clinipay_periodo = (
    filtrar_por_periodo(df_clinipay, inicio_periodo, fim_periodo)
    if periodo_valido else pd.DataFrame()
)
df_fluxo_periodo = (
    filtrar_por_periodo(df_fluxo_caixa, inicio_periodo, fim_periodo)
    if periodo_valido else pd.DataFrame()
)
if not df_fluxo_periodo.empty:
    arquivos_fluxo = df_fluxo_periodo.get(
        "_arquivo_origem",
        pd.Series("", index=df_fluxo_periodo.index),
    ).fillna("").map(normalizar_texto)
    tipos_fluxo = df_fluxo_periodo.get(
        "tipo_transacao",
        pd.Series("", index=df_fluxo_periodo.index),
    ).fillna("").map(normalizar_texto)
    tipos_lancamento_fluxo = df_fluxo_periodo.get(
        "tipo_lancamento",
        pd.Series("", index=df_fluxo_periodo.index),
    ).fillna("").map(normalizar_texto)
    descricoes_fluxo = df_fluxo_periodo.get(
        "memo",
        pd.Series("", index=df_fluxo_periodo.index),
    ).fillna("").map(normalizar_texto)
    mascara_conta_corrente_clinicorp = (
        arquivos_fluxo.str.contains("conta corrente", regex=False)
        | tipos_lancamento_fluxo.str.contains(
            r"recebimento\s+de\s+pagamento|pagamento\s+recebido\s+paciente",
            regex=True,
        )
        | (
            tipos_fluxo.str.contains("entrada|saida|saída|vendas", regex=True)
            & descricoes_fluxo.str.contains(
                r"reconciliacao\s+de\s+pagamento|reconciliação\s+de\s+pagamento|"
                r"confirmacao\s+pix|confirmação\s+pix|pagamento\s+de\s+tratamento|"
                r"pagamento\s+recebido\s+paciente",
                regex=True,
            )
        )
    )
    df_conta_corrente_clinicorp_periodo = df_fluxo_periodo[
        mascara_conta_corrente_clinicorp
    ].copy()
    df_fluxo_periodo = df_fluxo_periodo[
        ~mascara_conta_corrente_clinicorp
    ].copy()
else:
    df_conta_corrente_clinicorp_periodo = pd.DataFrame()
df_conta_azul_receber_periodo = (
    filtrar_por_periodo(df_conta_azul_receber, inicio_periodo, fim_periodo)
    if periodo_valido else pd.DataFrame()
)
recebimentos_conta_azul_periodo = (
    float(
        pd.to_numeric(
            df_conta_azul_receber_periodo["valor"],
            errors="coerce",
        ).fillna(0).sum()
    )
    if not df_conta_azul_receber_periodo.empty
    and "valor" in df_conta_azul_receber_periodo.columns
    else 0.0
)
usar_recebimentos_conta_azul = recebimentos_conta_azul_periodo > 0
df_infinity_periodo = (
    filtrar_por_periodo(df_infinity_pay, inicio_periodo, fim_periodo)
    if periodo_valido else pd.DataFrame()
)
df_belle_gerencial_periodo = (
    filtrar_por_periodo(
        st.session_state.df_belle_gerencial,
        inicio_periodo,
        fim_periodo,
    )
    if periodo_valido else pd.DataFrame()
)


def filtrar_vendas_validas(df: Optional[pd.DataFrame]) -> pd.DataFrame:
    if df is None or df.empty:
        return pd.DataFrame()
    if "venda_valida" not in df.columns:
        return df.copy()
    mascara = df["venda_valida"].fillna(True).astype(bool)
    return df[mascara].copy()


df_orcamentos_periodo = filtrar_vendas_validas(df_orcamentos_periodo)


def valor_coluna_gerencial(df: pd.DataFrame, coluna: str) -> float:
    if df is None or df.empty or coluna not in df.columns:
        return 0.0
    valores = pd.to_numeric(df[coluna], errors="coerce").dropna()
    return float(valores.max()) if not valores.empty else 0.0


def filtrar_recebimentos_belle_caixa_infinity(
    df: Optional[pd.DataFrame],
) -> pd.DataFrame:
    if df is None or df.empty or "conta_destino" not in df.columns:
        return pd.DataFrame()

    contas = df["conta_destino"].fillna("").map(normalizar_texto)
    mascara = contas.str.contains(
        r"caixa|infinite\s*pay|infinitepay|infinity\s*pay|infinitypay",
        regex=True,
    )
    return df[mascara].copy()


def filtrar_recebimentos_belle_por_conta(
    df: Optional[pd.DataFrame],
    padrao: str,
) -> pd.DataFrame:
    if df is None or df.empty or "conta_destino" not in df.columns:
        return pd.DataFrame()
    contas = df["conta_destino"].fillna("").map(normalizar_texto)
    return df[contas.str.contains(padrao, regex=True)].copy()


def filtrar_vendas_para_conciliacao_ofx(
    df: Optional[pd.DataFrame],
) -> pd.DataFrame:
    if df is None or df.empty:
        return pd.DataFrame()
    vendas = df.copy()
    if "conta_destino" not in vendas.columns:
        return vendas
    contas = vendas["conta_destino"].fillna("").map(normalizar_texto)
    formas = vendas.get(
        "forma",
        pd.Series("", index=vendas.index),
    ).fillna("").map(normalizar_texto)
    dinheiro_direto = (
        contas.str.contains(r"caixa|dinheiro", regex=True)
        | formas.str.contains(r"caixa|dinheiro", regex=True)
    )
    return vendas[~dinheiro_direto].copy()


def filtrar_transferencias_belle(df: Optional[pd.DataFrame]) -> pd.DataFrame:
    if df is None or df.empty or "observacao" not in df.columns:
        return pd.DataFrame()
    observacoes = df["observacao"].fillna("").map(normalizar_texto)
    mascara = observacoes.str.contains(r"transf|transferencia", regex=True)
    return df[mascara].copy()


def somar_recebido_belle(df: Optional[pd.DataFrame]) -> float:
    if df is None or df.empty:
        return 0.0
    coluna = "valor_recebido" if "valor_recebido" in df.columns else "valor"
    return float(pd.to_numeric(df[coluna], errors="coerce").fillna(0).sum())


def filtrar_recebimento_oficial_belle_periodo(
    df: Optional[pd.DataFrame],
    inicio: pd.Timestamp,
    fim: pd.Timestamp,
) -> pd.DataFrame:
    if df is None or df.empty:
        return pd.DataFrame()
    if "valor_recebimento_belle" not in df.columns:
        return pd.DataFrame()

    base = df.copy()
    coluna_data = (
        "data_recebimento_belle"
        if "data_recebimento_belle" in base.columns else "data"
    )
    datas = pd.to_datetime(base[coluna_data], errors="coerce")
    valores = pd.to_numeric(
        base["valor_recebimento_belle"],
        errors="coerce",
    ).fillna(0)
    transferencias = base.get(
        "transferencia_recebimento_belle",
        pd.Series(False, index=base.index),
    ).fillna(False).astype(bool)
    return base[
        datas.between(inicio, fim, inclusive="both")
        & (valores > 0)
        & ~transferencias
    ].copy()


def somar_recebimento_oficial_belle_periodo(
    df: Optional[pd.DataFrame],
    inicio: pd.Timestamp,
    fim: pd.Timestamp,
) -> float:
    base = filtrar_recebimento_oficial_belle_periodo(df, inicio, fim)
    if base.empty:
        return 0.0
    return float(
        pd.to_numeric(
            base["valor_recebimento_belle"],
            errors="coerce",
        ).fillna(0).sum()
    )


def somar_recebimento_maquininha(df: Optional[pd.DataFrame]) -> float:
    if df is None or df.empty:
        return 0.0
    base = df.copy()
    valor_liquido = pd.to_numeric(
        base.get("valor", 0),
        errors="coerce",
    ).fillna(0).abs()
    valor_bruto = pd.to_numeric(
        base.get("valor_bruto_maquininha", 0),
        errors="coerce",
    ).fillna(0).abs()
    antecipada = base.get(
        "antecipada_maquininha",
        pd.Series("", index=base.index),
    ).fillna("").map(normalizar_texto)
    mascara_antecipada = antecipada.str.contains(
        r"\bsim\b|antecipad",
        regex=True,
    )
    valor_receita = valor_liquido.mask(
        mascara_antecipada & (valor_bruto > 0),
        valor_bruto,
    )
    return float(valor_receita.sum())


def despesas_recebimentos_conta_azul(
    df: Optional[pd.DataFrame],
) -> pd.DataFrame:
    if df is None or df.empty:
        return pd.DataFrame()

    despesas = []
    origem = df.copy()
    descricao_origem = origem.get(
        "descricao",
        pd.Series("", index=origem.index),
    ).fillna("").astype(str).str.strip()
    forma_origem = origem.get(
        "conta_destino",
        pd.Series("", index=origem.index),
    ).fillna("").astype(str).str.strip()

    for coluna, categoria, descricao_padrao in [
        ("valor_desconto", "Descontos concedidos", "Desconto em recebimento"),
        ("valor_tarifa", "Tarifas de recebimento", "Tarifa de recebimento"),
    ]:
        if coluna not in origem.columns:
            continue
        valores = pd.to_numeric(origem[coluna], errors="coerce").fillna(0).abs()
        mascara = valores > 0
        if not mascara.any():
            continue
        quadro = pd.DataFrame({
            "data": origem.loc[mascara, "data"],
            "descricao": (
                descricao_padrao + " - "
                + descricao_origem.loc[mascara].replace("", "Conta Azul")
            ),
            "valor": valores.loc[mascara],
            "categoria": categoria,
            "forma": forma_origem.loc[mascara],
            "tipo": "Custo Variável",
            "fonte": "Conta Azul - Descontos e tarifas a receber",
        })
        despesas.append(quadro)

    if not despesas:
        return pd.DataFrame()

    resultado = pd.concat(despesas, ignore_index=True, sort=False)
    resultado["grupo_custo"] = classificar_grupo_custo(resultado)
    return resultado


def creditos_diretos_nao_registrados_conta_azul(
    df_receber: Optional[pd.DataFrame],
    df_banco: Optional[pd.DataFrame],
) -> pd.DataFrame:
    if (
        df_receber is None or df_receber.empty
        or df_banco is None or df_banco.empty
        or "valor" not in df_receber.columns
        or "valor" not in df_banco.columns
    ):
        return pd.DataFrame()

    creditos = df_banco[df_banco["valor"] > 0].copy()
    creditos = filtrar_movimentos_validos_recebimento(creditos)
    if creditos.empty:
        return pd.DataFrame()

    memo_norm = creditos.get(
        "memo",
        pd.Series("", index=creditos.index),
    ).fillna("").map(normalizar_texto)
    creditos = creditos[
        memo_norm.str.contains(r"pix\s+recebido|dep\s+cheque|deposito", regex=True)
    ].copy()
    if creditos.empty:
        return pd.DataFrame()

    creditos["data"] = pd.to_datetime(creditos["data"], errors="coerce").dt.date
    creditos["valor_match"] = (
        pd.to_numeric(creditos["valor"], errors="coerce").fillna(0).round(2)
    )
    creditos = creditos[
        (creditos["valor_match"] > 0)
        & (creditos["valor_match"] <= 1000)
    ].copy()
    if creditos.empty:
        return pd.DataFrame()

    recebimentos = df_receber.copy()
    recebimentos["data_match"] = (
        pd.to_datetime(recebimentos["data"], errors="coerce").dt.date
        if "data" in recebimentos.columns else pd.NaT
    )
    recebimentos["valor_match"] = (
        pd.to_numeric(recebimentos["valor"], errors="coerce").fillna(0).round(2)
    )
    chaves_receber = set(
        zip(recebimentos["data_match"], recebimentos["valor_match"])
    )
    mascara_extra = ~creditos.apply(
        lambda linha: (linha["data"], linha["valor_match"]) in chaves_receber,
        axis=1,
    )
    return creditos[mascara_extra].copy()


def calcular_taxas_maquininha(df: Optional[pd.DataFrame]) -> Dict[str, float]:
    if df is None or df.empty:
        return {
            "taxas": 0.0,
            "bruto": 0.0,
            "liquido": 0.0,
            "taxa_media": 0.0,
        }
    base = df.copy()
    valores = pd.to_numeric(base.get("valor", 0), errors="coerce").fillna(0)
    entradas = base[valores > 0].copy()
    liquido = float(pd.to_numeric(
        entradas.get("valor", 0),
        errors="coerce",
    ).fillna(0).sum())
    taxas_coluna = (
        pd.to_numeric(
            entradas.get("taxa_maquininha", pd.Series(0, index=entradas.index)),
            errors="coerce",
        )
        .fillna(0)
        .abs()
    )
    taxas = float(taxas_coluna.sum())
    bruto_calculado = (
        pd.to_numeric(entradas.get("valor", 0), errors="coerce").fillna(0).abs()
        + taxas_coluna
    )
    bruto = float(bruto_calculado.sum())

    if taxas <= 0:
        texto = (
            base.get("memo", pd.Series("", index=base.index)).fillna("").map(normalizar_texto)
            + " "
            + base.get("tipo_transacao", pd.Series("", index=base.index)).fillna("").map(normalizar_texto)
            + " "
            + base.get("detalhe", pd.Series("", index=base.index)).fillna("").map(normalizar_texto)
        )
        taxas_negativas = base[
            (valores < 0)
            & texto.str.contains(
                r"taxa|tarifa|mdr|desconto|custo|fee",
                regex=True,
            )
        ].copy()
        taxas = float(pd.to_numeric(
            taxas_negativas.get("valor", 0),
            errors="coerce",
        ).fillna(0).abs().sum())
        bruto = liquido + taxas

    taxa_media = (taxas / bruto * 100) if bruto > 0 else 0.0
    return {
        "taxas": taxas,
        "bruto": bruto,
        "liquido": liquido,
        "taxa_media": taxa_media,
    }


def calcular_taxas_clinipay(df: Optional[pd.DataFrame]) -> Dict[str, float]:
    if df is None or df.empty:
        return {
            "bruto": 0.0,
            "taxas": 0.0,
            "juros": 0.0,
            "liquido": 0.0,
            "taxa_media": 0.0,
            "juros_medio": 0.0,
        }
    base = df.copy()
    bruto = pd.to_numeric(
        base.get("valor_bruto_clinipay", base.get("valor", 0)),
        errors="coerce",
    ).fillna(0).abs()
    liquido = pd.to_numeric(base.get("valor", 0), errors="coerce").fillna(0).abs()
    taxas = pd.to_numeric(
        base.get("taxa_clinipay", pd.Series(0, index=base.index)),
        errors="coerce",
    ).fillna(0).abs()
    juros = pd.to_numeric(
        base.get("juros_clinipay", pd.Series(0, index=base.index)),
        errors="coerce",
    ).fillna(0).abs()
    total_bruto = float(bruto.sum())
    total_taxas = float(taxas.sum())
    total_juros = float(juros.sum())
    total_liquido = float(liquido.sum())
    base_juros = float((bruto - juros).clip(lower=0).sum())
    return {
        "bruto": total_bruto,
        "taxas": total_taxas,
        "juros": total_juros,
        "liquido": total_liquido,
        "taxa_media": (total_taxas / total_bruto * 100) if total_bruto > 0 else 0.0,
        "juros_medio": (total_juros / base_juros * 100) if base_juros > 0 else 0.0,
    }


def identificar_creditos_banco_da_infinity(
    df_saidas_infinity: Optional[pd.DataFrame],
    df_banco: Optional[pd.DataFrame],
    df_transferencias_belle: Optional[pd.DataFrame] = None,
    exigir_transferencia_belle: bool = False,
) -> pd.DataFrame:
    if (
        df_saidas_infinity is None or df_saidas_infinity.empty
        or df_banco is None or df_banco.empty
    ):
        return pd.DataFrame()

    creditos = df_banco[df_banco["valor"] > 0].copy()
    if creditos.empty:
        return pd.DataFrame()

    if exigir_transferencia_belle and (
        df_transferencias_belle is None or df_transferencias_belle.empty
    ):
        return pd.DataFrame()

    transferencias_por_valor = {}
    if df_transferencias_belle is not None and not df_transferencias_belle.empty:
        coluna_valor = (
            "valor_recebido"
            if "valor_recebido" in df_transferencias_belle.columns else "valor"
        )
        for valor in pd.to_numeric(
            df_transferencias_belle[coluna_valor],
            errors="coerce",
        ).dropna():
            chave = round(float(abs(valor)), 2)
            if chave <= 0:
                continue
            transferencias_por_valor[chave] = (
                transferencias_por_valor.get(chave, 0) + 1
            )

    usados = set()
    linhas = []
    for _, saida in df_saidas_infinity.iterrows():
        valor_saida = abs(float(saida["valor"]))
        chave_saida = round(valor_saida, 2)
        if transferencias_por_valor:
            if transferencias_por_valor.get(chave_saida, 0) <= 0:
                continue
        candidatos = creditos[
            (~creditos.index.isin(usados))
            & ((creditos["valor"].astype(float) - valor_saida).abs() <= 0.01)
        ]
        if candidatos.empty:
            continue
        if transferencias_por_valor:
            transferencias_por_valor[chave_saida] -= 1
        idx = candidatos.sort_values("data", na_position="last").index[0]
        usados.add(idx)
        linhas.append(creditos.loc[idx])

    return pd.DataFrame(linhas)


def remover_creditos_por_vendas_suspeitas(
    df_creditos: Optional[pd.DataFrame],
    df_suspeitas: Optional[pd.DataFrame],
) -> pd.DataFrame:
    if df_creditos is None or df_creditos.empty:
        return pd.DataFrame()
    if df_suspeitas is None or df_suspeitas.empty:
        return df_creditos.copy()

    creditos = df_creditos.copy()
    coluna_valor = (
        "valor_recebido" if "valor_recebido" in df_suspeitas.columns else "valor"
    )
    valores_suspeitos = pd.to_numeric(
        df_suspeitas[coluna_valor],
        errors="coerce",
    ).dropna()

    indices_remover = set()
    for valor in valores_suspeitos:
        valor = abs(float(valor))
        candidatos = creditos[
            (~creditos.index.isin(indices_remover))
            & ((creditos["valor"].astype(float).abs() - valor).abs() <= 0.01)
        ]
        if candidatos.empty:
            continue
        indices_remover.add(candidatos.index[0])

    if not indices_remover:
        return creditos
    return creditos.drop(index=list(indices_remover)).copy()


recebimento_liquido_gerencial_periodo = valor_coluna_gerencial(
    df_belle_gerencial_periodo,
    "recebimento_liquido_gerencial",
)
vendas_gerencial_periodo = valor_coluna_gerencial(
    df_belle_gerencial_periodo,
    "vendas_gerencial",
)
vendas_gerencial_total = valor_coluna_gerencial(
    st.session_state.df_belle_gerencial,
    "vendas_gerencial",
)
tem_gerencial_belle = (
    st.session_state.df_belle_gerencial is not None
    and not st.session_state.df_belle_gerencial.empty
)
vendas_contas_receber_periodo = (
    float(df_orcamentos_periodo["valor"].sum())
    if not df_orcamentos_periodo.empty else 0.0
)
vendas_referencia_periodo = (
    vendas_gerencial_periodo
    if tem_gerencial_belle
    else vendas_contas_receber_periodo
)
base_recebimentos_belle = df_orcamentos_periodo
recebimentos_oficiais_belle_periodo = (
    somar_recebimento_oficial_belle_periodo(
        st.session_state.df_belle_receber,
        inicio_periodo,
        fim_periodo,
    )
    if "belle" in sistema_vendas_ativo else 0.0
)
usar_recebimentos_oficiais_belle = (
    "belle" in sistema_vendas_ativo
    and recebimentos_oficiais_belle_periodo > 0
)
base_recebimentos_ofx_periodo = filtrar_vendas_para_conciliacao_ofx(
    base_recebimentos_belle
)
vendas_contas_receber_referencia = (
    float(base_recebimentos_belle["valor"].sum())
    if not base_recebimentos_belle.empty else 0.0
)
diferenca_vendas_gerencial_receber = (
    vendas_contas_receber_referencia - vendas_referencia_periodo
    if tem_gerencial_belle else 0.0
)
df_banco_base_periodo = df_banco_periodo
quadros_banco_base_periodo = [
    quadro
    for quadro in [
        df_banco_periodo,
        df_conta_corrente_clinicorp_periodo,
    ]
    if quadro is not None and not quadro.empty
]
df_banco_base_periodo = (
    pd.concat(quadros_banco_base_periodo, ignore_index=True, sort=False)
    if quadros_banco_base_periodo else pd.DataFrame()
)
(
    df_recebimentos_conta_corrente_periodo,
    df_creditos_ignorados_conta_corrente_periodo,
) = separar_recebimentos_conta_corrente_clinicorp(
    df_banco_base_periodo
)
recebimentos_conta_corrente_periodo = (
    float(
        pd.to_numeric(
            df_recebimentos_conta_corrente_periodo.get("valor", 0),
            errors="coerce",
        ).fillna(0).sum()
    )
    if not df_recebimentos_conta_corrente_periodo.empty else 0.0
)
usar_conta_corrente_clinicorp = (
    "clinicorp" in sistema_vendas_ativo
    and not df_recebimentos_conta_corrente_periodo.empty
)
df_infinity_base_periodo = normalizar_bruto_taxa_maquininha(
    df_infinity_periodo
)
if df_clinipay_periodo is not None and not df_clinipay_periodo.empty:
    tipo_relatorio_clinipay = (
        df_clinipay_periodo.get(
            "_tipo_relatorio_clinipay",
            pd.Series("recebiveis", index=df_clinipay_periodo.index),
        )
        .fillna("recebiveis")
        .map(normalizar_texto)
    )
    df_clinipay_base_periodo = df_clinipay_periodo[
        tipo_relatorio_clinipay.eq("recebiveis")
    ].copy()
    df_clinipay_extrato_periodo = df_clinipay_periodo[
        tipo_relatorio_clinipay.eq("extrato")
    ].copy()
else:
    df_clinipay_base_periodo = pd.DataFrame()
    df_clinipay_extrato_periodo = pd.DataFrame()
df_fluxo_base_periodo = df_fluxo_periodo
df_fluxo_base_periodo = remover_fluxo_duplicado_clinipay(
    df_fluxo_base_periodo,
    df_clinipay_base_periodo,
)
df_exc_base_periodo = df_exc_periodo
df_despesas_recebimentos_conta_azul = despesas_recebimentos_conta_azul(
    df_conta_azul_receber
)
df_despesas_recebimentos_conta_azul_periodo = (
    filtrar_por_periodo(
        df_despesas_recebimentos_conta_azul,
        inicio_periodo,
        fim_periodo,
    )
    if periodo_valido else pd.DataFrame()
)
if not df_despesas_recebimentos_conta_azul_periodo.empty:
    if df_exc_base_periodo is None or df_exc_base_periodo.empty:
        df_exc_base_periodo = pd.DataFrame()
    df_exc_base_periodo = pd.concat(
        [
            df_exc_base_periodo,
            df_despesas_recebimentos_conta_azul_periodo,
        ],
        ignore_index=True,
        sort=False,
    )
df_recebimentos_belle_periodo = filtrar_recebimentos_belle_caixa_infinity(
    base_recebimentos_belle
)
df_recebimentos_caixa_belle_periodo = filtrar_recebimentos_belle_por_conta(
    base_recebimentos_belle,
    r"caixa",
)
df_recebimentos_caixa_conta_azul_periodo = filtrar_recebimentos_belle_por_conta(
    df_conta_azul_receber_periodo,
    r"caixa",
)
if not df_recebimentos_caixa_conta_azul_periodo.empty:
    df_recebimentos_caixa_belle_periodo = pd.concat(
        [
            df_recebimentos_caixa_belle_periodo,
            df_recebimentos_caixa_conta_azul_periodo,
        ],
        ignore_index=True,
        sort=False,
    )
df_recebimentos_infinity_belle_periodo = filtrar_recebimentos_belle_por_conta(
    base_recebimentos_belle,
    r"infinite\s*pay|infinitepay|infinity\s*pay|infinitypay",
)
df_transferencias_belle_periodo = filtrar_transferencias_belle(
    base_recebimentos_belle
)
df_entradas_infinity_periodo = (
    filtrar_movimentos_validos_recebimento(
        df_infinity_base_periodo[df_infinity_base_periodo["valor"] > 0].copy()
    )
    if df_infinity_base_periodo is not None and not df_infinity_base_periodo.empty
    else pd.DataFrame()
)
df_saidas_infinity_periodo = (
    df_infinity_base_periodo[df_infinity_base_periodo["valor"] < 0].copy()
    if df_infinity_base_periodo is not None and not df_infinity_base_periodo.empty
    else pd.DataFrame()
)
recebimentos_caixa_belle_periodo = (
    somar_recebido_belle(df_recebimentos_caixa_belle_periodo)
)
recebimentos_infinity_sistema_periodo = (
    somar_recebido_belle(df_recebimentos_infinity_belle_periodo)
)
recebimentos_infinity_extrato_periodo = (
    somar_recebimento_maquininha(df_entradas_infinity_periodo)
    if not df_entradas_infinity_periodo.empty else 0.0
)
saidas_infinity_periodo = (
    float(df_saidas_infinity_periodo["valor"].abs().sum())
    if not df_saidas_infinity_periodo.empty else 0.0
)
taxas_maquininha_periodo = calcular_taxas_maquininha(df_infinity_base_periodo)
df_creditos_banco_infinity_periodo = identificar_creditos_banco_da_infinity(
    df_saidas_infinity_periodo,
    df_banco_base_periodo,
    df_transferencias_belle_periodo,
    exigir_transferencia_belle=True,
)
df_banco_conciliacao_periodo = (
    df_banco_base_periodo.drop(index=df_creditos_banco_infinity_periodo.index)
    if not df_creditos_banco_infinity_periodo.empty
    else df_banco_base_periodo
)
df_banco_conciliacao_periodo = remover_creditos_por_vendas_suspeitas(
    df_banco_conciliacao_periodo,
    df_vendas_suspeitas_infinity_periodo,
)
if recebimentos_infinity_extrato_periodo > 0:
    df_banco_conciliacao_periodo = remover_repasses_maquininha_do_banco(
        df_banco_conciliacao_periodo
    )
df_banco_conciliacao_periodo = remover_repasses_clinipay_do_banco(
    df_banco_conciliacao_periodo,
    df_clinipay_base_periodo,
)
df_creditos_diretos_conta_azul_periodo = (
    creditos_diretos_nao_registrados_conta_azul(
        df_conta_azul_receber_periodo,
        df_banco_conciliacao_periodo,
    )
    if usar_recebimentos_conta_azul else pd.DataFrame()
)
recebimentos_diretos_conta_azul_periodo = (
    float(df_creditos_diretos_conta_azul_periodo["valor"].sum())
    if not df_creditos_diretos_conta_azul_periodo.empty else 0.0
)
recebimentos_conta_azul_periodo = (
    recebimentos_conta_azul_periodo
    + recebimentos_diretos_conta_azul_periodo
)
creditos_banco_infinity_periodo = (
    float(df_creditos_banco_infinity_periodo["valor"].sum())
    if not df_creditos_banco_infinity_periodo.empty else 0.0
)
recebimentos_belle_periodo = (
    recebimentos_caixa_belle_periodo
    + (
        recebimentos_infinity_extrato_periodo
        if recebimentos_infinity_extrato_periodo > 0
        else recebimentos_infinity_sistema_periodo
    )
)
pagamentos_periodo_global = conciliar_pagamentos(
    df_exc_base_periodo,
    df_banco_base_periodo,
)
df_exc_com_despesas_recebimentos = df_exc
if not df_despesas_recebimentos_conta_azul.empty:
    if df_exc_com_despesas_recebimentos is None:
        df_exc_com_despesas_recebimentos = pd.DataFrame()
    df_exc_com_despesas_recebimentos = pd.concat(
        [
            df_exc_com_despesas_recebimentos,
            df_despesas_recebimentos_conta_azul,
        ],
        ignore_index=True,
        sort=False,
    )

vendas_aprovadas = vendas_referencia_periodo
conciliacao_geral = conciliar_recebimentos(
    base_recebimentos_ofx_periodo,
    df_banco_conciliacao_periodo,
    df_clinipay_base_periodo,
    df_fluxo_base_periodo,
)
recebimentos = (
    recebimentos_conta_azul_periodo
    if usar_recebimentos_conta_azul
    else recebimentos_conta_corrente_periodo
    if usar_conta_corrente_clinicorp
    else recebimentos_oficiais_belle_periodo
    if usar_recebimentos_oficiais_belle
    else (
        recebimentos_belle_periodo + conciliacao_geral["total"]
        if recebimentos_belle_periodo > 0
        else conciliacao_geral["total"]
    )
)
pagamentos_geral = conciliar_pagamentos(
    df_exc_com_despesas_recebimentos,
    df_ofx_movimentos,
)
despesas_geral = (
    pagamentos_geral["fixos"]
    + pagamentos_geral["variaveis"]
    + pagamentos_geral.get("pro_labore", 0.0)
)
saidas_lucro_geral = (
    pagamentos_geral["retiradas"]
    + pagamentos_geral.get("antecipacoes_lucro", 0.0)
)
resultado = (
    recebimentos
    - despesas_geral
    - saidas_lucro_geral
)

def criar_resumo_periodo() -> Dict[str, float]:
    despesas_periodo = (
        pagamentos_periodo_global["fixos"]
        + pagamentos_periodo_global["variaveis"]
        + pagamentos_periodo_global.get("pro_labore", 0.0)
    )
    retiradas_periodo = pagamentos_periodo_global["retiradas"]
    pro_labore_periodo = pagamentos_periodo_global.get("pro_labore", 0.0)
    antecipacoes_lucro_periodo = pagamentos_periodo_global.get(
        "antecipacoes_lucro",
        0.0,
    )
    saidas_lucro_periodo = (
        retiradas_periodo
        + antecipacoes_lucro_periodo
    )
    resultado_operacional_periodo = (
        recebimentos - despesas_periodo - saidas_lucro_periodo
    )
    resultado_final_periodo = resultado_operacional_periodo
    return {
        "recebimentos": float(recebimentos),
        "despesas": float(despesas_periodo),
        "custos_fixos": float(pagamentos_periodo_global["fixos"]),
        "custos_variaveis": float(pagamentos_periodo_global["variaveis"]),
        "retiradas": float(retiradas_periodo),
        "pro_labore": float(pro_labore_periodo),
        "antecipacoes_lucro": float(antecipacoes_lucro_periodo),
        "vendas": float(vendas_referencia_periodo),
        "resultado_operacional": float(resultado_operacional_periodo),
        "resultado_final": float(resultado_final_periodo),
    }


@st.dialog("💾 Salvar relatório", width="large")
def abrir_salvar_relatorio():
    st.caption(
        "Guarde uma cópia interna deste relatório para consultar e comparar "
        "com outros meses depois."
    )
    mes_padrao = str(st.session_state.periodo_inicio_global.month).zfill(2)
    ano_padrao = st.session_state.periodo_inicio_global.year
    cliente = st.selectbox("Cliente", CLIENTES_RELATORIO)
    col_mes, col_ano = st.columns([1.4, 1])
    with col_mes:
        mes_referencia = st.selectbox(
            "Mês de referência",
            [mes for mes, _ in MESES_RELATORIO],
            index=max(0, int(mes_padrao) - 1),
            format_func=lambda mes: dict(MESES_RELATORIO).get(mes, mes),
        )
    with col_ano:
        ano_referencia = st.number_input(
            "Ano",
            min_value=2020,
            max_value=2100,
            value=int(ano_padrao),
            step=1,
        )
    referencia = rotulo_mes_referencia(mes_referencia, int(ano_referencia))
    titulo = st.text_input(
        "Nome do relatório",
        value=f"Fechamento financeiro — {referencia}",
    )
    resumo = criar_resumo_periodo()
    resumo["mes_referencia"] = referencia
    resumo["cliente"] = cliente
    c1, c2, c3, c4 = st.columns(4)
    c1.metric("Recebimentos", fmt_brl(resumo["recebimentos"]))
    c2.metric("Despesas", fmt_brl(resumo["despesas"]))
    c3.metric("Vendas", fmt_brl(resumo["vendas"]))
    c4.metric("Resultado", fmt_brl(resumo["resultado_final"], sinal=True))

    if st.button("Salvar no histórico", type="primary", use_container_width=True):
        if not cliente.strip() or not titulo.strip():
            st.error("Informe o nome da cliente e do relatório.")
            return
        create_saved_report(
            user.id,
            cliente.strip(),
            titulo.strip(),
            criar_payload_compartilhado(user.id),
            json.dumps(resumo, ensure_ascii=False),
            str(st.session_state.periodo_inicio_global),
            str(st.session_state.periodo_fim_global),
        )
        st.session_state.mostrar_salvar_relatorio = False
        st.success("Relatório salvo no histórico.")
        st.rerun()


if st.session_state.mostrar_salvar_relatorio and not st.session_state.share_mode:
    st.session_state.mostrar_salvar_relatorio = False
    abrir_salvar_relatorio()

if (
    st.session_state.share_mode
    and st.session_state.pagina not in {
        "visao", "fechamento", "fornecedores", "comparativo", "detalhes", "saldo"
    }
):
    st.session_state.pagina = "visao"

# =========================
# PÁGINA: ADMIN
# =========================
if st.session_state.pagina == "admin":
    if not user.is_admin:
        st.error("❌ Acesso negado.")
        st.stop()

    st.markdown('<div class="page-title">⚙️ Painel Administrativo</div>', unsafe_allow_html=True)
    st.markdown('<div class="page-subtitle">Gerencie usuários e configurações do sistema.</div>', unsafe_allow_html=True)

    tab_criar, tab_listar, tab_senha, tab_deletar = st.tabs([
        "➕ Criar Usuário",
        "📋 Listar Usuários",
        "🔑 Recuperar Senha",
        "🗑️ Deletar Usuário",
    ])

    with tab_criar:
        st.markdown("### Criar novo usuário")
        col1, col2 = st.columns(2)
        with col1:
            novo_user = st.text_input("Nome de usuário", placeholder="exemplo_user")
        with col2:
            novo_pass = st.text_input("Senha inicial", type="password", placeholder="••••••••")
        col_btn1, col_btn2, _ = st.columns([1, 1, 2])
        with col_btn1:
            if st.button("✅ Criar usuário", type="primary", use_container_width=True):
                is_valid, msg = validar_entrada_usuario(novo_user, novo_pass)
                if not is_valid:
                    st.error(msg)
                else:
                    try:
                        success = create_user(novo_user, novo_pass)
                        if success:
                            logger.info(f"New user created by admin: {novo_user}")
                            st.success(f"✅ Usuário '{novo_user}' criado!")
                            st.info(
                                "Informe a senha inicial ao usuário por um canal "
                                "seguro. Ela não será exibida novamente."
                            )
                        else:
                            st.error(f"❌ Usuário '{novo_user}' já existe.")
                    except Exception as e:
                        st.error(f"❌ Erro: {str(e)}")
        with col_btn2:
            if st.button("🔄 Limpar formulário", use_container_width=True):
                st.rerun()

    with tab_listar:
        st.markdown("### Usuários cadastrados")
        try:
            usuarios = get_all_users()
            if usuarios:
                df_users = pd.DataFrame(
                    usuarios,
                    columns=["ID", "Usuário", "Administrador"],
                )
                df_users["Administrador"] = df_users["Administrador"].map(
                    {0: "Não", 1: "Sim"}
                )
                st.dataframe(df_users, use_container_width=True, hide_index=True)
                st.info(f"📊 Total: {len(usuarios)} usuário(s)")
            else:
                st.warning("Nenhum usuário encontrado.")
        except Exception as e:
            st.error(f"❌ Erro: {str(e)}")

    with tab_senha:
        st.markdown("### Emitir código de recuperação")
        st.caption(
            "O usuário poderá trocar a própria senha na tela de login. "
            "O código é de uso único e expira em 15 minutos."
        )
        try:
            usuarios = get_all_users()
            usernames = [u[1] for u in usuarios]
            if usernames:
                reset_username = st.selectbox(
                    "Usuário",
                    usernames,
                    key="reset_username_admin",
                )
                admin_password_reset = st.text_input(
                    "Confirme sua senha de administrador",
                    type="password",
                    key="admin_password_reset",
                )
                if st.button(
                    "Gerar código temporário",
                    type="primary",
                    key="gerar_codigo_senha",
                ):
                    if not verify_admin_password(user.id, admin_password_reset):
                        st.error("Senha de administrador incorreta.")
                    else:
                        codigo = create_password_reset_code(reset_username)
                        if codigo:
                            st.session_state.codigo_recuperacao_exibido = (
                                reset_username,
                                codigo,
                            )
                        else:
                            st.error("Não foi possível gerar o código.")

                codigo_exibido = st.session_state.get(
                    "codigo_recuperacao_exibido"
                )
                if codigo_exibido:
                    usuario_codigo, codigo = codigo_exibido
                    st.success(
                        f"Código para **{usuario_codigo}**. Envie-o ao usuário "
                        "por um canal seguro."
                    )
                    st.code(codigo, language=None)
            else:
                st.info("Nenhum usuário cadastrado.")
        except Exception as e:
            logger.exception("Erro ao gerar código de recuperação")
            st.error(f"❌ Erro: {str(e)}")

    with tab_deletar:
        st.markdown("### Deletar usuário")
        st.warning(
            "⚠️ Esta ação remove também contas, links compartilhados e códigos "
            "de recuperação do usuário."
        )
        try:
            usuarios = get_all_users()
            if usuarios:
                usernames = [u[1] for u in usuarios if not bool(u[2])]
                if usernames:
                    user_to_delete = st.selectbox("Selecione o usuário", usernames)
                    if (
                        st.session_state.usuario_exclusao_pendente
                        != user_to_delete
                    ):
                        if st.button(
                            "Continuar para confirmação",
                            type="secondary",
                            use_container_width=True,
                        ):
                            st.session_state.usuario_exclusao_pendente = (
                                user_to_delete
                            )
                            st.rerun()
                    else:
                        st.error(
                            f"Confirmação final para excluir **{user_to_delete}**."
                        )
                        confirmacao_nome = st.text_input(
                            f'Digite "{user_to_delete}" para confirmar',
                            key="confirmacao_nome_exclusao",
                        )
                        confirmacao_senha = st.text_input(
                            "Digite sua senha de administrador",
                            type="password",
                            key="confirmacao_senha_exclusao",
                        )
                        col_confirmar, col_cancelar = st.columns(2)
                        with col_confirmar:
                            if st.button(
                                "Excluir definitivamente",
                                type="primary",
                                use_container_width=True,
                            ):
                                if confirmacao_nome != user_to_delete:
                                    st.error(
                                        "O nome digitado não corresponde ao usuário."
                                    )
                                elif not verify_admin_password(
                                    user.id,
                                    confirmacao_senha,
                                ):
                                    st.error("Senha de administrador incorreta.")
                                else:
                                    delete_user_by_admin(user_to_delete)
                                    st.session_state.usuario_exclusao_pendente = None
                                    st.success(
                                        f"✅ '{user_to_delete}' foi excluído."
                                    )
                                    st.rerun()
                        with col_cancelar:
                            if st.button(
                                "Cancelar",
                                use_container_width=True,
                                key="cancelar_exclusao_usuario",
                            ):
                                st.session_state.usuario_exclusao_pendente = None
                                st.rerun()
                else:
                    st.info("Nenhum usuário disponível para deletar.")
        except Exception as e:
            st.error(f"❌ Erro: {str(e)}")

# =========================
# PÁGINA: VISÃO FINANCEIRA
# =========================
elif st.session_state.pagina == "visao":
    st.markdown('<div class="page-title">Visão Financeira</div>', unsafe_allow_html=True)
    st.markdown(
        '<div class="page-subtitle">O período global do menu lateral atualiza toda a análise.</div>',
        unsafe_allow_html=True,
    )

    if not periodo_valido:
        st.error("❌ A data inicial não pode ser posterior à data final.")

    st.markdown(
        f"""
        <div class="period-summary">
            <span><strong>Período selecionado</strong> · {periodo_label}</span>
            <span>🟢 Receita &nbsp;&nbsp; 🔴 Despesa</span>
        </div>
        """,
        unsafe_allow_html=True,
    )
    vendas_aprovadas_periodo = (
        vendas_referencia_periodo
    )
    conciliacao_periodo = conciliar_recebimentos(
        base_recebimentos_ofx_periodo,
        df_banco_conciliacao_periodo,
        df_clinipay_base_periodo,
        df_fluxo_base_periodo,
    )
    recebimentos_periodo = (
        recebimentos_conta_azul_periodo
        if usar_recebimentos_conta_azul
        else recebimentos_conta_corrente_periodo
        if usar_conta_corrente_clinicorp
        else recebimentos_oficiais_belle_periodo
        if usar_recebimentos_oficiais_belle
        else (
            recebimentos_belle_periodo + conciliacao_periodo["total"]
            if recebimentos_belle_periodo > 0
            else conciliacao_periodo["total"]
        )
    )
    antecipacoes_periodo = calcular_antecipacoes(df_antecipacao_periodo)
    antecipacoes_maquininha_periodo = calcular_antecipacoes_maquininha(
        df_infinity_base_periodo
    )
    antecipacoes_periodo = {
        chave: (
            float(antecipacoes_periodo.get(chave, 0.0))
            + float(antecipacoes_maquininha_periodo.get(chave, 0.0))
        )
        for chave in ["recebido", "custo", "liquido"]
    }
    pagamentos_periodo = pagamentos_periodo_global
    custos_fixos_periodo = pagamentos_periodo["fixos"]
    custos_variaveis_periodo = pagamentos_periodo["variaveis"]
    retiradas_periodo = pagamentos_periodo["retiradas"]
    pro_labore_periodo = pagamentos_periodo.get("pro_labore", 0.0)
    antecipacoes_lucro_periodo = pagamentos_periodo.get(
        "antecipacoes_lucro",
        0.0,
    )
    despesas_periodo = (
        custos_fixos_periodo
        + custos_variaveis_periodo
        + pro_labore_periodo
    )
    saidas_lucro_periodo = (
        retiradas_periodo
        + antecipacoes_lucro_periodo
    )
    resultado_operacional_periodo = (
        recebimentos_periodo - despesas_periodo - saidas_lucro_periodo
    )
    resultado_periodo = resultado_operacional_periodo
    resultado_antes_retiradas_periodo = recebimentos_periodo - despesas_periodo

    _, topo_assistente = st.columns([4, 1.15])
    with topo_assistente:
        st.markdown('<div class="assistant-top-button">', unsafe_allow_html=True)
        if st.button(
            "💬 Assistente",
            use_container_width=True,
            type="primary",
            key="abrir_assistente_financeiro_topo",
        ):
            abrir_assistente(
                periodo_label,
                recebimentos_periodo,
                vendas_aprovadas_periodo,
                despesas_periodo,
                custos_fixos_periodo,
                custos_variaveis_periodo,
                retiradas_periodo,
                antecipacoes_lucro_periodo,
            )
        st.markdown('</div>', unsafe_allow_html=True)

    k1, k2, k3, k4 = st.columns(4)
    with k1:
        st.markdown(f"""
            <div class="kpi-card green">
                <div class="kpi-label">Recebimentos</div>
                <div class="kpi-value green">{fmt_brl(recebimentos_periodo)}</div>
            <div class="kpi-footer">{"Conta corrente oficial" if usar_conta_corrente_clinicorp else "Caixa, maquininhas e banco direto" if recebimentos_belle_periodo > 0 else "Recebimentos identificados das vendas"}</div>
            </div>""", unsafe_allow_html=True)
    with k2:
        st.markdown(f"""
        <div class="kpi-card neutral">
            <div class="kpi-label">Despesas</div>
            <div class="kpi-value red">{fmt_brl_saida(despesas_periodo)}</div>
            <div class="kpi-footer">Custos fixos + variáveis + pró-labore</div>
        </div>""", unsafe_allow_html=True)
    with k3:
        st.markdown(f"""
        <div class="kpi-card blue">
            <div class="kpi-label">Vendas Aprovadas</div>
            <div class="kpi-value blue">{fmt_brl(vendas_aprovadas_periodo)}</div>
            <div class="kpi-footer">{"Gerencial de resultados Belle" if tem_gerencial_belle else fonte_vendas_ativa}</div>
        </div>""", unsafe_allow_html=True)
    with k4:
        cor    = "green" if resultado_periodo >= 0 else "red"
        margem = (
            resultado_periodo / recebimentos_periodo * 100
            if recebimentos_periodo > 0 else 0
        )
        st.markdown(f"""
        <div class="kpi-card purple">
            <div class="kpi-label">Resultado final</div>
            <div class="kpi-value {cor}">{fmt_brl(resultado_periodo, sinal=True)}</div>
            <div class="kpi-footer">
                {fmt_brl(resultado_antes_retiradas_periodo)} antes das retiradas ·
                {fmt_brl_saida(saidas_lucro_periodo)} em retiradas · Margem: {margem:.1f}%
            </div>
        </div>""", unsafe_allow_html=True)

    st.markdown("<br>", unsafe_allow_html=True)
    st.markdown(
        """
        <div style="font-size:0.7rem;font-weight:700;letter-spacing:0.07em;
                    color:#4A4A7A;text-transform:uppercase;margin-bottom:0.8rem">
            Antecipação de Recebíveis
        </div>
        """,
        unsafe_allow_html=True,
    )
    ant1, ant2, ant3 = st.columns([1, 1, 1.5])
    with ant1:
        st.markdown(
            f"""
            <div class="kpi-card green">
                <div class="kpi-label">Valor antecipado (bruto)</div>
                <div class="kpi-value green">{fmt_brl(antecipacoes_periodo["recebido"])}</div>
                <div class="kpi-footer">Total antes das taxas</div>
            </div>
            """,
            unsafe_allow_html=True,
        )
    with ant2:
        st.markdown(
            f"""
            <div class="kpi-card neutral">
                <div class="kpi-label">Custo da antecipação</div>
                <div class="kpi-value red">{fmt_brl_saida(antecipacoes_periodo["custo"])}</div>
                <div class="kpi-footer">Taxas e débitos identificados</div>
            </div>
            """,
            unsafe_allow_html=True,
        )
    with ant3:
        valores_antecipacao = [
            antecipacoes_periodo["recebido"],
            antecipacoes_periodo["custo"],
        ]
        fig_antecipacao = go.Figure(
            go.Bar(
                x=["Valor bruto", "Custo pago"],
                y=valores_antecipacao,
                marker_color=["#1D9E75", "#F09595"],
                text=[fmt_brl(v) for v in valores_antecipacao],
                textposition="auto",
            )
        )
        fig_antecipacao.update_layout(
            height=185,
            margin=dict(l=0, r=0, t=8, b=0),
            paper_bgcolor="#11112A",
            plot_bgcolor="#11112A",
            showlegend=False,
            xaxis=dict(showgrid=False, tickfont=dict(color="#A0A0C0")),
            yaxis=dict(showgrid=True, gridcolor="#1E1E3A", tickfont=dict(color="#4A4A7A")),
        )
        st.plotly_chart(
            fig_antecipacao,
            use_container_width=True,
            config={"displayModeBar": False},
        )

    st.caption(
        f"Valor líquido recebido no período: "
        f"**{fmt_brl(antecipacoes_periodo['liquido'])}**. "
        "A identificação é feita pelas descrições de antecipação no extrato OFX."
    )

    st.markdown("<br>", unsafe_allow_html=True)
    col_graf, col_comp = st.columns([1.6, 1])

    with col_graf:
        st.markdown("""
        <div style="font-size:0.7rem;font-weight:700;letter-spacing:0.07em;
                    color:#4A4A7A;text-transform:uppercase;margin-bottom:0.8rem">
            Fluxo Diário &nbsp;&nbsp;
            <span style="color:#1D9E75">● Entrada</span> &nbsp;
            <span style="color:#F09595">● Saída</span>
        </div>""", unsafe_allow_html=True)

        creditos_periodo = conciliacao_periodo["creditos_conciliados"]
        contas_pagas_periodo = pagamentos_periodo["contas_conciliadas"]
        if not creditos_periodo.empty and not contas_pagas_periodo.empty:
            entradas = creditos_periodo.groupby(
                creditos_periodo["data"].dt.date
            )["valor"].sum().reset_index()
            saidas = contas_pagas_periodo.groupby(
                contas_pagas_periodo["data"].dt.date
            )["valor"].sum().reset_index()
            entradas.columns = saidas.columns = ["data", "valor"]
            fig = go.Figure()
            fig.add_trace(go.Scatter(x=entradas["data"], y=entradas["valor"], mode="lines",
                line=dict(color="#1D9E75", width=2), fill="tozeroy", fillcolor="rgba(29,158,117,0.12)"))
            fig.add_trace(go.Scatter(x=saidas["data"], y=saidas["valor"], mode="lines",
                line=dict(color="#F09595", width=2), fill="tozeroy", fillcolor="rgba(240,149,149,0.12)"))
            fig.update_layout(
                height=220, margin=dict(l=0,r=0,t=0,b=0),
                paper_bgcolor="#11112A", plot_bgcolor="#11112A", showlegend=False,
                xaxis=dict(showgrid=False, showline=False, tickfont=dict(size=10, color="#4A4A7A")),
                yaxis=dict(showgrid=True, gridcolor="#1E1E3A", showline=False, tickfont=dict(size=10, color="#4A4A7A")),
            )
            st.plotly_chart(fig, use_container_width=True, config={"displayModeBar": False})
        else:
            st.markdown("""
            <div style="height:220px;background:#11112A;border-radius:10px;border:1px solid #1E1E3A;
                        display:flex;align-items:center;justify-content:center;
                        color:#4A4A7A;font-size:0.82rem">
                📊 Importe os arquivos para ver o fluxo diário
            </div>""", unsafe_allow_html=True)

    with col_comp:
        st.markdown("""
        <div style="font-size:0.7rem;font-weight:700;letter-spacing:0.07em;
                    color:#4A4A7A;text-transform:uppercase;margin-bottom:0.8rem">
            Custos Operacionais
        </div>""", unsafe_allow_html=True)

        itens_comp = [
            ("Custo Fixo",        custos_fixos_periodo,     "#7F77DD"),
            ("Custo Variável",    custos_variaveis_periodo, "#F09595"),
            ("Pró-labore",        pro_labore_periodo,       "#EF9F27"),
        ]
        total_comp = sum(v for _, v, _ in itens_comp)
        total_retiradas_comp = (
            retiradas_periodo
            + antecipacoes_lucro_periodo
        )

        if total_comp <= 0:
            st.markdown('<div style="font-size:0.78rem;color:#4A4A7A;text-align:center;padding:2rem">Sem dados de despesas</div>', unsafe_allow_html=True)
        else:
            for label, valor, cor in itens_comp:
                if valor > 0:
                    pct = valor / total_comp * 100
                    st.markdown(f"""
                    <div class="comp-row">
                        <span>{label}</span>
                <span style="color:#F09595;font-family:'DM Mono',monospace;font-size:0.78rem">{fmt_brl_saida(valor)}</span>
                    </div>
                    <div class="comp-bar-bg">
                        <div class="comp-bar-fill" style="width:{pct:.0f}%;background:{cor}"></div>
                    </div>""", unsafe_allow_html=True)

        if retiradas_periodo > 0:
            pct_retirada = (
                retiradas_periodo / total_retiradas_comp * 100
                if total_retiradas_comp > 0 else 0
            )
            st.markdown(f"""
            <div class="comp-row" style="margin-top:0.9rem">
                <span>Retirada de lucro</span>
                <span style="color:#EF9F27;font-family:'DM Mono',monospace;font-size:0.78rem">{fmt_brl_saida(retiradas_periodo)}</span>
            </div>
            <div class="comp-bar-bg">
                <div class="comp-bar-fill" style="width:{pct_retirada:.0f}%;background:#EF9F27"></div>
            </div>""", unsafe_allow_html=True)

        if antecipacoes_lucro_periodo > 0:
            pct_antecipacao = (
                antecipacoes_lucro_periodo / total_retiradas_comp * 100
                if total_retiradas_comp > 0 else 0
            )
            st.markdown(f"""
            <div class="comp-row" style="margin-top:0.9rem">
                <span>Antecipação de lucro</span>
                <span style="color:#EF9F27;font-family:'DM Mono',monospace;font-size:0.78rem">{fmt_brl_saida(antecipacoes_lucro_periodo)}</span>
            </div>
            <div class="comp-bar-bg">
                <div class="comp-bar-fill" style="width:{pct_antecipacao:.0f}%;background:#EF9F27"></div>
            </div>""", unsafe_allow_html=True)

        st.markdown("""
        <div style="font-size:0.7rem;font-weight:700;letter-spacing:0.07em;
                    color:#4A4A7A;text-transform:uppercase;margin-top:1.2rem;margin-bottom:0.8rem">
            Formas de Recebimento
        </div>""", unsafe_allow_html=True)

        if recebimentos_periodo > 0:
            total_rec = recebimentos_periodo
            if total_rec > 0:
                pix_val = conciliacao_periodo["pix"]
                cart_val = conciliacao_periodo["cartao"]
                antecipacao_val = conciliacao_periodo["antecipacao"]
                bancario_val = conciliacao_periodo.get("bancario", 0.0)
                clinipay_val = conciliacao_periodo["clinipay"]
                direto_val = conciliacao_periodo["direto"]
                for lbl, val, cor in [
                    ("PIX conciliado", pix_val, "#1D9E75"),
                    ("Cartão conciliado", cart_val, "#7F77DD"),
                    ("Antecipação conciliada", antecipacao_val, "#D85A30"),
                    ("Crédito bancário", bancario_val, "#62628F"),
                    ("Clinipay", clinipay_val, "#85B7EB"),
                    ("Pagamento direto", direto_val, "#EF9F27"),
                ]:
                    if val > 0:
                        pct = val / total_rec * 100
                        st.markdown(f"""
                        <div class="comp-row">
                            <span>{lbl}</span>
                            <span style="color:#1D9E75;font-family:'DM Mono',monospace;font-size:0.78rem">{fmt_brl(val)}</span>
                        </div>
                        <div class="comp-bar-bg">
                            <div class="comp-bar-fill" style="width:{pct:.0f}%;background:{cor}"></div>
                        </div>""", unsafe_allow_html=True)
        else:
            mensagem_recebimentos = (
                "Nenhum recebimento identificado no período"
                if not df_banco_base_periodo.empty else "Importe o OFX para ver"
            )
            st.markdown(
                f'<div style="font-size:0.78rem;color:#4A4A7A">{mensagem_recebimentos}</div>',
                unsafe_allow_html=True,
            )

# =========================
# PÁGINA: FECHAMENTO
# =========================
elif st.session_state.pagina == "fechamento":
    st.markdown('<div class="page-title">Fechamento mensal</div>', unsafe_allow_html=True)
    st.markdown(
        f'<div class="page-subtitle">DRE simplificado · {periodo_label}.</div>',
        unsafe_allow_html=True,
    )
    pagamentos_fechamento = pagamentos_periodo_global
    custos_fixos_fechamento = pagamentos_fechamento["fixos"]
    custos_variaveis_fechamento = pagamentos_fechamento["variaveis"]
    retiradas_fechamento = pagamentos_fechamento["retiradas"]
    pro_labore_fechamento = pagamentos_fechamento.get("pro_labore", 0.0)
    antecipacoes_lucro_fechamento = pagamentos_fechamento.get(
        "antecipacoes_lucro",
        0.0,
    )
    conciliacao_fechamento = conciliar_recebimentos(
        base_recebimentos_ofx_periodo,
        df_banco_conciliacao_periodo,
        df_clinipay_base_periodo,
        df_fluxo_base_periodo,
    )
    recebimentos_vendas = (
        recebimentos_conta_azul_periodo
        if usar_recebimentos_conta_azul
        else recebimentos_conta_corrente_periodo
        if usar_conta_corrente_clinicorp
        else (
            recebimentos_belle_periodo + conciliacao_fechamento["total"]
            if recebimentos_belle_periodo > 0
            else conciliacao_fechamento["total"]
        )
    )
    vendas_fechamento = (
        vendas_referencia_periodo
    )
    total_despesas = (
        custos_fixos_fechamento
        + custos_variaveis_fechamento
        + pro_labore_fechamento
    )
    saidas_lucro_fechamento = (
        retiradas_fechamento
        + antecipacoes_lucro_fechamento
    )
    resultado_operacional = (
        recebimentos_vendas - total_despesas - saidas_lucro_fechamento
    )
    resultado_mes = resultado_operacional
    cor_operacional = "#1D9E75" if resultado_operacional >= 0 else "#F09595"
    cor_resultado   = "#1D9E75" if resultado_mes >= 0 else "#F09595"
    sinal           = "+" if resultado_mes >= 0 else ""
    sinal_operacional = "+" if resultado_operacional >= 0 else ""

    st.markdown(f"""
    <table class="dre-table">
      <thead><tr><th>Demonstrativo de Resultado</th><th>Valor</th></tr></thead>
      <tbody>
        <tr><td>Total de vendas aprovadas</td><td style="color:#85B7EB">{fmt_brl(vendas_fechamento)}</td></tr>
        <tr><td>Total recebido dessas vendas</td><td style="color:#1D9E75">{fmt_brl(recebimentos_vendas)}</td></tr>
        <tr class="section-title"><td colspan="2">Despesas operacionais</td></tr>
        <tr><td>&nbsp;&nbsp;&nbsp;Custos fixos</td><td style="color:#F09595">{fmt_brl(custos_fixos_fechamento)}</td></tr>
        <tr><td>&nbsp;&nbsp;&nbsp;Custos variáveis</td><td style="color:#F09595">{fmt_brl(custos_variaveis_fechamento)}</td></tr>
        <tr><td>&nbsp;&nbsp;&nbsp;Pró-labore</td><td style="color:#F09595">{fmt_brl_saida(pro_labore_fechamento)}</td></tr>
        <tr class="total-row"><td>Total despesas operacionais</td><td>{fmt_brl_saida(total_despesas)}</td></tr>
        <tr class="section-title"><td colspan="2">Distribuição de lucro</td></tr>
        <tr><td>&nbsp;&nbsp;&nbsp;Retirada de lucro</td><td style="color:#F09595">{fmt_brl_saida(retiradas_fechamento)}</td></tr>
        <tr><td>&nbsp;&nbsp;&nbsp;Antecipação de lucro</td><td style="color:#F09595">{fmt_brl_saida(antecipacoes_lucro_fechamento)}</td></tr>
        <tr class="resultado-row"><td>Resultado (recebimentos - despesas - retiradas)</td>
            <td style="color:{cor_resultado}">{sinal}{fmt_brl(resultado_mes)}</td></tr>
      </tbody>
    </table>""", unsafe_allow_html=True)

# =========================
# PÁGINA: FORNECEDORES
# =========================
elif st.session_state.pagina == "fornecedores":
    st.markdown('<div class="page-title">Fornecedores</div>', unsafe_allow_html=True)
    st.markdown(
        f'<div class="page-subtitle">Materiais, medicamentos, insumos, exames e itens ligados à produção · {periodo_label}.</div>',
        unsafe_allow_html=True,
    )

    despesas_fornecedor = pagamentos_periodo_global[
        "contas_conciliadas"
    ].copy()
    if despesas_fornecedor.empty:
        st.info("📥 Importe contas a pagar para visualizar os fornecedores.")
    else:
        despesas_fornecedor = filtrar_fornecedores_producao(despesas_fornecedor)
        if despesas_fornecedor.empty:
            st.info(
                "Nenhum gasto de produção encontrado no período. Este painel "
                "considera materiais e medicamentos, remédios, insumos clínicos, "
                "exames, material médico e material odontológico."
            )
        else:
            if "fornecedor" in despesas_fornecedor.columns:
                nomes_fornecedores = despesas_fornecedor["fornecedor"]
            elif "descricao" in despesas_fornecedor.columns:
                nomes_fornecedores = despesas_fornecedor["descricao"]
            else:
                nomes_fornecedores = pd.Series(
                    "Sem fornecedor identificado",
                    index=despesas_fornecedor.index,
                )
            despesas_fornecedor["fornecedor"] = (
                nomes_fornecedores
                .fillna("")
                .astype(str)
                .str.strip()
                .replace("", "Sem fornecedor identificado")
            )
            resumo_fornecedores = (
                despesas_fornecedor.groupby("fornecedor", as_index=False)
                .agg(
                    valor=("valor", "sum"),
                    lancamentos=("valor", "count"),
                    ultima_data=("data", "max"),
                    categoria=("categoria", lambda s: ", ".join(
                        s.dropna().astype(str).replace("", pd.NA).dropna().unique()[:2]
                    )),
                )
                .sort_values("valor", ascending=False)
            )
            total_fornecedores = float(resumo_fornecedores["valor"].sum())
            maior_fornecedor = resumo_fornecedores.iloc[0]
            ticket_medio = (
                total_fornecedores / int(resumo_fornecedores["lancamentos"].sum())
                if resumo_fornecedores["lancamentos"].sum() else 0.0
            )

            f1, f2, f3, f4 = st.columns(4)
            with f1:
                st.markdown(
                    f"""
                    <div class="kpi-card red">
                        <div class="kpi-label">Gasto de produção</div>
                        <div class="kpi-value red">{fmt_brl_saida(total_fornecedores)}</div>
                        <div class="kpi-footer">Itens diretamente ligados à produção</div>
                    </div>
                    """,
                    unsafe_allow_html=True,
                )
            with f2:
                st.markdown(
                    f"""
                        <div class="kpi-card blue">
                            <div class="kpi-label">Fornecedores</div>
                            <div class="kpi-value blue">{len(resumo_fornecedores)}</div>
                            <div class="kpi-footer">Com gasto de produção</div>
                        </div>
                    """,
                    unsafe_allow_html=True,
                )
            with f3:
                st.markdown(
                    f"""
                    <div class="kpi-card purple">
                        <div class="kpi-label">Maior gasto</div>
                        <div class="kpi-value red">{fmt_brl_saida(float(maior_fornecedor["valor"]))}</div>
                        <div class="kpi-footer">{texto_html(maior_fornecedor["fornecedor"], 42)}</div>
                    </div>
                    """,
                    unsafe_allow_html=True,
                )
            with f4:
                st.markdown(
                    f"""
                    <div class="kpi-card neutral">
                        <div class="kpi-label">Ticket médio</div>
                        <div class="kpi-value">{fmt_brl(ticket_medio)}</div>
                        <div class="kpi-footer">Média por lançamento</div>
                    </div>
                    """,
                    unsafe_allow_html=True,
                )

            st.markdown("<br>", unsafe_allow_html=True)
            st.markdown(
                """
                <div style="font-size:0.7rem;font-weight:700;letter-spacing:0.07em;
                            color:#4A4A7A;text-transform:uppercase;margin-bottom:0.8rem">
                    Principais fornecedores da produção
                </div>
                """,
                unsafe_allow_html=True,
            )
            cards_top = resumo_fornecedores.head(6).reset_index(drop=True)
            for inicio_cards in range(0, len(cards_top), 3):
                cols_cards = st.columns(3)
                for col_card, (_, fornecedor) in zip(
                    cols_cards,
                    cards_top.iloc[inicio_cards:inicio_cards + 3].iterrows(),
                ):
                    participacao = (
                        float(fornecedor["valor"]) / total_fornecedores * 100
                        if total_fornecedores > 0 else 0.0
                    )
                    with col_card:
                        st.markdown(
                            f"""
                            <div class="kpi-card neutral" style="margin-bottom:1rem">
                                <div class="kpi-label">{texto_html(fornecedor["fornecedor"], 48)}</div>
                                <div class="kpi-value red">{fmt_brl_saida(float(fornecedor["valor"]))}</div>
                                <div class="kpi-footer">
                                    {int(fornecedor["lancamentos"])} lançamento(s) · {participacao:.1f}% do gasto
                                    <br>{texto_html(fornecedor.get("categoria", "") or "Produção clínica", 64)}
                                </div>
                            </div>
                            """,
                            unsafe_allow_html=True,
                        )

            top_fornecedores = resumo_fornecedores.head(12).copy()
            fig_fornecedores = go.Figure(
                go.Bar(
                    x=top_fornecedores["valor"],
                    y=top_fornecedores["fornecedor"],
                    orientation="h",
                    marker_color="#F09595",
                    text=[fmt_brl(v) for v in top_fornecedores["valor"]],
                    textposition="auto",
                )
            )
            fig_fornecedores.update_layout(
                height=420,
                margin=dict(l=0, r=0, t=10, b=0),
                paper_bgcolor="#11112A",
                plot_bgcolor="#11112A",
                showlegend=False,
                xaxis=dict(showgrid=True, gridcolor="#1E1E3A", tickfont=dict(color="#8E91AD")),
                yaxis=dict(autorange="reversed", tickfont=dict(color="#CFD1DF")),
            )
            st.plotly_chart(
                fig_fornecedores,
                use_container_width=True,
                config={"displayModeBar": False},
            )

            busca_fornecedor = st.text_input(
                "Buscar fornecedor",
                placeholder="Digite parte do nome",
            )
            tabela_fornecedores = resumo_fornecedores.copy()
            if busca_fornecedor:
                tabela_fornecedores = tabela_fornecedores[
                    tabela_fornecedores["fornecedor"].str.contains(
                        busca_fornecedor,
                        case=False,
                        regex=False,
                        na=False,
                    )
                ]
            tabela_fornecedores["participacao"] = (
                tabela_fornecedores["valor"] / total_fornecedores * 100
            )
            tabela_fornecedores["ultima_data"] = pd.to_datetime(
                tabela_fornecedores["ultima_data"],
                errors="coerce",
            )
            st.dataframe(
                tabela_fornecedores.rename(
                    columns={
                        "fornecedor": "Fornecedor",
                        "valor": "Valor",
                        "lancamentos": "Lançamentos",
                        "ultima_data": "Última data",
                        "participacao": "% do gasto",
                    }
                ),
                use_container_width=True,
                hide_index=True,
                column_config={
                    "Valor": st.column_config.NumberColumn(format="R$ %.2f"),
                    "% do gasto": st.column_config.NumberColumn(format="%.1f%%"),
                    "Última data": st.column_config.DateColumn(format="DD/MM/YYYY"),
                },
            )

# =========================
# PÁGINA: RELATÓRIOS SALVOS
# =========================
elif st.session_state.pagina == "relatorios":
    st.markdown('<div class="page-title">Relatórios salvos</div>', unsafe_allow_html=True)
    st.markdown(
        '<div class="page-subtitle">Histórico interno por cliente, período e fechamento.</div>',
        unsafe_allow_html=True,
    )
    relatorios_salvos = list_saved_reports(user.id)
    if not relatorios_salvos:
        st.info("Nenhum relatório salvo ainda. Use o botão Salvar Relatório no menu lateral.")
    else:
        clientes = sorted({r[1] for r in relatorios_salvos})
        cliente_filtro = st.selectbox(
            "Buscar por cliente",
            ["Todas"] + clientes,
        )

        relatorios_por_cliente: Dict[str, Dict[str, List[dict]]] = {}
        for report_id, cliente, titulo, resumo_json, inicio, fim, criado_em in relatorios_salvos:
            if cliente_filtro != "Todas" and cliente != cliente_filtro:
                continue
            resumo = resumo_do_json(resumo_json)
            periodo_salvo = fmt_periodo_salvo(inicio, fim)
            mes_referencia = resumo.get("mes_referencia") or periodo_salvo
            relatorios_por_cliente.setdefault(cliente, {}).setdefault(
                mes_referencia,
                [],
            ).append({
                "id": report_id,
                "titulo": titulo,
                "resumo": resumo,
                "periodo": periodo_salvo,
                "criado_em": criado_em,
            })

        for cliente, meses in relatorios_por_cliente.items():
            total_cliente = sum(
                len(relatorios_mes)
                for relatorios_mes in meses.values()
            )
            with st.expander(
                f"📁 {cliente} · {total_cliente} relatório(s)",
                expanded=(cliente_filtro != "Todas"),
            ):
                for mes_referencia, relatorios_mes in meses.items():
                    with st.expander(
                        f"📂 {mes_referencia} · {len(relatorios_mes)} arquivo(s)",
                        expanded=True,
                    ):
                        for relatorio in relatorios_mes:
                            resumo = relatorio["resumo"]
                            st.markdown(
                                f"#### 📄 {texto_html(relatorio['titulo'])}"
                            )
                            st.caption(
                                f"Período: {relatorio['periodo']} · salvo em "
                                f"{relatorio['criado_em']}"
                            )
                            r1, r2, r3, r4, r5, r6 = st.columns([1, 1, 1, 1, .7, .7])
                            r1.metric(
                                "Recebimentos",
                                fmt_brl(resumo.get("recebimentos", 0.0)),
                            )
                            r2.metric(
                                "Despesas",
                                fmt_brl(resumo.get("despesas", 0.0)),
                            )
                            r3.metric(
                                "Vendas",
                                fmt_brl(resumo.get("vendas", 0.0)),
                            )
                            r4.metric(
                                "Resultado",
                                fmt_brl(
                                    resumo.get("resultado_final", 0.0),
                                    sinal=True,
                                ),
                            )
                            with r5:
                                if st.button(
                                    "Abrir",
                                    key=f"abrir_relatorio_{relatorio['id']}",
                                    use_container_width=True,
                                ):
                                    st.session_state.relatorio_salvo_para_abrir = relatorio["id"]
                                    st.rerun()
                            with r6:
                                if st.button(
                                    "Excluir",
                                    key=f"excluir_relatorio_{relatorio['id']}",
                                    use_container_width=True,
                                ):
                                    delete_saved_report(relatorio["id"], user.id)
                                    st.rerun()
                            st.markdown("---")

# =========================
# PÁGINA: COMPARATIVO
# =========================
elif st.session_state.pagina == "comparativo":
    st.markdown('<div class="page-title">Comparativo</div>', unsafe_allow_html=True)
    st.markdown(
        '<div class="page-subtitle">Compare recebimentos, despesas, resultado e vendas entre relatórios salvos.</div>',
        unsafe_allow_html=True,
    )
    relatorios_salvos = list_saved_reports(user.id)
    linhas_comparativo = []
    for report_id, cliente, titulo, resumo_json, inicio, fim, criado_em in relatorios_salvos:
        resumo = resumo_do_json(resumo_json)
        if not resumo:
            continue
        linhas_comparativo.append({
            "id": report_id,
            "cliente": cliente,
            "titulo": titulo,
            "periodo": resumo.get("mes_referencia") or fmt_periodo_salvo(inicio, fim),
            "inicio": pd.to_datetime(inicio, errors="coerce"),
            "fim": pd.to_datetime(fim, errors="coerce"),
            "recebimentos": resumo.get("recebimentos", 0.0),
            "despesas": resumo.get("despesas", 0.0),
            "resultado": resumo.get("resultado_final", 0.0),
            "vendas": resumo.get("vendas", 0.0),
        })
    df_comparativo = pd.DataFrame(linhas_comparativo)
    if df_comparativo.empty:
        st.info("Salve ao menos um relatório para montar o comparativo.")
    else:
        clientes = sorted(df_comparativo["cliente"].dropna().unique().tolist())
        if not clientes:
            st.info("Nenhum cliente encontrado nos relatórios salvos.")
            st.stop()
        cliente_comp = st.selectbox("Cliente", clientes)
        df_filtrado_comp = df_comparativo.copy()
        df_filtrado_comp = df_filtrado_comp[
            df_filtrado_comp["cliente"] == cliente_comp
        ]
        df_filtrado_comp = df_filtrado_comp.sort_values(["inicio", "fim", "titulo"])

        if df_filtrado_comp.empty:
            st.info("Nenhum relatório salvo para este cliente.")
        else:
            c1, c2, c3, c4 = st.columns(4)
            c1.metric("Recebimentos", fmt_brl(df_filtrado_comp["recebimentos"].sum()))
            c2.metric("Despesas", fmt_brl(df_filtrado_comp["despesas"].sum()))
            c3.metric("Vendas", fmt_brl(df_filtrado_comp["vendas"].sum()))
            c4.metric("Resultado", fmt_brl(df_filtrado_comp["resultado"].sum(), sinal=True))

            def grafico_comparativo(
                titulo: str,
                coluna: str,
                cor: str,
                cores_por_valor: bool = False,
            ) -> go.Figure:
                valores = df_filtrado_comp[coluna]
                cores = (
                    ["#2FC792" if valor >= 0 else "#F09595" for valor in valores]
                    if cores_por_valor else cor
                )
                fig = go.Figure()
                fig.add_trace(
                    go.Bar(
                        x=df_filtrado_comp["periodo"],
                        y=valores,
                        marker_color=cores,
                        text=[
                            fmt_brl(valor, sinal=(coluna == "resultado"))
                            for valor in valores
                        ],
                        textposition="auto",
                    )
                )
                fig.update_layout(
                    title=dict(
                        text=titulo,
                        font=dict(color="#E2E8F0", size=15),
                    ),
                    height=310,
                    margin=dict(l=0, r=0, t=42, b=0),
                    paper_bgcolor="#11112A",
                    plot_bgcolor="#11112A",
                    showlegend=False,
                    xaxis=dict(
                        showgrid=False,
                        tickfont=dict(color="#A0A0C0"),
                    ),
                    yaxis=dict(
                        showgrid=True,
                        gridcolor="#1E1E3A",
                        tickfont=dict(color="#4A4A7A"),
                    ),
                )
                return fig

            linha_graf_1 = st.columns(2)
            with linha_graf_1[0]:
                st.plotly_chart(
                    grafico_comparativo(
                        "Recebimentos",
                        "recebimentos",
                        "#2FC792",
                    ),
                    use_container_width=True,
                    config={"displayModeBar": False},
                )
            with linha_graf_1[1]:
                st.plotly_chart(
                    grafico_comparativo(
                        "Despesas",
                        "despesas",
                        "#F09595",
                    ),
                    use_container_width=True,
                    config={"displayModeBar": False},
                )
            linha_graf_2 = st.columns(2)
            with linha_graf_2[0]:
                st.plotly_chart(
                    grafico_comparativo(
                        "Resultado",
                        "resultado",
                        "#7167DC",
                        cores_por_valor=True,
                    ),
                    use_container_width=True,
                    config={"displayModeBar": False},
                )
            with linha_graf_2[1]:
                st.plotly_chart(
                    grafico_comparativo(
                        "Vendas",
                        "vendas",
                        "#68A9ED",
                    ),
                    use_container_width=True,
                    config={"displayModeBar": False},
                )
        st.dataframe(
            df_filtrado_comp[
                ["cliente", "periodo", "titulo", "recebimentos", "despesas", "vendas", "resultado"]
            ].rename(
                columns={
                    "cliente": "Cliente",
                    "periodo": "Período",
                    "titulo": "Relatório",
                    "recebimentos": "Recebimentos",
                    "despesas": "Despesas",
                    "vendas": "Vendas",
                    "resultado": "Resultado",
                }
            ),
            use_container_width=True,
            hide_index=True,
            column_config={
                "Recebimentos": st.column_config.NumberColumn(format="R$ %.2f"),
                "Despesas": st.column_config.NumberColumn(format="R$ %.2f"),
                "Vendas": st.column_config.NumberColumn(format="R$ %.2f"),
                "Resultado": st.column_config.NumberColumn(format="R$ %.2f"),
            },
        )

# =========================
# PÁGINA: DETALHES
# =========================
elif st.session_state.pagina == "detalhes":
    st.markdown('<div class="page-title">Detalhes</div>', unsafe_allow_html=True)
    subtitulo_detalhes = (
        "Visualização dos dados incluídos neste relatório."
        if st.session_state.share_mode
        else "Visualize os dados importados por relatório."
    )
    st.markdown(
        f'<div class="page-subtitle">{subtitulo_detalhes}</div>',
        unsafe_allow_html=True,
    )
    st.caption(f"Registros filtrados pelo período global: {periodo_label}.")

    tab_contas, tab_receb, tab_vendas, tab_infinity = st.tabs([
        "🔴 Contas Pagas",
        "🟢 Recebimentos",
        "🔵 Vendas",
        "💳 Maquininhas",
    ])

    with tab_contas:
        if df_exc is None or df_exc.empty:
            st.info("📥 Importe o arquivo Excel para ver as contas pagas.")
        else:
            col_busca, col_tipo, col_forma, _ = st.columns([2, 1.5, 1.5, 1])
            with col_busca:
                busca = st.text_input("", placeholder="🔍 Buscar...", label_visibility="collapsed")
            with col_tipo:
                filtro_tipo = st.selectbox(
                    "",
                    [
                        "Todos os tipos",
                        "Custo Fixo",
                        "Custo Variável",
                        "Retirada de Lucro",
                        "Pró-labore",
                        "Antecipação de Lucro",
                    ],
                    label_visibility="collapsed",
                )
            with col_forma:
                filtro_forma = st.selectbox("", ["Todas as formas","Transferência","Boleto","Dinheiro","Cartão"], label_visibility="collapsed")

            df_filtrado = pagamentos_periodo_global[
                "contas_conciliadas"
            ].copy()
            for coluna_padrao, valor_padrao in {
                "descricao": "",
                "categoria": "",
                "forma": "",
                "valor": 0.0,
            }.items():
                if coluna_padrao not in df_filtrado.columns:
                    df_filtrado[coluna_padrao] = valor_padrao
            df_filtrado["valor"] = pd.to_numeric(
                df_filtrado["valor"],
                errors="coerce",
            ).fillna(0.0)
            if "grupo_custo" not in df_filtrado.columns:
                df_filtrado["grupo_custo"] = classificar_grupo_custo(df_filtrado)
            if busca:
                mask = (
                    df_filtrado["descricao"].str.lower().str.contains(busca.lower(), regex=False, na=False) |
                    df_filtrado["categoria"].str.lower().str.contains(busca.lower(), regex=False, na=False)
                )
                df_filtrado = df_filtrado[mask]
            if filtro_tipo  != "Todos os tipos":
                df_filtrado = df_filtrado[df_filtrado["grupo_custo"] == filtro_tipo]
            if filtro_forma != "Todas as formas":
                df_filtrado = df_filtrado[df_filtrado["forma"].str.lower().str.contains(filtro_forma.lower(), regex=False, na=False)]

            totais_tipo = {
                "Custo Fixo": float(
                    df_filtrado.loc[
                        df_filtrado["grupo_custo"] == "Custo Fixo",
                        "valor",
                    ].sum()
                ),
                "Custo Variável": float(
                    df_filtrado.loc[
                        df_filtrado["grupo_custo"] == "Custo Variável",
                        "valor",
                    ].sum()
                ),
                "Retirada de Lucro": float(
                    df_filtrado.loc[
                        df_filtrado["grupo_custo"] == "Retirada de Lucro",
                        "valor",
                    ].sum()
                ),
                "Pró-labore": float(
                    df_filtrado.loc[
                        df_filtrado["grupo_custo"] == "Pró-labore",
                        "valor",
                    ].sum()
                ),
                "Antecipação de Lucro": float(
                    df_filtrado.loc[
                        df_filtrado["grupo_custo"] == "Antecipação de Lucro",
                        "valor",
                    ].sum()
                ),
            }
            cols_totais_contas = st.columns(len(totais_tipo))
            for col_total, (rotulo_total, valor_total) in zip(
                cols_totais_contas,
                totais_tipo.items(),
            ):
                with col_total:
                    st.markdown(
                        f"""
                        <div class="kpi-card neutral" style="margin-bottom:1rem">
                            <div class="kpi-label">{rotulo_total}</div>
                            <div class="kpi-value red">{fmt_brl_saida(valor_total)}</div>
                            <div class="kpi-footer">Total gasto no período filtrado</div>
                        </div>
                        """,
                        unsafe_allow_html=True,
                    )

            st.markdown(f'<div style="text-align:right;font-size:0.78rem;color:#4A4A7A;margin-bottom:0.5rem">{len(df_filtrado)} registro(s)</div>', unsafe_allow_html=True)

            grupos = {
                "Custo Fixo": (
                    "blue",
                    df_filtrado[df_filtrado["grupo_custo"] == "Custo Fixo"],
                ),
                "Custo Variável": (
                    "green",
                    df_filtrado[df_filtrado["grupo_custo"] == "Custo Variável"],
                ),
                "Retirada de Lucro": (
                    "orange",
                    df_filtrado[df_filtrado["grupo_custo"] == "Retirada de Lucro"],
                ),
                "Pró-labore": (
                    "orange",
                    df_filtrado[df_filtrado["grupo_custo"] == "Pró-labore"],
                ),
                "Antecipação de Lucro": (
                    "orange",
                    df_filtrado[
                        df_filtrado["grupo_custo"] == "Antecipação de Lucro"
                    ],
                ),
            }
            for nome_grupo, (cor, df_g) in grupos.items():
                total_g = df_g["valor"].sum()
                st.markdown(f"""
                <div class="group-header {cor}">
                    <span>▸ {nome_grupo} ({len(df_g)})</span>
                    <span style="font-family:'DM Mono',monospace">{fmt_brl_saida(total_g)}</span>
                </div>""", unsafe_allow_html=True)
                if df_g.empty:
                    st.markdown('<table class="fin-table"><tbody><tr class="empty-row"><td colspan="5">Nenhum registro</td></tr></tbody></table>', unsafe_allow_html=True)
                else:
                    rows = ""
                    for _, row in df_g.iterrows():
                        data_str = row["data"].strftime("%d/%m/%y") if pd.notna(row["data"]) else "—"
                        fornec = texto_html(row.get("descricao", ""), 55) or "—"
                        rows += f"<tr><td>{data_str}</td><td>{fornec}</td><td class='valor-neg'>{fmt_brl_saida(row['valor'])}</td><td>{tag_forma(row['forma'])}</td><td>{tag_categoria(row['categoria'])}</td></tr>"
                    st.markdown(f'<table class="fin-table"><thead><tr><th>Vencimento</th><th>Fornecedor</th><th>Valor</th><th>Forma</th><th>Categoria</th></tr></thead><tbody>{rows}</tbody></table>', unsafe_allow_html=True)

    with tab_receb:
        if usar_recebimentos_conta_azul:
            df_rec = df_conta_azul_receber_periodo.copy()
            if not df_creditos_diretos_conta_azul_periodo.empty:
                df_diretos = df_creditos_diretos_conta_azul_periodo.copy()
                df_diretos["descricao"] = df_diretos.get(
                    "memo",
                    pd.Series("Banco direto", index=df_diretos.index),
                )
                df_diretos["conta_destino"] = "Banco direto"
                df_diretos["tipo_recebimento"] = "credito bancario"
                df_rec = pd.concat(
                    [df_rec, df_diretos],
                    ignore_index=True,
                    sort=False,
                )

            if not df_rec.empty:
                df_rec["forma_recebimento"] = df_rec.apply(
                    classificar_forma_recebimento,
                    axis=1,
                )
            formas = (
                df_rec["forma_recebimento"].fillna("Outros")
                if "forma_recebimento" in df_rec.columns
                else pd.Series(dtype="object")
            )
            valores_rec = pd.to_numeric(
                df_rec.get("valor", pd.Series(0, index=df_rec.index)),
                errors="coerce",
            ).fillna(0)
            total_caixa = float(valores_rec[formas.isin(["Dinheiro"])].sum())
            total_maquininhas = float(
                valores_rec[formas.isin(["Maquininhas", "Cartão"])].sum()
            )
            total_banco = float(
                valores_rec[formas.isin(["Banco direto", "PIX", "Boleto"])].sum()
            )
            total_outros = float(
                valores_rec[
                    ~formas.isin([
                        "Dinheiro", "Maquininhas", "Cartão",
                        "Banco direto", "PIX", "Boleto",
                    ])
                ].sum()
            )

            st.markdown(
                f"""
                <div class="kpi-card green" style="margin-bottom:1rem">
                    <div class="kpi-label">Recebimentos</div>
                    <div class="kpi-value">{fmt_brl(recebimentos_conta_azul_periodo)}</div>
                    <div class="kpi-footer">
                        Fonte oficial: Contas a receber do Conta Azul
                    </div>
                </div>
                """,
                unsafe_allow_html=True,
            )
            rec_caixa, rec_maquininha, rec_banco, rec_outros = st.columns(4)
            cards_recebimentos = [
                (rec_caixa, "Caixa", total_caixa, "green", "Entradas em dinheiro/caixa"),
                (rec_maquininha, "Maquininhas", total_maquininhas, "blue", "Cartões identificados no Conta Azul"),
                (rec_banco, "Banco direto", total_banco, "green", "PIX, boleto e créditos bancários"),
                (rec_outros, "Outros", total_outros, "purple", "Formas não classificadas"),
            ]
            for coluna, titulo, valor, cor_card, rodape in cards_recebimentos:
                valor_classe = "blue" if cor_card == "blue" else "green"
                with coluna:
                    st.markdown(
                        f"""
                        <div class="kpi-card {cor_card}" style="margin-bottom:1rem">
                            <div class="kpi-label">{texto_html(titulo)}</div>
                            <div class="kpi-value {valor_classe}">{fmt_brl(valor)}</div>
                            <div class="kpi-footer">{texto_html(rodape)}</div>
                        </div>
                        """,
                        unsafe_allow_html=True,
                    )

            formas_disponiveis = sorted(
                df_rec.get(
                    "forma_recebimento",
                    pd.Series(dtype="object"),
                ).dropna().unique().tolist()
            )
            filtro_receb_forma = st.selectbox(
                "Forma de recebimento",
                ["Todas"] + formas_disponiveis,
                key="filtro_forma_recebimentos_conta_azul",
            )
            if filtro_receb_forma != "Todas" and "forma_recebimento" in df_rec.columns:
                df_rec = df_rec[df_rec["forma_recebimento"] == filtro_receb_forma]

            st.markdown(
                f'<div style="text-align:right;font-size:0.78rem;color:#4A4A7A;margin-bottom:0.5rem">{len(df_rec)} registro(s)</div>',
                unsafe_allow_html=True,
            )
            rows = ""
            for _, row in df_rec.sort_values(
                "data",
                ascending=False,
                na_position="last",
            ).iterrows():
                data_str = row["data"].strftime("%d/%m/%y") if pd.notna(row["data"]) else "—"
                memo = texto_html(
                    row.get("descricao", "") or row.get("memo", ""),
                    70,
                )
                conta = tag_categoria(row.get("conta_destino", "Conta Azul"))
                forma = tag_categoria(row.get("forma_recebimento", "Outros"))
                rows += (
                    f"<tr><td>{data_str}</td><td>{memo}</td><td>{forma}</td>"
                    f"<td>{conta}</td><td class='valor-pos'>{fmt_brl(row.get('valor', 0))}</td></tr>"
                )
            st.markdown(
                '<table class="fin-table"><thead><tr><th>Data</th><th>Descrição</th>'
                '<th>Forma</th><th>Conta</th><th>Valor</th></tr></thead>'
                f'<tbody>{rows}</tbody></table>',
                unsafe_allow_html=True,
            )
        elif (
            usar_conta_corrente_clinicorp
            or usar_recebimentos_oficiais_belle
            or recebimentos_belle_periodo > 0
        ):
            if usar_conta_corrente_clinicorp:
                df_rec = pd.DataFrame()
            elif recebimentos_infinity_extrato_periodo > 0:
                df_infinity_rec = df_entradas_infinity_periodo.copy()
                df_infinity_rec["descricao"] = df_infinity_rec["memo"]
                df_rec = pd.concat(
                    [
                        df_recebimentos_caixa_belle_periodo,
                        df_infinity_rec,
                    ],
                    ignore_index=True,
                    sort=False,
                )
            else:
                df_rec = df_recebimentos_belle_periodo.copy()
            conciliacao_banco_detalhes = conciliar_recebimentos(
                base_recebimentos_ofx_periodo,
                df_banco_conciliacao_periodo,
                df_clinipay_base_periodo,
                df_fluxo_base_periodo,
            )
            total_recebimentos_detalhes = (
                recebimentos_conta_corrente_periodo
                if usar_conta_corrente_clinicorp
                else recebimentos_oficiais_belle_periodo
                if usar_recebimentos_oficiais_belle
                else recebimentos_belle_periodo + conciliacao_banco_detalhes["total"]
            )
            st.markdown(
                f"""
                <div class="kpi-card green" style="margin-bottom:1rem">
                    <div class="kpi-label">Recebimentos</div>
                    <div class="kpi-value">{fmt_brl(total_recebimentos_detalhes)}</div>
                    <div class="kpi-footer">
                        {"Conta corrente oficial · cards abaixo compõem este total" if usar_conta_corrente_clinicorp else "Contas a receber do Belle" if usar_recebimentos_oficiais_belle else "Caixa, maquininhas e banco direto"}
                    </div>
                </div>
                """,
                unsafe_allow_html=True,
            )
            valor_caixa_detalhes = recebimentos_caixa_belle_periodo
            valor_maquininhas_detalhes = (
                recebimentos_infinity_extrato_periodo
                if recebimentos_infinity_extrato_periodo > 0
                else recebimentos_infinity_sistema_periodo
            )
            valor_banco_direto_detalhes = conciliacao_banco_detalhes["total"]
            if usar_conta_corrente_clinicorp and not df_recebimentos_conta_corrente_periodo.empty:
                conta_corrente_rec = df_recebimentos_conta_corrente_periodo.copy()
                texto_conta_corrente = (
                    conta_corrente_rec.get("memo", pd.Series("", index=conta_corrente_rec.index)).fillna("").map(normalizar_texto)
                    + " "
                    + conta_corrente_rec.get("nome", pd.Series("", index=conta_corrente_rec.index)).fillna("").map(normalizar_texto)
                    + " "
                    + conta_corrente_rec.get("tipo_lancamento", pd.Series("", index=conta_corrente_rec.index)).fillna("").map(normalizar_texto)
                    + " "
                    + conta_corrente_rec.get("tipo_transacao", pd.Series("", index=conta_corrente_rec.index)).fillna("").map(normalizar_texto)
                )
                valores_conta_corrente = pd.to_numeric(
                    conta_corrente_rec.get("valor", 0),
                    errors="coerce",
                ).fillna(0)
                mascara_caixa_cc = texto_conta_corrente.str.contains(
                    r"\bcaixa\b|dinheiro|especie|espécie",
                    regex=True,
                )
                mascara_maquininha_cc = texto_conta_corrente.str.contains(
                    r"cartao|cartão|credito|crédito|debito|débito|visa|master|elo|rede",
                    regex=True,
                )
                valor_caixa_detalhes = float(valores_conta_corrente[mascara_caixa_cc].sum())
                valor_maquininhas_detalhes = float(valores_conta_corrente[mascara_maquininha_cc].sum())
                valor_banco_direto_detalhes = float(
                    valores_conta_corrente[
                        ~(mascara_caixa_cc | mascara_maquininha_cc)
                    ].sum()
                )
            valor_clinicorp_detalhes = (
                max(
                    total_recebimentos_detalhes
                    - valor_caixa_detalhes
                    - valor_maquininhas_detalhes
                    - valor_banco_direto_detalhes,
                    0.0,
                )
                if usar_conta_corrente_clinicorp else 0.0
            )
            cards_recebimentos = [
                (
                    rotulo_caixa_recebimentos,
                    valor_caixa_detalhes,
                    "green",
                    "Entradas em dinheiro/caixa",
                ),
                (
                    "Maquininhas",
                    valor_maquininhas_detalhes,
                    "blue",
                    "Cartões dentro da conta corrente" if usar_conta_corrente_clinicorp else "Extratos de cartão conciliados",
                ),
                (
                    "Banco direto",
                    valor_banco_direto_detalhes,
                    "green",
                    "Pix, boletos e entradas diretas" if usar_conta_corrente_clinicorp else "Créditos localizados no OFX",
                ),
            ]
            if usar_conta_corrente_clinicorp:
                cards_recebimentos.append(
                    (
                        "Clinicorp",
                        valor_clinicorp_detalhes,
                        "purple",
                        "Demais entradas na conta corrente",
                    )
                )
            cards_recebimentos = [
                card for card in cards_recebimentos
                if float(card[1] or 0) > 0 or card[0] in {rotulo_caixa_recebimentos, "Banco direto"}
            ]
            cols_recebimentos = st.columns(min(4, max(1, len(cards_recebimentos))))
            for idx_card, (titulo, valor, cor_card, rodape) in enumerate(cards_recebimentos):
                valor_classe = "blue" if cor_card == "blue" else "green"
                coluna = cols_recebimentos[idx_card % len(cols_recebimentos)]
                with coluna:
                    st.markdown(
                        f"""
                        <div class="kpi-card {cor_card}" style="margin-bottom:1rem">
                            <div class="kpi-label">{texto_html(titulo)}</div>
                            <div class="kpi-value {valor_classe}">{fmt_brl(valor)}</div>
                            <div class="kpi-footer">{texto_html(rodape)}</div>
                        </div>
                        """,
                        unsafe_allow_html=True,
                    )
            df_banco_rec = (
                df_recebimentos_conta_corrente_periodo.copy()
                if usar_conta_corrente_clinicorp
                else conciliacao_banco_detalhes["creditos_conciliados"].copy()
            )
            if usar_conta_corrente_clinicorp and not df_banco_rec.empty:
                df_banco_rec["tipo_recebimento"] = "Conta corrente"
            if not df_rec.empty:
                df_rec["forma_recebimento"] = df_rec.apply(
                    classificar_forma_recebimento,
                    axis=1,
                )
            if not df_banco_rec.empty:
                df_banco_rec["forma_recebimento"] = df_banco_rec.apply(
                    classificar_forma_recebimento,
                    axis=1,
                )
            formas_disponiveis = sorted(
                set(df_rec.get("forma_recebimento", pd.Series(dtype="object")).dropna())
                | set(df_banco_rec.get("forma_recebimento", pd.Series(dtype="object")).dropna())
            )
            filtro_receb_forma = st.selectbox(
                "Forma de recebimento",
                ["Todas"] + formas_disponiveis,
                key="filtro_forma_recebimentos_detalhes",
            )
            if filtro_receb_forma != "Todas":
                if "forma_recebimento" in df_rec.columns:
                    df_rec = df_rec[df_rec["forma_recebimento"] == filtro_receb_forma]
                if "forma_recebimento" in df_banco_rec.columns:
                    df_banco_rec = df_banco_rec[
                        df_banco_rec["forma_recebimento"] == filtro_receb_forma
                    ]

            total_filtrado_receb = len(df_rec) + len(df_banco_rec)
            st.markdown(f'<div style="text-align:right;font-size:0.78rem;color:#4A4A7A;margin-bottom:0.5rem">{total_filtrado_receb} registro(s)</div>', unsafe_allow_html=True)
            rows = ""
            if not df_rec.empty and "data" in df_rec.columns:
                for _, row in df_rec.sort_values(
                    "data", ascending=False, na_position="last"
                ).iterrows():
                    data_str = row["data"].strftime("%d/%m/%y") if pd.notna(row["data"]) else "—"
                    memo = texto_html(row.get("descricao", ""), 70)
                    conta = tag_categoria(row.get("conta_destino", ""))
                    forma = tag_categoria(row.get("forma_recebimento", "Outros"))
                    valor_recebido = row.get("valor_recebido", row["valor"])
                    if pd.isna(valor_recebido):
                        valor_recebido = row["valor"]
                    rows += f"<tr><td>{data_str}</td><td>{memo}</td><td>{forma}</td><td>{conta}</td><td class='valor-pos'>{fmt_brl(valor_recebido)}</td></tr>"
                st.markdown(f'<table class="fin-table"><thead><tr><th>Data</th><th>Descrição</th><th>Forma</th><th>Conta</th><th>Valor</th></tr></thead><tbody>{rows}</tbody></table>', unsafe_allow_html=True)
            if not df_banco_rec.empty:
                rows = ""
                for _, row in df_banco_rec.iterrows():
                    data_str = row["data"].strftime("%d/%m/%y") if pd.notna(row["data"]) else "—"
                    memo = texto_html(row.get("memo", ""), 70)
                    forma = tag_categoria(row.get("forma_recebimento", "Banco direto"))
                    rows += f"<tr><td>{data_str}</td><td>{memo}</td><td>{forma}</td><td class='valor-pos'>{fmt_brl(row['valor'])}</td></tr>"
                st.markdown(f'<table class="fin-table"><thead><tr><th>Data</th><th>Recebimento bancário</th><th>Forma</th><th>Valor</th></tr></thead><tbody>{rows}</tbody></table>', unsafe_allow_html=True)
        elif (
            (df_banco_base_periodo.empty)
            and (df_clinipay_base_periodo.empty)
            and (df_fluxo_base_periodo.empty)
            or base_recebimentos_belle.empty
        ):
            st.info(
                "📥 Importe a planilha de orçamentos e ao menos um extrato "
                "(bancário, Clinipay ou Fluxo de Caixa) para conciliar "
                "os recebimentos."
            )
        else:
            conciliacao_detalhes = conciliar_recebimentos(
                base_recebimentos_ofx_periodo,
                df_banco_conciliacao_periodo,
                df_clinipay_base_periodo,
                df_fluxo_base_periodo,
            )
            df_rec = (
                df_recebimentos_conta_corrente_periodo.copy()
                if usar_conta_corrente_clinicorp
                else conciliacao_detalhes["creditos_conciliados"].copy()
            )
            if usar_conta_corrente_clinicorp and not df_rec.empty:
                df_rec["tipo_recebimento"] = "Conta corrente"
            if df_rec.empty:
                st.info("Nenhum recebimento conciliado com as vendas aprovadas.")
            else:
                st.markdown(
                    f"""
                    <div class="kpi-card green" style="margin-bottom:1rem">
                        <div class="kpi-label">{"Recebimentos da conta corrente" if usar_conta_corrente_clinicorp else "Recebimentos conciliados"}</div>
                        <div class="kpi-value">{fmt_brl(recebimentos_conta_corrente_periodo if usar_conta_corrente_clinicorp else conciliacao_detalhes["total"])}</div>
                        <div class="kpi-footer">
                            {"Fonte oficial do Clinicorp" if usar_conta_corrente_clinicorp else "Recebimentos identificados das vendas aprovadas"}
                        </div>
                    </div>
                    """,
                    unsafe_allow_html=True,
                )
                df_rec["forma_recebimento"] = df_rec.apply(
                    classificar_forma_recebimento,
                    axis=1,
                )
                formas_disponiveis = sorted(
                    df_rec["forma_recebimento"].dropna().unique().tolist()
                )
                filtro_receb_forma = st.selectbox(
                    "Forma de recebimento",
                    ["Todas"] + formas_disponiveis,
                    key="filtro_forma_recebimentos_conciliados",
                )
                if filtro_receb_forma != "Todas":
                    df_rec = df_rec[
                        df_rec["forma_recebimento"] == filtro_receb_forma
                    ]
                st.markdown(f'<div style="text-align:right;font-size:0.78rem;color:#4A4A7A;margin-bottom:0.5rem">{len(df_rec)} registro(s)</div>', unsafe_allow_html=True)
                if df_rec.empty:
                    st.info("Nenhum recebimento nesta forma selecionada.")
                else:
                    items_per_page = 50
                    total_pages    = (len(df_rec) - 1) // items_per_page + 1
                    page           = st.number_input("Página", min_value=1, max_value=total_pages, value=1)
                    df_page        = df_rec.iloc[(page-1)*items_per_page : page*items_per_page]
                    rows = ""
                    for _, row in df_page.iterrows():
                        data_str = row["data"].strftime("%d/%m/%y") if pd.notna(row["data"]) else "—"
                        memo = texto_html(row.get("memo", ""), 70)
                        forma = tag_categoria(row.get("forma_recebimento", "Outros"))
                        rows += f"<tr><td>{data_str}</td><td>{memo}</td><td>{forma}</td><td class='valor-pos'>{fmt_brl(row['valor'])}</td></tr>"
                    st.markdown(f'<table class="fin-table"><thead><tr><th>Data</th><th>Descrição</th><th>Forma</th><th>Valor</th></tr></thead><tbody>{rows}</tbody></table>', unsafe_allow_html=True)

    with tab_vendas:
        if tem_gerencial_belle:
            st.markdown(
                f"""
                <div class="kpi-card blue" style="margin-bottom:1rem">
                    <div class="kpi-label">Total de vendas feitas</div>
                    <div class="kpi-value blue">{fmt_brl(vendas_referencia_periodo)}</div>
                    <div class="kpi-footer">
                        Fonte oficial: Venda de Planos no Gerencial Belle
                    </div>
                </div>
                """,
                unsafe_allow_html=True,
            )
            col_vg, col_cr, col_dif = st.columns(3)
            col_vg.metric("Gerencial", fmt_brl(vendas_referencia_periodo))
            col_cr.metric("Contas a receber", fmt_brl(vendas_contas_receber_referencia))
            col_dif.metric(
                "Diferença",
                fmt_brl(diferenca_vendas_gerencial_receber, sinal=True),
            )
            rows = (
                f"<tr><td>{inicio_periodo.strftime('%d/%m/%Y')}</td>"
                f"<td>Venda de Planos</td>"
                f"<td class='valor-pos'>{fmt_brl(vendas_referencia_periodo)}</td>"
                f"<td>{tag_categoria('Gerencial')}</td>"
                f"<td>Gerencial de resultados Belle</td></tr>"
            )
            st.markdown(
                '<table class="fin-table"><thead><tr>'
                '<th>Data</th><th>Descrição</th><th>Valor</th>'
                '<th>Status</th><th>Origem</th></tr></thead>'
                f'<tbody>{rows}</tbody></table>',
                unsafe_allow_html=True,
            )
        elif df_orcamentos is None or df_orcamentos.empty:
            st.info(
                "📥 Importe a planilha de orçamentos para visualizar as vendas aprovadas."
            )
        else:
            vendas_detalhes = (
                float(base_recebimentos_belle["valor"].sum())
                if not base_recebimentos_belle.empty else 0.0
            )
            st.markdown(
                f"""
                <div class="kpi-card blue" style="margin-bottom:1rem">
                    <div class="kpi-label">Total de vendas aprovadas</div>
                    <div class="kpi-value blue">{fmt_brl(vendas_detalhes)}</div>
                    <div class="kpi-footer">
                        {len(base_recebimentos_belle)} orçamento(s) na base selecionada
                    </div>
                </div>
                """,
                unsafe_allow_html=True,
            )
            if (
                "servico_vendido" in base_recebimentos_belle.columns
                and normalizar_texto(fonte_vendas_ativa)
                in {"orcamentos do clinicorp", "vendas do conta azul"}
            ):
                fonte_vendas_norm = normalizar_texto(fonte_vendas_ativa)
                vendas_por_servico = base_recebimentos_belle.copy()
                vendas_por_servico["servico_vendido"] = (
                    vendas_por_servico["servico_vendido"]
                    .fillna("")
                    .astype(str)
                    .str.strip()
                    .replace("", "Serviço não identificado")
                )
                tipo_item_servico = vendas_por_servico.get(
                    "tipo_item_vendido",
                    pd.Series("", index=vendas_por_servico.index),
                ).fillna("").astype(str)
                if fonte_vendas_norm == "vendas do conta azul":
                    vendas_por_servico["servico_vendido"] = [
                        agrupar_servico_vendido(servico, tipo_item)
                        for servico, tipo_item in zip(
                            vendas_por_servico["servico_vendido"],
                            tipo_item_servico,
                        )
                    ]
                else:
                    vendas_por_servico["servico_vendido"] = (
                        vendas_por_servico["servico_vendido"]
                        .apply(agrupar_procedimento_clinicorp)
                    )
                coluna_valor_servico = (
                    "valor_liquido"
                    if "valor_liquido" in vendas_por_servico.columns
                    else "valor"
                )
                vendas_por_servico["valor_card_servico"] = pd.to_numeric(
                    vendas_por_servico[coluna_valor_servico],
                    errors="coerce",
                ).fillna(0).abs()
                resumo_servicos = (
                    vendas_por_servico.groupby("servico_vendido", as_index=False)
                    .agg(
                        valor=("valor_card_servico", "sum"),
                        quantidade=("valor_card_servico", "count"),
                    )
                    .sort_values("valor", ascending=False)
                )
                if not resumo_servicos.empty:
                    st.markdown(
                        """
                        <div style="font-size:0.7rem;font-weight:700;letter-spacing:0.07em;
                                    color:#4A4A7A;text-transform:uppercase;margin:0.3rem 0 0.8rem">
                            Serviços vendidos
                        </div>
                        """,
                        unsafe_allow_html=True,
                    )
                    if fonte_vendas_norm == "vendas do conta azul":
                        categorias_servico = [
                            "PRODUTO",
                            "CONSULTA",
                            "IMPLANTES",
                            "INJETÁVEIS",
                            "PROGRAMA",
                            "TIRZEPATIDA",
                        ]
                        servicos_por_nome = {
                            str(row["servico_vendido"]): row
                            for _, row in resumo_servicos.iterrows()
                        }
                        cards_servicos = []
                        for categoria_servico in categorias_servico:
                            linha_servico = servicos_por_nome.get(categoria_servico)
                            cards_servicos.append({
                                "servico_vendido": categoria_servico,
                                "valor": (
                                    float(linha_servico["valor"])
                                    if linha_servico is not None else 0.0
                                ),
                                "quantidade": (
                                    int(linha_servico["quantidade"])
                                    if linha_servico is not None else 0
                                ),
                            })
                        cards_servicos.append({
                            "servico_vendido": "TOTAL VENDAS",
                            "valor": float(resumo_servicos["valor"].sum()),
                            "quantidade": int(resumo_servicos["quantidade"].sum()),
                        })
                    else:
                        cards_servicos = resumo_servicos.to_dict("records")
                    for inicio_servicos in range(0, len(cards_servicos), 4):
                        cols_servicos = st.columns(4)
                        for col_servico, servico in zip(
                            cols_servicos,
                            cards_servicos[inicio_servicos:inicio_servicos + 4],
                        ):
                            participacao = (
                                float(servico["valor"]) / vendas_detalhes * 100
                                if vendas_detalhes > 0 else 0.0
                            )
                            with col_servico:
                                st.markdown(
                                    f"""
                                    <div class="kpi-card blue" style="margin-bottom:1rem">
                                        <div class="kpi-label">{texto_html(servico["servico_vendido"], 52)}</div>
                                        <div class="kpi-value blue">{fmt_brl(float(servico["valor"]))}</div>
                                        <div class="kpi-footer">
                                            {int(servico["quantidade"])} venda(s) · {participacao:.1f}% do total
                                        </div>
                                    </div>
                                    """,
                                    unsafe_allow_html=True,
                                )
            vendas_tabela = base_recebimentos_belle.copy()
            for coluna_padrao, valor_padrao in {
                "data": pd.NaT,
                "descricao": "",
                "valor": 0.0,
                "status": "Aprovado",
                "fonte_venda": "Venda",
            }.items():
                if coluna_padrao not in vendas_tabela.columns:
                    vendas_tabela[coluna_padrao] = valor_padrao
            vendas_tabela["data"] = pd.to_datetime(
                vendas_tabela["data"],
                errors="coerce",
            )
            vendas_tabela["valor"] = pd.to_numeric(
                vendas_tabela["valor"],
                errors="coerce",
            ).fillna(0.0)
            if (
                "servico_vendido" in vendas_tabela.columns
                and normalizar_texto(fonte_vendas_ativa)
                in {"orcamentos do clinicorp", "vendas do conta azul"}
            ):
                fonte_vendas_norm = normalizar_texto(fonte_vendas_ativa)
                servicos_tabela = (
                    vendas_tabela["servico_vendido"]
                    .fillna("")
                    .astype(str)
                    .str.strip()
                )
                tipos_tabela = vendas_tabela.get(
                    "tipo_item_vendido",
                    pd.Series("", index=vendas_tabela.index),
                ).fillna("").astype(str)
                if fonte_vendas_norm == "vendas do conta azul":
                    vendas_tabela["tipo_servico_card"] = [
                        agrupar_servico_vendido(servico, tipo_item)
                        for servico, tipo_item in zip(servicos_tabela, tipos_tabela)
                    ]
                    opcoes_tipo_servico = [
                        "Todos",
                        "PRODUTO",
                        "CONSULTA",
                        "IMPLANTES",
                        "INJETÁVEIS",
                        "PROGRAMA",
                        "TIRZEPATIDA",
                        "Outros",
                    ]
                else:
                    vendas_tabela["tipo_servico_card"] = (
                        servicos_tabela.apply(agrupar_procedimento_clinicorp)
                    )
                    opcoes_tipo_servico = [
                        "Todos",
                        *sorted(
                            vendas_tabela["tipo_servico_card"]
                            .dropna()
                            .astype(str)
                            .unique()
                            .tolist()
                        ),
                    ]
                filtro_tipo_servico = st.selectbox(
                    "Tipo de serviço",
                    opcoes_tipo_servico,
                    key="filtro_tipo_servico_vendas",
                )
                if filtro_tipo_servico != "Todos":
                    vendas_tabela = vendas_tabela[
                        vendas_tabela["tipo_servico_card"] == filtro_tipo_servico
                    ].copy()
                st.markdown(
                    f'<div style="text-align:right;font-size:0.78rem;color:#4A4A7A;margin-bottom:0.5rem">{len(vendas_tabela)} venda(s)</div>',
                    unsafe_allow_html=True,
                )
            rows = ""
            for _, row in vendas_tabela.sort_values(
                "data", ascending=False, na_position="last"
            ).iterrows():
                data_str = (
                    row["data"].strftime("%d/%m/%Y")
                    if pd.notna(row["data"]) else "—"
                )
                descricao = (
                    texto_html(row["descricao"], 80)
                    or "Orçamento aprovado"
                )
                fonte_venda = html.escape(
                    str(row.get("fonte_venda", "Venda"))
                )
                rows += (
                    f"<tr><td>{data_str}</td><td>{descricao}</td>"
                    f"<td class='valor-pos'>{fmt_brl(row['valor'])}</td>"
                    f"<td>{tag_categoria(row['status'])}</td>"
                    f"<td>{fonte_venda}</td></tr>"
                )
            st.markdown(
                '<table class="fin-table"><thead><tr>'
                '<th>Data</th><th>Cliente / Orçamento</th><th>Valor</th>'
                '<th>Status</th><th>Origem</th></tr></thead>'
                f'<tbody>{rows}</tbody></table>',
                unsafe_allow_html=True,
            )

    with tab_infinity:
        if (
            (df_infinity_base_periodo is None or df_infinity_base_periodo.empty)
            and (df_clinipay_base_periodo is None or df_clinipay_base_periodo.empty)
        ):
            st.info(
                "📥 Importe o extrato da maquininha ou o relatório Clinipay "
                "para validar entradas, taxas e boletos."
            )
        else:
            if df_infinity_base_periodo is None or df_infinity_base_periodo.empty:
                df_maquininhas_view = pd.DataFrame(
                    columns=[
                        "data", "valor", "memo", "hora", "maquininha_operadora",
                        "ponto_venda", "fluxo_infinity", "bandeira",
                        "tipo_transacao", "parcelas", "taxa_maquininha",
                    ]
                )
            else:
                df_maquininhas_view = df_infinity_base_periodo.copy()
            col_op, col_band, col_tipo_maq, col_parc = st.columns(4)
            with col_op:
                opcoes_operadora = ["Todas"] + sorted(
                    [
                        valor for valor in
                        df_maquininhas_view["maquininha_operadora"]
                        .fillna("Maquininha")
                        .astype(str)
                        .unique()
                        if valor
                    ]
                )
                filtro_operadora = st.selectbox(
                    "Maquininha",
                    opcoes_operadora,
                    key="filtro_maquininha_operadora",
                )
            with col_band:
                opcoes_bandeira = ["Todas"] + sorted(
                    [
                        valor for valor in
                        df_maquininhas_view["bandeira"]
                        .fillna("")
                        .astype(str)
                        .unique()
                        if valor and valor.lower() != "nan"
                    ]
                )
                filtro_bandeira = st.selectbox(
                    "Bandeira",
                    opcoes_bandeira,
                    key="filtro_maquininha_bandeira",
                )
            with col_tipo_maq:
                opcoes_tipo = ["Todos"] + sorted(
                    [
                        valor for valor in
                        df_maquininhas_view["tipo_transacao"]
                        .fillna("")
                        .astype(str)
                        .unique()
                        if valor and valor.lower() != "nan"
                    ]
                )
                filtro_tipo_maq = st.selectbox(
                    "Tipo",
                    opcoes_tipo,
                    key="filtro_maquininha_tipo",
                )
            with col_parc:
                opcoes_parcelas = ["Todas"] + sorted(
                    [
                        valor for valor in
                        df_maquininhas_view["parcelas"]
                        .fillna("")
                        .astype(str)
                        .unique()
                        if valor and valor.lower() != "nan"
                    ]
                )
                filtro_parcelas = st.selectbox(
                    "Parcelas",
                    opcoes_parcelas,
                    key="filtro_maquininha_parcelas",
                )

            if filtro_operadora != "Todas":
                df_maquininhas_view = df_maquininhas_view[
                    df_maquininhas_view["maquininha_operadora"].fillna("") == filtro_operadora
                ]
            if filtro_bandeira != "Todas":
                df_maquininhas_view = df_maquininhas_view[
                    df_maquininhas_view["bandeira"].fillna("") == filtro_bandeira
                ]
            if filtro_tipo_maq != "Todos":
                df_maquininhas_view = df_maquininhas_view[
                    df_maquininhas_view["tipo_transacao"].fillna("") == filtro_tipo_maq
                ]
            if filtro_parcelas != "Todas":
                df_maquininhas_view = df_maquininhas_view[
                    df_maquininhas_view["parcelas"].fillna("") == filtro_parcelas
                ]

            taxas_maquininha_view = calcular_taxas_maquininha(df_maquininhas_view)
            entradas_view = df_maquininhas_view[
                pd.to_numeric(df_maquininhas_view["valor"], errors="coerce").fillna(0) > 0
            ].copy()
            recebimentos_maquininha_view = float(
                pd.to_numeric(entradas_view.get("valor", 0), errors="coerce")
                .fillna(0)
                .sum()
            )
            entradas_maquininha_total = (
                df_infinity_base_periodo[
                    pd.to_numeric(
                        df_infinity_base_periodo.get("valor", 0),
                        errors="coerce",
                    ).fillna(0) > 0
                ].copy()
                if df_infinity_base_periodo is not None
                and not df_infinity_base_periodo.empty
                else pd.DataFrame()
            )
            total_entradas_maquininha = float(
                pd.to_numeric(
                    entradas_maquininha_total.get(
                        "valor",
                        pd.Series(0, index=entradas_maquininha_total.index),
                    ),
                    errors="coerce",
                ).fillna(0).sum()
            )
            saidas_maquininha_view = float(
                pd.to_numeric(
                    df_maquininhas_view.loc[
                        pd.to_numeric(
                            df_maquininhas_view["valor"],
                            errors="coerce",
                        ).fillna(0) < 0,
                        "valor",
                    ],
                    errors="coerce",
                ).fillna(0).abs().sum()
            )
            diferenca_entradas = (
                total_entradas_maquininha
                - recebimentos_maquininha_view
            )
            diferenca_saidas = (
                creditos_banco_infinity_periodo - saidas_maquininha_view
            )
            percentual_taxa_maquininha = taxas_maquininha_view["taxa_media"]
            percentual_recebido_maquininha = (
                100 - percentual_taxa_maquininha
                if taxas_maquininha_view["bruto"] > 0 else 0.0
            )
            c1, c2, c3, c4 = st.columns(4)
            delta_entradas_cls = (
                "positive" if diferenca_entradas > 0
                else "negative" if diferenca_entradas < 0
                else "neutral"
            )
            delta_saidas_cls = (
                "positive" if diferenca_saidas > 0
                else "negative" if diferenca_saidas < 0
                else "neutral"
            )
            with c1:
                st.markdown(
                    f"""
                        <div class="kpi-card green" style="margin-bottom:1rem">
                            <div class="kpi-label">Entradas maquininhas</div>
                            <div class="kpi-value green">{fmt_brl(total_entradas_maquininha)}</div>
                            <div class="kpi-footer">Total do extrato importado</div>
                        </div>
                    """,
                    unsafe_allow_html=True,
                )
            with c2:
                st.markdown(
                    f"""
                        <div class="kpi-card blue" style="margin-bottom:1rem">
                            <div class="kpi-label">Extrato filtrado</div>
                            <div class="kpi-value blue">{fmt_brl(recebimentos_maquininha_view)}</div>
                            <div class="kpi-footer">Recorte atual dos filtros</div>
                            <div class="machine-delta {delta_entradas_cls}">{fmt_brl(diferenca_entradas, sinal=True)}</div>
                        </div>
                    """,
                    unsafe_allow_html=True,
                )
            with c3:
                st.markdown(
                    f"""
                        <div class="kpi-card red" style="margin-bottom:1rem">
                            <div class="kpi-label">Taxas maquininhas</div>
                            <div class="kpi-value red">{fmt_brl_saida(taxas_maquininha_view["taxas"])}</div>
                            <div class="kpi-footer">Paga {percentual_taxa_maquininha:.2f}% · recebe {percentual_recebido_maquininha:.2f}% do bruto</div>
                        </div>
                    """,
                    unsafe_allow_html=True,
                )
            with c4:
                st.markdown(
                    f"""
                        <div class="kpi-card neutral" style="margin-bottom:1rem">
                            <div class="kpi-label">Saídas maquininhas</div>
                            <div class="kpi-value red">{fmt_brl_saida(saidas_maquininha_view)}</div>
                            <div class="kpi-footer">Repasses e saídas identificadas</div>
                        </div>
                    """,
                    unsafe_allow_html=True,
                )
            resumo_operadoras = []
            for operadora, grupo in df_maquininhas_view.groupby(
                df_maquininhas_view["maquininha_operadora"].fillna("Maquininha")
            ):
                resumo = calcular_taxas_maquininha(grupo)
                resumo_operadoras.append({
                    "operadora": operadora,
                    "bruto": resumo["bruto"],
                    "taxas": resumo["taxas"],
                    "liquido": resumo["liquido"],
                    "taxa_media": resumo["taxa_media"],
                    "lancamentos": len(grupo),
                })
            if resumo_operadoras:
                cards_operadoras = sorted(
                    resumo_operadoras,
                    key=lambda item: item["bruto"],
                    reverse=True,
                )[:4]
                cols_operadoras = st.columns(len(cards_operadoras))
                for col_operadora, item in zip(cols_operadoras, cards_operadoras):
                    with col_operadora:
                        st.markdown(
                            f"""
                            <div class="kpi-card neutral" style="margin-bottom:1rem">
                                <div class="kpi-label">{texto_html(item["operadora"], 32)}</div>
                                <div class="kpi-value green">{fmt_brl(item["liquido"])}</div>
                                <div class="kpi-footer">
                                    Bruto {fmt_brl(item["bruto"])} · taxa {item["taxa_media"]:.2f}%
                                </div>
                            </div>
                            """,
                            unsafe_allow_html=True,
                        )

            rows = ""
            for _, row in df_maquininhas_view.sort_values(
                "data", ascending=False, na_position="last"
            ).iterrows():
                data_str = (
                    row["data"].strftime("%d/%m/%y")
                    if pd.notna(row["data"]) else "—"
                )
                descricao = texto_html(row.get("memo", ""), 80)
                classe = "valor-pos" if row["valor"] > 0 else "valor-neg"
                taxa_linha = pd.to_numeric(
                    pd.Series([row.get("taxa_maquininha", 0)]),
                    errors="coerce",
                ).fillna(0).iloc[0]
                valor_liquido_linha = pd.to_numeric(
                    pd.Series([row.get("valor", 0)]),
                    errors="coerce",
                ).fillna(0).abs().iloc[0]
                valor_bruto_linha = valor_liquido_linha + max(taxa_linha, 0)
                percentual_taxa_linha = (
                    taxa_linha / valor_bruto_linha * 100
                    if valor_bruto_linha > 0 and taxa_linha > 0 else 0.0
                )
                rows += (
                    f"<tr><td>{data_str}</td>"
                    f"<td>{texto_html(row.get('hora', ''))}</td>"
                    f"<td>{texto_html(row.get('maquininha_operadora', 'Maquininha'), 24)}</td>"
                    f"<td>{texto_html(row.get('ponto_venda', ''), 12) or '—'}</td>"
                    f"<td>{descricao}</td>"
                    f"<td>{tag_categoria(row.get('fluxo_infinity', ''))}</td>"
                    f"<td>{texto_html(row.get('bandeira', ''), 18) or '—'}</td>"
                    f"<td>{texto_html(row.get('tipo_transacao', ''), 22) or '—'}</td>"
                    f"<td>{texto_html(row.get('parcelas', ''), 10) or '—'}</td>"
                    f"<td class='valor-pos'>{fmt_brl(valor_bruto_linha) if valor_bruto_linha > 0 else '—'}</td>"
                    f"<td>{percentual_taxa_linha:.2f}%</td>"
                    f"<td class='valor-neg'>{fmt_brl_saida(taxa_linha) if taxa_linha > 0 else '—'}</td>"
                    f"<td class='{classe}'>{fmt_brl_saida(row['valor']) if row['valor'] < 0 else fmt_brl(row['valor'])}</td></tr>"
                )
            st.markdown(
                '<table class="fin-table"><thead><tr>'
                '<th>Data</th><th>Hora</th><th>Maquininha</th><th>PDV</th><th>Descrição</th>'
                '<th>Fluxo</th><th>Bandeira</th><th>Tipo</th><th>Parcelas</th>'
                '<th>Bruto</th><th>% Taxa</th><th>Taxa</th><th>Líquido</th></tr></thead>'
                f'<tbody>{rows}</tbody></table>',
                unsafe_allow_html=True,
            )

            if df_clinipay_base_periodo is not None and not df_clinipay_base_periodo.empty:
                df_clinipay_view = df_clinipay_base_periodo.copy()
                resumo_clinipay = calcular_taxas_clinipay(df_clinipay_view)
                st.markdown(
                    """
                    <div style="font-size:0.7rem;font-weight:700;letter-spacing:0.07em;
                                color:#4A4A7A;text-transform:uppercase;margin:1.4rem 0 0.8rem">
                        Boletos e Clinipay
                    </div>
                    """,
                    unsafe_allow_html=True,
                )
                bc1, bc2, bc3, bc4, bc5 = st.columns(5)
                with bc1:
                    st.markdown(
                        f"""
                        <div class="kpi-card green" style="margin-bottom:1rem">
                            <div class="kpi-label">Bruto boletos/Clinipay</div>
                            <div class="kpi-value green">{fmt_brl(resumo_clinipay["bruto"])}</div>
                            <div class="kpi-footer">Valor original liquidado</div>
                        </div>
                        """,
                        unsafe_allow_html=True,
                    )
                with bc2:
                    st.markdown(
                        f"""
                        <div class="kpi-card red" style="margin-bottom:1rem">
                            <div class="kpi-label">Custo de liquidação</div>
                            <div class="kpi-value red">{fmt_brl_saida(resumo_clinipay["taxas"])}</div>
                            <div class="kpi-footer">Taxa média {resumo_clinipay["taxa_media"]:.2f}%</div>
                        </div>
                        """,
                        unsafe_allow_html=True,
                    )
                with bc3:
                    st.markdown(
                        f"""
                        <div class="kpi-card neutral" style="margin-bottom:1rem">
                            <div class="kpi-label">Juros / multa</div>
                            <div class="kpi-value blue">{fmt_brl(resumo_clinipay["juros"])}</div>
                            <div class="kpi-footer">Juros médio calculado {resumo_clinipay["juros_medio"]:.2f}%</div>
                        </div>
                        """,
                        unsafe_allow_html=True,
                    )
                with bc4:
                    st.markdown(
                        f"""
                        <div class="kpi-card blue" style="margin-bottom:1rem">
                            <div class="kpi-label">Líquido recebido</div>
                            <div class="kpi-value blue">{fmt_brl(resumo_clinipay["liquido"])}</div>
                            <div class="kpi-footer">Bruto menos taxas do Clinipay</div>
                        </div>
                        """,
                        unsafe_allow_html=True,
                    )
                with bc5:
                    st.markdown(
                        f"""
                        <div class="kpi-card purple" style="margin-bottom:1rem">
                            <div class="kpi-label">Lançamentos</div>
                            <div class="kpi-value">{len(df_clinipay_view)}</div>
                            <div class="kpi-footer">Boletos, boletopix e Clinipay</div>
                        </div>
                        """,
                        unsafe_allow_html=True,
                    )

                rows_clinipay = ""
                for _, row in df_clinipay_view.sort_values(
                    "data", ascending=False, na_position="last"
                ).iterrows():
                    data_str = (
                        row["data"].strftime("%d/%m/%y")
                        if pd.notna(row["data"]) else "—"
                    )
                    bruto_linha = pd.to_numeric(
                        pd.Series([
                            row.get("valor_bruto_clinipay", row.get("valor", 0))
                        ]),
                        errors="coerce",
                    ).fillna(0).abs().iloc[0]
                    taxa_linha = pd.to_numeric(
                        pd.Series([row.get("taxa_clinipay", 0)]),
                        errors="coerce",
                    ).fillna(0).abs().iloc[0]
                    juros_linha = pd.to_numeric(
                        pd.Series([row.get("juros_clinipay", 0)]),
                        errors="coerce",
                    ).fillna(0).abs().iloc[0]
                    liquido_linha = pd.to_numeric(
                        pd.Series([row.get("valor", 0)]),
                        errors="coerce",
                    ).fillna(0).abs().iloc[0]
                    perc_taxa = (
                        taxa_linha / bruto_linha * 100
                        if bruto_linha > 0 and taxa_linha > 0 else 0.0
                    )
                    base_juros = bruto_linha - juros_linha
                    perc_juros = (
                        juros_linha / base_juros * 100
                        if base_juros > 0 and juros_linha > 0 else 0.0
                    )
                    rows_clinipay += (
                        f"<tr><td>{data_str}</td>"
                        f"<td>{texto_html(row.get('forma_clinipay', 'Clinipay'), 24)}</td>"
                        f"<td>{texto_html(row.get('memo', ''), 72)}</td>"
                        f"<td>{texto_html(row.get('parcela_clinipay', ''), 10) or '—'}</td>"
                        f"<td class='valor-pos'>{fmt_brl(bruto_linha)}</td>"
                        f"<td>{perc_taxa:.2f}%</td>"
                        f"<td class='valor-neg'>{fmt_brl_saida(taxa_linha) if taxa_linha > 0 else '—'}</td>"
                        f"<td>{perc_juros:.2f}%</td>"
                        f"<td class='valor-pos'>{fmt_brl(juros_linha) if juros_linha > 0 else '—'}</td>"
                        f"<td class='valor-pos'>{fmt_brl(liquido_linha)}</td></tr>"
                    )
                st.markdown(
                    '<table class="fin-table"><thead><tr>'
                    '<th>Data</th><th>Forma</th><th>Descrição</th><th>Parcela</th>'
                    '<th>Bruto</th><th>% Taxa</th><th>Taxa</th>'
                    '<th>% Juros</th><th>Juros</th><th>Líquido</th></tr></thead>'
                    f'<tbody>{rows_clinipay}</tbody></table>',
                    unsafe_allow_html=True,
                )

# =========================
# PÁGINA: SALDO BANCÁRIO
# =========================
elif st.session_state.pagina == "saldo":
    col_title, col_btn = st.columns([5, 1])
    with col_title:
        st.markdown('<div class="page-title">Saldo Bancário</div>', unsafe_allow_html=True)
        subtitulo_saldo = (
            "Acompanhe os saldos incluídos neste relatório."
            if st.session_state.share_mode
            else "Gerencie suas contas e acompanhe a variação de saldo."
        )
        st.markdown(
            f'<div class="page-subtitle">{subtitulo_saldo}</div>',
            unsafe_allow_html=True,
        )
        st.markdown(
            f"""
            <div class="period-summary">
                <span><strong>Período selecionado</strong> · {periodo_label}</span>
                <span>O filtro é compartilhado com a Visão Financeira</span>
            </div>
            """,
            unsafe_allow_html=True,
        )
    with col_btn:
        st.markdown("<br>", unsafe_allow_html=True)
        if (
            not st.session_state.share_mode
            and st.button("➕  Nova conta", type="primary")
        ):
            st.session_state.mostrar_nova_conta = not st.session_state.mostrar_nova_conta

    if st.session_state.mostrar_nova_conta and not st.session_state.share_mode:
        with st.container():
            st.markdown("""
            <div style="background:#11112A;border:1px solid #1E1E3A;border-radius:12px;
                        padding:1.2rem 1.4rem;margin-bottom:1rem">
                <div style="font-size:0.85rem;font-weight:600;color:#E2E8F0;margin-bottom:0.8rem">
                    Nova conta ou investimento</div>
            </div>""", unsafe_allow_html=True)
            nc1, nc2 = st.columns(2)
            nc_tipo, nc3, nc4 = st.columns([1.2, 1, 1])
            with nc1: nova_nome   = st.text_input("Nome da conta", placeholder="Ex: Conta Corrente")
            with nc2: nova_banco  = st.text_input("Instituição",         placeholder="Ex: Sicoob")
            with nc_tipo:
                novo_tipo_conta = st.selectbox(
                    "Tipo",
                    ["Conta bancária", "Investimento"],
                )
            with nc3: nova_inicial = st.number_input("Saldo Inicial (R$)", value=0.0, format="%.2f")
            with nc4: nova_final   = st.number_input("Saldo Final (R$)", value=0.0, format="%.2f")
            cb1, cb2, _ = st.columns([1, 1, 4])
            with cb1:
                if st.button("Salvar conta", type="primary"):
                    if not nova_nome or not nova_banco:
                        st.error("❌ Preencha nome e banco.")
                    else:
                        try:
                            add_conta(
                                user.id,
                                nova_nome,
                                nova_banco,
                                nova_inicial,
                                nova_final,
                                novo_tipo_conta,
                            )
                            st.session_state.mostrar_nova_conta = False
                            st.success("✅ Conta adicionada!")
                            st.rerun()
                        except Exception as e:
                            st.error(f"❌ Erro: {str(e)}")
            with cb2:
                if st.button("Cancelar"):
                    st.session_state.mostrar_nova_conta = False
                    st.rerun()

    if st.session_state.share_mode:
        contas = st.session_state.share_accounts
    else:
        try:
            contas = get_contas(user.id)
        except Exception as e:
            st.error(f"❌ Erro: {str(e)}")
            contas = []

    saldo_ofx_estimado = False
    if df_ofx is not None and not df_ofx.empty and periodo_valido:
        extrato_completo = df_ofx_movimentos.copy()
        extrato_completo["data"] = pd.to_datetime(
            extrato_completo["data"],
            errors="coerce",
        )
        extrato_completo = extrato_completo.dropna(subset=["data"])
        movimento_periodo = float(
            extrato_completo.loc[
                extrato_completo["data"].between(
                    inicio_periodo,
                    fim_periodo,
                    inclusive="both",
                ),
                "valor",
            ].sum()
        )
        linhas_saldo = df_ofx.copy()
        if "_tipo_conta" in linhas_saldo.columns:
            tipo_saldo = linhas_saldo["_tipo_conta"].fillna("").map(normalizar_texto)
            linhas_investimento = linhas_saldo[tipo_saldo == "investimento"].copy()
            linhas_saldo = linhas_saldo[tipo_saldo != "investimento"].copy()
        else:
            linhas_investimento = pd.DataFrame()
        if "_data_saldo_ofx" in linhas_saldo.columns:
            linhas_saldo["_ordem_saldo"] = pd.to_datetime(
                linhas_saldo["_data_saldo_ofx"],
                format="%Y%m%d",
                errors="coerce",
            )
            linhas_saldo = linhas_saldo.sort_values(
                "_ordem_saldo",
                ascending=False,
                na_position="last",
            )
        contas_importadas = []

        def chave_conta_ofx(df: pd.DataFrame) -> pd.Series:
            banco = (
                df["_banco_ofx"] if "_banco_ofx" in df.columns
                else pd.Series("Banco", index=df.index)
            ).fillna("Banco").astype(str).str.strip()
            conta = (
                df["_conta_ofx"] if "_conta_ofx" in df.columns
                else pd.Series("", index=df.index)
            ).fillna("").astype(str).str.strip()
            arquivo = (
                df["_arquivo_origem"] if "_arquivo_origem" in df.columns
                else pd.Series("", index=df.index)
            ).fillna("").astype(str).str.strip()
            conta_base = conta.where(
                conta.ne("") & conta.map(normalizar_texto).ne("conta ofx"),
                arquivo,
            )
            conta_base = conta_base.where(conta_base.ne(""), "Conta OFX")
            return banco + "||" + conta_base

        if not linhas_saldo.empty:
            linhas_saldo["_grupo_conta_ofx"] = chave_conta_ofx(linhas_saldo)
            extrato_completo["_grupo_conta_ofx"] = chave_conta_ofx(extrato_completo)

            for idx_grupo, (grupo_conta, linhas_conta) in enumerate(
                linhas_saldo.groupby("_grupo_conta_ofx", sort=False)
            ):
                primeira_linha = linhas_conta.iloc[0]
                extrato_conta = extrato_completo[
                    extrato_completo["_grupo_conta_ofx"] == grupo_conta
                ]
                movimento_periodo_conta = float(
                    extrato_conta.loc[
                        extrato_conta["data"].between(
                            inicio_periodo,
                            fim_periodo,
                            inclusive="both",
                        ),
                        "valor",
                    ].sum()
                )
                saldo_oficial = primeira_linha.get("_saldo_ofx")
                tem_saldo_oficial = pd.notna(saldo_oficial)
                if tem_saldo_oficial:
                    saldo_final_extrato = float(saldo_oficial)
                    movimentos_desde_inicio = float(
                        extrato_conta.loc[
                            extrato_conta["data"] >= inicio_periodo,
                            "valor",
                        ].sum()
                    )
                    movimentos_apos_fim = float(
                        extrato_conta.loc[
                            extrato_conta["data"] > fim_periodo,
                            "valor",
                        ].sum()
                    )
                    saldo_inicial_ofx = saldo_final_extrato - movimentos_desde_inicio
                    saldo_final_ofx = saldo_final_extrato - movimentos_apos_fim
                else:
                    saldo_inicial_ofx = float(
                        extrato_conta.loc[
                            extrato_conta["data"] < inicio_periodo,
                            "valor",
                        ].sum()
                    )
                    saldo_final_ofx = saldo_inicial_ofx + movimento_periodo_conta
                conta_ofx = str(
                    primeira_linha.get("_conta_ofx", "Conta OFX") or "Conta OFX"
                ).strip()
                banco_ofx = nome_banco_brasileiro(str(
                    primeira_linha.get("_banco_ofx", "Banco") or "Banco"
                ).strip())
                arquivo_ofx = str(
                    primeira_linha.get("_arquivo_origem", "") or ""
                ).strip()
                if conta_ofx and normalizar_texto(conta_ofx) != "conta ofx":
                    sufixo_conta = conta_ofx[-4:] if len(conta_ofx) > 4 else conta_ofx
                    nome_conta_ofx = f"{banco_ofx} •••• {sufixo_conta}"
                else:
                    nome_conta_ofx = banco_ofx if banco_ofx != "Banco" else "Conta OFX"
                contas_importadas.append([
                    f"ofx_{idx_grupo}",
                    nome_conta_ofx,
                    banco_ofx,
                    saldo_inicial_ofx,
                    saldo_final_ofx,
                    "Conta bancária",
                ])
                saldo_ofx_estimado = saldo_ofx_estimado or not tem_saldo_oficial

        if not linhas_investimento.empty:
            investimentos_validos = linhas_investimento.dropna(
                subset=["_saldo_ofx"]
            )
            for idx, linha_invest in investimentos_validos.iterrows():
                saldo_invest = float(linha_invest.get("_saldo_ofx", 0.0))
                nome_invest = str(
                    linha_invest.get("_conta_ofx", "Investimentos")
                    or "Investimentos"
                )
                banco_invest = nome_banco_brasileiro(str(
                    linha_invest.get("_banco_ofx", "Banco") or "Banco"
                ))
                contas_importadas.append([
                    f"invest_excel_{idx}",
                    nome_invest,
                    banco_invest,
                    saldo_invest,
                    saldo_invest,
                    "Investimento",
                ])

        contas = list(contas) + contas_importadas

    if saldo_ofx_estimado:
        st.info(
            "O arquivo OFX não informou o saldo bancário oficial. Para não deixar "
            "o quadro vazio, os saldos são estimados pela movimentação acumulada "
            "do arquivo até as datas selecionadas."
        )

    if contas and any(c[0] != "ofx" for c in contas):
        st.caption(
            "Contas cadastradas manualmente não possuem lançamentos datados; por "
            "isso, seus valores permanecem fixos ao alterar o período."
        )

    bancos_disponiveis = sorted(
        {
            str(c[2] or "Banco")
            for c in contas
            if len(c) > 2 and str(c[2] or "").strip()
        }
    )
    filtro_banco_saldo = "Todos os bancos"
    if bancos_disponiveis:
        filtro_banco_saldo = st.selectbox(
            "Banco",
            ["Todos os bancos", *bancos_disponiveis],
            key="filtro_banco_saldo",
        )

    contas_filtradas = [
        c for c in contas
        if filtro_banco_saldo == "Todos os bancos"
        or str(c[2] or "Banco") == filtro_banco_saldo
    ]

    contas_bancarias = [
        c for c in contas_filtradas
        if normalizar_texto(tipo_conta_registro(c)) != "investimento"
    ]
    contas_investimento = [
        c for c in contas_filtradas
        if normalizar_texto(tipo_conta_registro(c)) == "investimento"
    ]
    saldo_banco_final = sum(c[4] for c in contas_bancarias) if contas_bancarias else 0.0
    saldo_invest_final = (
        sum(c[4] for c in contas_investimento) if contas_investimento else 0.0
    )
    total_inicial = sum(c[3] for c in contas_filtradas) if contas_filtradas else 0.0
    total_final   = saldo_banco_final + saldo_invest_final
    variacao      = total_final - total_inicial
    cor_var       = "green" if variacao >= 0 else "red"
    sinal_var     = "+" if variacao >= 0 else ""

    k1, k2, k3, k4 = st.columns(4)
    with k1:
        st.markdown(f'<div class="kpi-card green"><div class="kpi-label">Saldo em banco</div><div class="kpi-value">{fmt_brl(saldo_banco_final)}</div><div class="kpi-footer">Contas bancárias e OFX</div></div>', unsafe_allow_html=True)
    with k2:
        st.markdown(f'<div class="kpi-card blue"><div class="kpi-label">Investimentos</div><div class="kpi-value blue">{fmt_brl(saldo_invest_final)}</div><div class="kpi-footer">Aplicações e reservas</div></div>', unsafe_allow_html=True)
    with k3:
        st.markdown(f'<div class="kpi-card purple"><div class="kpi-label">Saldo total</div><div class="kpi-value">{fmt_brl(total_final)}</div><div class="kpi-footer">Banco + investimentos</div></div>', unsafe_allow_html=True)
    with k4:
        st.markdown(f'<div class="kpi-card neutral"><div class="kpi-label">Variação do Período</div><div class="kpi-value {cor_var}">{sinal_var}{fmt_brl(variacao)}</div><div class="kpi-footer">Final – Inicial</div></div>', unsafe_allow_html=True)

    st.markdown("<br>", unsafe_allow_html=True)

    if not contas_filtradas:
        st.markdown('<table class="fin-table"><thead><tr><th>Conta</th><th>Instituição</th><th>Tipo</th><th>Saldo Inicial</th><th>Saldo Final</th><th>Variação</th></tr></thead><tbody><tr class="empty-row"><td colspan="6">Nenhuma conta cadastrada.</td></tr></tbody></table>', unsafe_allow_html=True)
    else:
        rows = ""
        for c in contas_filtradas:
            cid, nome, banco, s_ini, s_fin = c[:5]
            tipo_conta = tipo_conta_registro(c)
            var     = s_fin - s_ini
            cls_var = "variacao-pos" if var >= 0 else "variacao-neg"
            sinal_c = "+" if var >= 0 else ""
            nome_seguro = html.escape(str(nome))
            banco_seguro = html.escape(str(banco))
            rows += f"<tr><td><strong style='color:#E2E8F0'>{nome_seguro}</strong></td><td>{banco_seguro}</td><td>{tag_categoria(tipo_conta)}</td><td class='valor-pos'>{fmt_brl(s_ini)}</td><td class='valor-pos'>{fmt_brl(s_fin)}</td><td class='{cls_var}'>{sinal_c}{fmt_brl(var)}</td></tr>"
        st.markdown(f'<table class="fin-table"><thead><tr><th>Conta</th><th>Instituição</th><th>Tipo</th><th>Saldo Inicial</th><th>Saldo Final</th><th>Variação</th></tr></thead><tbody>{rows}</tbody></table>', unsafe_allow_html=True)
